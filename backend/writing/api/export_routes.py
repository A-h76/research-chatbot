"""Writing export routes extracted from server.py."""

from __future__ import annotations

import io
import json

from flask import Blueprint, jsonify, request, send_file, session


def create_writing_export_blueprint(
    *,
    SessionLocal,
    Note,
    UserFile,
    PaperAnalysis,
    Conversation,
    select_fn,
    login_required,
    WritingDocument=None,
    ReviewerRun=None,
    ReviewerFinding=None,
):
    bp = Blueprint("writing_export_routes", __name__)

    def _export_as_markdown(content_str: str, title: str = "") -> bytes:
        header = f"# {title}\n\n" if title else ""
        return (header + content_str).encode("utf-8")

    def _export_as_docx(content_str: str, title: str = "") -> bytes:
        try:
            from io import BytesIO

            import docx as _docx

            doc = _docx.Document()
            if title:
                doc.add_heading(title, 0)
            for para in content_str.split("\n\n"):
                if para.strip():
                    doc.add_paragraph(para.strip())
            buf = BytesIO()
            doc.save(buf)
            return buf.getvalue()
        except ImportError:
            return content_str.encode("utf-8")

    @bp.route("/api/export/notes", methods=["POST"])
    @login_required
    def export_notes():
        data = request.get_json(silent=True) or {}
        uid = session["user_id"]
        fmt = str(data.get("format") or "md").lower()
        project_id = data.get("project_id")
        note_ids = data.get("note_ids")

        db = SessionLocal()
        try:
            stmt = select_fn(Note).where(Note.user_id == uid)
            if project_id is not None:
                stmt = stmt.where(Note.project_id == project_id)
            if note_ids:
                stmt = stmt.where(Note.id.in_([int(i) for i in note_ids]))
            notes = db.execute(stmt.order_by(Note.updated_at.desc())).scalars().all()
        finally:
            db.close()

        sections = []
        for note in notes:
            if note.title:
                sections.append(f"## {note.title}\n\n{note.content or ''}")
            else:
                sections.append(note.content or "")
        body = "\n\n---\n\n".join(sections)

        if fmt == "docx":
            blob = _export_as_docx(body, "Research Notes")
            mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            fname = "notes.docx"
        elif fmt == "txt":
            blob = body.encode("utf-8")
            mime = "text/plain"
            fname = "notes.txt"
        else:
            blob = _export_as_markdown(body, "Research Notes")
            mime = "text/markdown"
            fname = "notes.md"

        return send_file(io.BytesIO(blob), mimetype=mime, as_attachment=True, download_name=fname)

    @bp.route("/api/export/analysis/<int:file_id>", methods=["GET"])
    @login_required
    def export_analysis(file_id):
        uid = session["user_id"]
        fmt = request.args.get("format", "md").lower()

        db = SessionLocal()
        try:
            uf = db.get(UserFile, file_id)
            if not uf or uf.user_id != uid:
                return jsonify({"error": "not_found"}), 404

            pa = db.execute(select_fn(PaperAnalysis).where(PaperAnalysis.file_id == file_id)).scalar_one_or_none()
            if not pa or pa.status != "done":
                return (
                    jsonify(
                        {
                            "error": "analysis_not_ready",
                            "detail": "Analysis not yet complete.",
                        }
                    ),
                    400,
                )

            analysis = json.loads(pa.data) if pa.data else {}
        finally:
            db.close()

        title = uf.title or uf.name
        lines = [f"# Paper Analysis: {title}", ""]
        field_labels = {
            "executive_summary": "Executive Summary",
            "abstract_explained": "Abstract Explained",
            "research_objective": "Research Objective",
            "problem_statement": "Problem Statement",
            "methodology": "Methodology",
            "dataset": "Dataset",
            "experiments": "Experiments",
            "results": "Results",
            "key_contributions": "Key Contributions",
            "strengths": "Strengths",
            "limitations": "Limitations",
            "future_work": "Future Work",
            "keywords": "Keywords",
            "important_terms": "Important Terms",
        }
        for key, label in field_labels.items():
            val = analysis.get(key)
            if not val:
                continue
            lines.append(f"## {label}")
            if isinstance(val, list):
                for item in val:
                    lines.append(f"- {item}")
            elif isinstance(val, dict):
                for term, defn in val.items():
                    lines.append(f"**{term}**: {defn}")
            else:
                lines.append(str(val))
            lines.append("")

        body = "\n".join(lines)

        if fmt == "docx":
            blob = _export_as_docx(body, f"Analysis: {title}")
            mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            fname = f"analysis-{file_id}.docx"
        elif fmt == "txt":
            blob = body.encode("utf-8")
            mime = "text/plain"
            fname = f"analysis-{file_id}.txt"
        else:
            blob = body.encode("utf-8")
            mime = "text/markdown"
            fname = f"analysis-{file_id}.md"

        return send_file(io.BytesIO(blob), mimetype=mime, as_attachment=True, download_name=fname)

    @bp.route("/api/export/chat/<int:cid>", methods=["GET"])
    @login_required
    def export_chat(cid):
        uid = session["user_id"]
        fmt = request.args.get("format", "md").lower()

        db = SessionLocal()
        try:
            convo = db.execute(
                select_fn(Conversation).where(
                    Conversation.id == cid,
                    Conversation.user_id == uid,
                )
            ).scalar_one_or_none()
            if not convo:
                return jsonify({"error": "not_found"}), 404
            messages = list(convo.messages)
        finally:
            db.close()

        lines = [f"# {convo.title or 'Conversation'}", ""]
        for m in messages:
            role = "**You**" if m.role == "user" else "**Assistant**"
            lines.append(f"{role}\n\n{m.content or ''}")
            lines.append("\n---\n")

        body = "\n".join(lines)

        if fmt == "txt":
            blob = body.encode("utf-8")
            mime = "text/plain"
            fname = f"chat-{cid}.txt"
        else:
            blob = body.encode("utf-8")
            mime = "text/markdown"
            fname = f"chat-{cid}.md"

        return send_file(io.BytesIO(blob), mimetype=mime, as_attachment=True, download_name=fname)

    @bp.route("/api/writing/documents/<int:document_id>/export", methods=["POST"])
    @login_required
    def export_grounded_lit_review(document_id: int):
        """Server-gated lit-review export (#18): MD (+ optional BibTeX) with Reviewer gate.

        Body: ``{ writing, format?, writing_version?, title?, body? }``
        ``format``: ``markdown`` (default) or ``markdown_bibtex``.
        """
        if WritingDocument is None:
            return jsonify({"error": "not_configured"}), 503

        from backend.evidence.writing.export_markdown import (
            build_bibtex_from_writing,
            build_literature_review_markdown,
            can_export_grounded_lit_review,
            compute_export_traceability,
            merge_persisted_review_into_writing,
        )
        from backend.evidence.writing.reviewer_persistence import serialize_run

        uid = session["user_id"]
        data = request.get_json(silent=True) or {}
        writing = data.get("writing")
        if not isinstance(writing, dict):
            return (
                jsonify(
                    {
                        "error": "writing_required",
                        "detail": "Body must include the grounded writing snapshot.",
                    }
                ),
                400,
            )

        fmt = str(data.get("format") or "markdown").strip().lower()
        want_bib = fmt in {"markdown_bibtex", "md_bib", "md+bib"}

        db = SessionLocal()
        try:
            doc = db.execute(
                select_fn(WritingDocument).where(
                    WritingDocument.id == int(document_id),
                    WritingDocument.user_id == uid,
                )
            ).scalar_one_or_none()
            if not doc:
                return jsonify({"error": "not_found"}), 404

            # Merge latest persisted Reviewer run when available (B-514).
            if ReviewerRun is not None and ReviewerFinding is not None:
                run = (
                    db.execute(
                        select_fn(ReviewerRun)
                        .where(
                            ReviewerRun.document_id == int(document_id),
                            ReviewerRun.user_id == uid,
                        )
                        .order_by(ReviewerRun.created_at.desc(), ReviewerRun.id.desc())
                        .limit(1)
                    )
                    .scalars()
                    .first()
                )
                if run is not None:
                    findings = (
                        db.execute(
                            select_fn(ReviewerFinding)
                            .where(ReviewerFinding.run_id == int(run.id))
                            .order_by(ReviewerFinding.id.asc())
                        )
                        .scalars()
                        .all()
                    )
                    payload = serialize_run(run, findings=list(findings))
                    writing = merge_persisted_review_into_writing(
                        writing, payload.get("review")
                    )

            ok, reason = can_export_grounded_lit_review(writing)
            if not ok:
                return (
                    jsonify(
                        {
                            "error": "export_blocked",
                            "detail": reason or "Export blocked by Research Reviewer",
                            "reason": reason,
                        }
                    ),
                    403,
                )

            title = str(data.get("title") or getattr(doc, "title", None) or "Literature review")
            body = str(data.get("body") if data.get("body") is not None else (doc.content or ""))
            if not body.strip():
                return (
                    jsonify(
                        {
                            "error": "empty_draft",
                            "detail": "Draft is empty — generate and insert a grounded review first",
                        }
                    ),
                    400,
                )

            writing_version = data.get("writing_version")
            md = build_literature_review_markdown(
                title=title,
                body=body,
                writing=writing,
                writing_version=str(writing_version) if writing_version else None,
            )
            bib = build_bibtex_from_writing(writing) if want_bib else ""
            safe = "".join(
                c if c.isalnum() or c in "-_" else "-"
                for c in title.lower().replace(" ", "-")
            ).strip("-")[:60] or "literature-review"
            base = f"{safe}-{document_id}"
            trace = compute_export_traceability(writing)
            return jsonify(
                {
                    "ok": True,
                    "markdown": md,
                    "bibtex": bib if (want_bib and bib.strip()) else None,
                    "filename_base": base,
                    "traceability": trace,
                    "document_id": document_id,
                }
            )
        finally:
            db.close()

    return bp
