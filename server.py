"""
Personal AI / Dhund — ChatGPT-style chatbot backend (Phase 1)
Flask + Google OAuth + Postgres/SQLite + OpenAI Responses API (streaming)
+ Projects + selective memory + auto titles + web search
+ File uploads (PDF/Word/image/text) + vision + RAG + citation manager.
"""

import base64
import binascii
import io
import json
import logging
import math
import os
import shutil
import threading
import time
import uuid
from datetime import datetime, timedelta, timezone
from functools import wraps
from urllib.parse import urlparse

from dotenv import load_dotenv

load_dotenv()

import click
from authlib.integrations.flask_client import OAuth
from flask import (
    Flask,
    Response,
    abort,
    g,
    jsonify,
    redirect,
    render_template,
    request,
    send_file,
    send_from_directory,
    session,
    url_for,
)
from openai import OpenAI
from sqlalchemy import (
    BigInteger,
    Boolean,
    Column,
    DateTime,
    Float,
    ForeignKey,
    Integer,
    String,
    Text,
    create_engine,
    delete,
    func,
    select,
)
from sqlalchemy import text as sqltext
from sqlalchemy.orm import declarative_base, relationship, sessionmaker

import storage
from quotas import QuotaService, create_usage_log_model, EntitlementService
from feature_flags import (
    FeatureFlagService,
    create_feature_flag_model,
    FLAG_DISCOVER_SEARCH,
    FLAG_WRITING_INTELLIGENCE,
)

# ------------------------------------------------------------------ config
OPENAI_API_KEY = os.environ.get("OPENAI_API_KEY", "")
DATABASE_URL = os.environ.get("DATABASE_URL") or "sqlite:///chat_dev.db"
if DATABASE_URL.startswith("postgres://"):  # Neon compatibility
    DATABASE_URL = DATABASE_URL.replace("postgres://", "postgresql://", 1)

GOOGLE_CLIENT_ID = os.environ.get("GOOGLE_CLIENT_ID", "")
GOOGLE_CLIENT_SECRET = os.environ.get("GOOGLE_CLIENT_SECRET", "")
# DEV_AUTO_LOGIN: set to any non-empty string to skip Google OAuth in development.
# When set, visiting /login automatically signs in as a local dev user.
# NEVER set this in production.
DEV_AUTO_LOGIN = os.environ.get("DEV_AUTO_LOGIN", "")
ALLOWED_EMAILS = [e.strip().lower() for e in os.environ.get("ALLOWED_EMAILS", "").split(",") if e.strip()]
# Optional lockdown: when true, only allowlisted or invited emails may sign up.
# Open signup is the default for production.
BETA_INVITE_ONLY = (os.environ.get("BETA_INVITE_ONLY", "") or "").strip().lower() in {
    "1",
    "true",
    "yes",
    "on",
}
CLOSED_BETA = BETA_INVITE_ONLY or (
    os.environ.get("CLOSED_BETA", "").strip().lower() in {"1", "true", "yes", "on"}
)

# Defaults kept to models with confident, verified pricing (see
# backend/ai/cost_ledger.py's PRICING table and its own note on why
# gpt-5-family is deliberately excluded there) — not a claim that gpt-5
# doesn't exist or can't be used; a user can still pick it manually from
# the live-fetched dropdown. Only what's used automatically changed.
DEFAULT_MODEL = os.environ.get("DEFAULT_MODEL", "gpt-5-mini")
UTILITY_MODEL = os.environ.get("UTILITY_MODEL", "gpt-5-mini")
EMBED_MODEL = os.environ.get("EMBED_MODEL", "text-embedding-3-small")
FALLBACK_MODELS = [
    m.strip() for m in os.environ.get("MODELS", "gpt-5.5,gpt-5-mini,gpt-5.5-pro").split(",") if m.strip()
]

UPLOAD_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "uploads")
os.makedirs(UPLOAD_DIR, exist_ok=True)  # only used for throwaway temp files now
MAX_FILE_MB = int(os.environ.get("MAX_FILE_MB", "25"))

# Optional job-status cache (database-design.md §5's job:{id}:status key).
# Never the source of truth — Postgres's upload_jobs row always is; a
# missing/unreachable Redis just means every read falls through to it.
REDIS_URL = os.environ.get("REDIS_URL", "")
JOB_STATUS_CACHE_TTL_SECONDS = 3600

# Transactional email — single source of truth in backend.config.email
from backend.config.email import (
    APP_BASE_URL,
    AUTH_EMAIL_FROM,
    EMAIL_FROM,
    NOREPLY_EMAIL_FROM,
    NOTIFICATIONS_EMAIL_FROM,
    PUBLIC_SITE_URL,
    RESEND_API_KEY,
    SUPPORT_EMAIL,
)

IS_PRODUCTION = (
    os.environ.get("FLASK_ENV", "").lower() == "production" or os.environ.get("APP_ENV", "").lower() == "production"
)
# worker.py's default poll interval is 2s (WORKER_POLL_INTERVAL) and it
# heartbeats every iteration — 60s is ~30 missed cycles, generous enough
# that a normal GC pause or slow job doesn't false-positive as "down".
WORKER_HEALTH_THRESHOLD_SECONDS = int(os.environ.get("WORKER_HEALTH_THRESHOLD_SECONDS", "60"))

# PR1: refuse production boot without explicit secrets (never silent random keys).
from security.startup import require_production_secrets, resolve_flask_secret_key, resolve_limiter_storage_uri
from security.authz import project_owned_by_user, resolve_owned_project_id
from security.metrics_access import check_metrics_access
from security.headers import apply_security_headers
from security.session_ttl import (
    enforce_session_ttl,
    mark_session_login,
    session_absolute_seconds,
)

require_production_secrets(os.environ, is_production=IS_PRODUCTION)

app = Flask(__name__)
app.secret_key = resolve_flask_secret_key(os.environ, is_production=IS_PRODUCTION)

# Optional Sentry — no-op unless SENTRY_DSN is set (see docs/security-baseline-v1-deploy-checklist.md).
from security.sentry_init import init_sentry

init_sentry(os.environ, flask_app=app)


@app.route("/api/health")
def api_health():
    """Liveness — no DB. Used by Railway/Cloudflare probes; must stay cheap."""
    return jsonify({"ok": True, "service": "web"}), 200


# Railway/Render terminate TLS at the edge and forward X-Forwarded-Proto/Host.
# Without ProxyFix, url_for(_external=True) and OAuth redirects become http://…
from werkzeug.middleware.proxy_fix import ProxyFix

app.wsgi_app = ProxyFix(app.wsgi_app, x_for=1, x_proto=1, x_host=1, x_prefix=1)
app.config["MAX_CONTENT_LENGTH"] = MAX_FILE_MB * 1024 * 1024
# Secure session-cookie defaults (Secure flag only in production/HTTPS).
# Absolute cookie lifetime matches SESSION_ABSOLUTE_HOURS; idle is enforced
# in enforce_session_ttl (PR4).
_abs_secs = session_absolute_seconds(os.environ) or (12 * 3600)
app.config.update(
    SESSION_COOKIE_HTTPONLY=True,
    SESSION_COOKIE_SAMESITE="Lax",
    SESSION_COOKIE_SECURE=IS_PRODUCTION,
    PERMANENT_SESSION_LIFETIME=timedelta(seconds=_abs_secs),
    SESSION_REFRESH_EACH_REQUEST=False,  # idle sliding is handled explicitly
)
client = OpenAI(api_key=OPENAI_API_KEY)

# ------------------------------------------------------------------ JWT (auth/ package)
# Extension alongside session/OAuth login, not a replacement — the
# existing @login_required + session["user_id"] path is untouched. This
# is for future API/programmatic clients that can't hold a browser
# session cookie. JWT_SECRET_KEY defaults to the same secret as Flask
# sessions when unset — production already required FLASK_SECRET_KEY above.
app.config.update(
    JWT_SECRET_KEY=os.environ.get("JWT_SECRET_KEY", app.secret_key),
    JWT_ACCESS_TOKEN_EXPIRES=timedelta(minutes=int(os.environ.get("JWT_ACCESS_TOKEN_EXPIRES_MIN", "15"))),
    JWT_REFRESH_TOKEN_EXPIRES=timedelta(days=int(os.environ.get("JWT_REFRESH_TOKEN_EXPIRES_DAYS", "30"))),
    JWT_ALGORITHM="HS256",
)
from flask_jwt_extended import JWTManager

jwt_manager = JWTManager(app)
from auth import JWTError, create_get_current_user, create_jwt, decode_jwt
from auth.decorators import set_jwt_session_version_checker
from auth.jwt_utils import session_version_matches

# ------------------------------------------------------------------ logging
from observability import (
    HTTP_REQUEST_DURATION_SECONDS,
    HTTP_REQUESTS_TOTAL,
    UPLOAD_QUEUE_LENGTH,
    configure_logging,
    correlation_id_var,
    record_ai_call,
    render_metrics,
)

configure_logging()
security_log = logging.getLogger("security")
email_log = logging.getLogger("email")


def log_security_event(event, **fields):
    """Structured audit trail for security-relevant actions."""
    detail = " ".join(f"{k}={v}" for k, v in fields.items())
    security_log.info("event=%s %s", event, detail)


# ------------------------------------------------------------------ email service
from backend.services.email import EmailEvent, TransactionalEmailService

# Keep EmailService name for older call sites / tests.
EmailService = TransactionalEmailService


def _html_to_text(html):
    from backend.services.email.renderer import html_to_text

    return html_to_text(html)


def _redact_email_secrets(body: str) -> str:
    """Strip magic-link / reset tokens from console email logs (Phase 3)."""
    import re

    text = body or ""
    text = re.sub(r"(token=)[^&\s\"'<>]+", r"\1[REDACTED]", text, flags=re.I)
    text = re.sub(r"(/auth/magic-link\?[^\s\"'<>]+)", "/auth/magic-link?[REDACTED]", text, flags=re.I)
    return text


email_service = TransactionalEmailService.from_env()


# ------------------------------------------------------------------ rate limiting + CSRF
from flask_limiter import Limiter
from flask_limiter.errors import RateLimitExceeded
from flask_limiter.util import get_remote_address

# Opt-in per-route limits. Redis when REDIS_URL is reachable; memory:// otherwise
# (shared limits across workers require Redis — see security/startup.py).
_LIMITER_STORAGE = resolve_limiter_storage_uri(REDIS_URL, is_production=IS_PRODUCTION)
limiter = Limiter(
    key_func=get_remote_address,
    app=app,
    default_limits=[],
    storage_uri=_LIMITER_STORAGE,
    strategy="fixed-window",
)


@app.errorhandler(RateLimitExceeded)
def _rate_limit_exceeded(e):
    uid = session.get("user_id")
    log_security_event(
        "rate_limit_exceeded",
        path=request.path,
        method=request.method,
        user_id=uid or "",
        remote=get_remote_address(),
        description=getattr(e, "description", "") or str(e),
    )
    return (
        jsonify(
            {
                "error": "rate_limit_exceeded",
                "message": "Too many requests. Please try again later.",
            }
        ),
        429,
    )

SAFE_METHODS = {"GET", "HEAD", "OPTIONS"}

# `npm run dev` (Vite) proxies /api/* to this process but the browser's own
# Origin header on state-changing requests reflects Vite's origin, not this
# one — request.host/APP_BASE_URL are always localhost:5000. Without this,
# every POST/PUT/PATCH/DELETE from the SPA in local dev 403s.
# Vite falls back to 5174, 5175, … when 5173 is already taken, so allow a
# small localhost port range. Gated on non-production so prod (Flask serves
# the built frontend, genuinely same-origin) never widens this check.
# Note: do not gate only on DEV_AUTO_LOGIN — local Vite + Google/session
# login is a common setup with DEV_AUTO_LOGIN unset.
_DEV_VITE_PORTS = range(5173, 5183)
DEV_FRONTEND_ORIGINS = (
    {f"localhost:{p}" for p in _DEV_VITE_PORTS} | {f"127.0.0.1:{p}" for p in _DEV_VITE_PORTS}
    if not IS_PRODUCTION
    else set()
)


@app.before_request
def _enforce_session_ttl():
    """Idle + absolute session timeout (PR4)."""
    if "user_id" not in session:
        return
    expired, reason = enforce_session_ttl(session, environ=os.environ)
    if not expired:
        return
    session.clear()
    log_security_event("session_expired", reason=reason, path=request.path)
    if request.path.startswith("/api/"):
        return (
            jsonify(
                {
                    "error": "session_expired",
                    "reason": reason,
                    "message": "Your session has expired. Please sign in again.",
                }
            ),
            401,
        )
    return redirect(url_for("login_page"))


@app.before_request
def _enforce_session_version():
    """Logout-all: bump User.session_version to invalidate other cookies."""
    uid = session.get("user_id")
    if not uid:
        return
    # Ops models may not be ready during very early import hooks — skip safely.
    try:
        db = SessionLocal()
        try:
            user = db.get(User, uid)
            if not user:
                return
            current = int(getattr(user, "session_version", 0) or 0)
            stamped = session.get("session_version")
            if stamped is None:
                session["session_version"] = current
                return
            if int(stamped) != current:
                session.clear()
                log_security_event("session_expired", reason="revoked", user_id=uid)
                if request.path.startswith("/api/"):
                    return jsonify({"error": "session_expired", "reason": "revoked"}), 401
                return redirect(url_for("login_page"))
        finally:
            db.close()
    except Exception:
        return


@app.before_request
def csrf_protect():
    """Same-origin check for state-changing /api/* and /auth/* calls —
    defense-in-depth on top of SameSite=Lax cookies.
    Non-browser clients (no Origin/Referer) pass."""
    if request.method in SAFE_METHODS:
        return
    path = request.path or ""
    if not (path.startswith("/api/") or path.startswith("/auth/")):
        return
    src = request.headers.get("Origin") or request.headers.get("Referer")
    if not src:
        return
    src_host = urlparse(src).netloc
    allowed = {request.host, urlparse(APP_BASE_URL).netloc} | DEV_FRONTEND_ORIGINS
    if src_host not in allowed:
        log_security_event("csrf_blocked", path=request.path, origin=src_host)
        return jsonify({"error": "csrf_origin_mismatch"}), 403


_request_log = logging.getLogger("request")


@app.before_request
def _start_request_observability():
    """Correlation id + start time for the logging/metrics after_request
    hook below. Accepts an inbound X-Request-ID so a request can be
    traced across a reverse proxy / calling service, mints one otherwise."""
    g.request_id = request.headers.get("X-Request-ID") or uuid.uuid4().hex
    correlation_id_var.set(g.request_id)
    g._request_start = time.monotonic()


@app.after_request
def _finish_request_observability(response):
    response.headers["X-Request-ID"] = getattr(g, "request_id", "")
    # url_rule.rule is the route TEMPLATE ("/api/files/<int:fid>"), not
    # the resolved path — using the resolved path would put every
    # distinct file/conversation/etc. id in its own metric label, an
    # unbounded cardinality that never stops growing.
    route = request.url_rule.rule if request.url_rule else "unmatched"
    duration = time.monotonic() - getattr(g, "_request_start", time.monotonic())
    HTTP_REQUESTS_TOTAL.labels(method=request.method, route=route, status=response.status_code).inc()
    HTTP_REQUEST_DURATION_SECONDS.labels(method=request.method, route=route).observe(duration)
    _request_log.info(
        "request",
        extra={
            "method": request.method,
            "route": route,
            "status": response.status_code,
            "duration_ms": round(duration * 1000, 1),
        },
    )
    apply_security_headers(
        response,
        is_production=IS_PRODUCTION,
        environ=os.environ,
        request_path=request.path,
    )
    return response


# ------------------------------------------------------------------ dynamic model list
_EXCLUDE = (
    "embedding",
    "whisper",
    "tts",
    "audio",
    "realtime",
    "image",
    "sora",
    "moderation",
    "transcribe",
    "davinci",
    "babbage",
    "instruct",
    "dall-e",
)
_INCLUDE_PREFIX = ("gpt-", "chatgpt-", "o1", "o3", "o4", "chat-latest")
_model_cache = {"ts": 0.0, "models": []}
_model_lock = threading.Lock()

# Model-capability guards for the temperature / reasoning-effort controls.
# Reasoning models (o-series + gpt-5 family) reject `temperature` — OpenAI
# returns 400 if it is sent. Use `reasoning_effort` instead for those models.
# gpt-5.5-pro (and similar Pro SKUs) do not support Responses streaming —
# chat must use a non-stream call and synthesize SSE for the client.
# Re-verify against current OpenAI docs if these ever misbehave — single named
# constants so it's a one-line fix.
REASONING_EFFORT_PREFIXES = ("o1", "o3", "o4", "gpt-5")
NO_TEMPERATURE_PREFIXES = ("o1", "o3", "o4", "gpt-5")
NO_STREAMING_PREFIXES = ("gpt-5.5-pro", "gpt-5.4-pro")
# Pro reasoning SKUs reject ``low``; OpenAI allows medium | high | xhigh.
PRO_REASONING_EFFORTS = frozenset({"medium", "high", "xhigh"})


def supports_reasoning_effort(model):
    return model.startswith(REASONING_EFFORT_PREFIXES)


def supports_temperature(model):
    return not model.startswith(NO_TEMPERATURE_PREFIXES)


def supports_streaming(model):
    return not (model or "").startswith(NO_STREAMING_PREFIXES)


def normalize_reasoning_effort(model, effort):
    """Clamp effort to values the selected model accepts; None means default."""
    if not effort:
        return None
    e = str(effort).strip().lower()
    if (model or "").startswith(NO_STREAMING_PREFIXES) or (model or "").endswith("-pro"):
        # Only clamp known Pro aliases that reject ``low``.
        if (model or "").startswith(("gpt-5.5-pro", "gpt-5.4-pro", "gpt-5.2-pro")):
            if e == "low" or e not in PRO_REASONING_EFFORTS:
                return "medium"
            return e
    if e not in ("low", "medium", "high", "xhigh"):
        return None
    return e


def _responses_output_text(resp) -> str:
    """Best-effort plain text from a non-stream Responses API result."""
    direct = getattr(resp, "output_text", None)
    if isinstance(direct, str) and direct:
        return direct
    parts: list[str] = []
    for item in getattr(resp, "output", None) or []:
        if getattr(item, "type", "") != "message":
            continue
        for block in getattr(item, "content", None) or []:
            if getattr(block, "type", "") in ("output_text", "text"):
                t = getattr(block, "text", None)
                if isinstance(t, str) and t:
                    parts.append(t)
    return "".join(parts)


def get_models(force=False):
    with _model_lock:
        if not force and _model_cache["models"] and time.time() - _model_cache["ts"] < 600:
            return _model_cache["models"]
        try:
            raw = client.models.list().data
            models = sorted(
                (m.id for m in raw if m.id.startswith(_INCLUDE_PREFIX) and not any(x in m.id for x in _EXCLUDE)),
                key=lambda mid: next((-m.created for m in raw if m.id == mid), 0),
            )
            if models:
                _model_cache.update(ts=time.time(), models=models)
                return models
        except Exception:
            pass
        return _model_cache["models"] or FALLBACK_MODELS


# ------------------------------------------------------------------ database
Base = declarative_base()
# Fail fast on unreachable Postgres — default TCP waits produced Cloudflare 524
# (origin accepts the socket, gunicorn workers never finish importing server.py).
_ENGINE_KWARGS: dict = {
    "pool_pre_ping": True,
    # Closed-beta target: several researchers online together.
    "pool_size": int(os.environ.get("DB_POOL_SIZE", "10")),
    "max_overflow": int(os.environ.get("DB_MAX_OVERFLOW", "20")),
}
if DATABASE_URL.startswith("postgresql"):
    _ENGINE_KWARGS["connect_args"] = {
        "connect_timeout": int(os.environ.get("DB_CONNECT_TIMEOUT", "10")),
    }
engine = create_engine(DATABASE_URL, **_ENGINE_KWARGS)
SessionLocal = sessionmaker(bind=engine, expire_on_commit=False)

# Production schema comes from migrations/*.sql (entrypoint + run_migrations.py).
# Running ORM create_all/ensure_columns at import on every Gunicorn worker can
# hang the process before any HTTP handler is registered.
_ORM_BOOTSTRAP_SCHEMA = (not IS_PRODUCTION) or DATABASE_URL.startswith("sqlite")


class User(Base):
    __tablename__ = "users"
    id = Column(Integer, primary_key=True)
    email = Column(String(320), unique=True, nullable=False)
    name = Column(String(200))
    picture = Column(String(500))
    custom_instructions = Column(Text, default="")
    # How this account was first created — 'google' | 'magic' | 'dev'. Set
    # once at creation, never overwritten by a later login via a different
    # method (an existing Google user who later uses a magic link logs
    # into the same account; their auth_provider stays 'google').
    auth_provider = Column(String(20), default="google")
    created_at = Column(DateTime, default=lambda: datetime.now(timezone.utc))

    # ── Quotas (quotas/service.py) ──────────────────────────────────────
    # Current storage usage is NOT duplicated here — StorageUsage.bytes_used
    # (below) is already the live, actively-maintained source of truth;
    # only the per-user *limit* is new. Both /api/files and
    # /api/documents/upload check against this same column now (falling
    # back to DEFAULT_STORAGE_LIMIT_BYTES) — they used to disagree, one
    # via the standalone MAX_STORAGE_MB env var (5000 MB default), the
    # other via this column's own default (~1000 MB) — see server.py's
    # upload_file() for why the increment side still doesn't share code
    # despite the check side now agreeing on the same limit.
    storage_limit_bytes = Column(BigInteger, default=QuotaService.DEFAULT_STORAGE_LIMIT_BYTES)
    monthly_token_used = Column(Integer, default=0)
    monthly_token_limit = Column(Integer, default=QuotaService.DEFAULT_TOKEN_LIMIT)
    quota_reset_at = Column(DateTime, nullable=True)

    # Prompt Engine admin routes (migrations/0016, backend/prompts/routes.py)
    # — default false for everyone; no signup flow ever sets this, the
    # first admin is always a manual DB update.
    is_admin = Column(Boolean, default=False)

    # Closed-beta ops (migrations/0025) — extend, don't replace auth.
    status = Column(String(30), default="active")  # pending_verification|active|suspended|deleted
    email_verified = Column(Boolean, default=False)
    email_verified_at = Column(DateTime, nullable=True)
    password_hash = Column(String(255), nullable=True)
    plan = Column(String(30), default="beta")  # free|beta|student|pro
    session_version = Column(Integer, default=0)
    monthly_cost_used = Column(Float, default=0.0)
    monthly_cost_limit = Column(Float, default=20.0)
    last_login_at = Column(DateTime, nullable=True)
    onboarding_completed_at = Column(DateTime, nullable=True)
    research_role = Column(String(40), nullable=True)
    research_fields = Column(Text, nullable=True)  # comma-separated field ids
    institution = Column(String(200), nullable=True)
    research_goal = Column(String(40), nullable=True)
    experience_level = Column(String(20), nullable=True)
    # Legacy; prefer structured columns above
    onboarding_json = Column(Text, nullable=True)


class Project(Base):
    __tablename__ = "projects"
    id = Column(Integer, primary_key=True)
    user_id = Column(Integer, ForeignKey("users.id"), nullable=False)
    name = Column(String(100), nullable=False)
    emoji = Column(String(16), default="📁")
    description = Column(Text, default="")  # what this research project is about
    instructions = Column(Text, default="")
    created_at = Column(DateTime, default=lambda: datetime.now(timezone.utc))


class Conversation(Base):
    __tablename__ = "conversations"
    id = Column(Integer, primary_key=True)
    user_id = Column(Integer, ForeignKey("users.id"), nullable=False)
    project_id = Column(Integer, ForeignKey("projects.id"), nullable=True)
    # When set, this is a *paper chat*: retrieval is hard-scoped to this one
    # document and the assistant is told to answer from it alone.
    file_id = Column(Integer, ForeignKey("files.id"), nullable=True)
    title = Column(String(200), default="New chat")
    title_generated = Column(Integer, default=0)
    model = Column(String(100), default=DEFAULT_MODEL)
    temperature = Column(Float, nullable=True)
    reasoning_effort = Column(String(20), nullable=True)
    memory_enabled = Column(Integer, default=1)
    created_at = Column(DateTime, default=lambda: datetime.now(timezone.utc))
    updated_at = Column(
        DateTime,
        default=lambda: datetime.now(timezone.utc),
        onupdate=lambda: datetime.now(timezone.utc),
    )
    messages = relationship(
        "Message",
        back_populates="conversation",
        cascade="all, delete-orphan",
        order_by="Message.id",
    )


class Message(Base):
    __tablename__ = "messages"
    id = Column(Integer, primary_key=True)
    conversation_id = Column(Integer, ForeignKey("conversations.id"), nullable=False)
    role = Column(String(20), nullable=False)
    content = Column(Text, nullable=False)
    sources = Column(Text)  # JSON list of web sources
    attachments = Column(Text)  # JSON list [{id,name,mime,kind}]
    created_at = Column(DateTime, default=lambda: datetime.now(timezone.utc))
    conversation = relationship("Conversation", back_populates="messages")


class Memory(Base):
    """User / research memories.

    Sprint C: research memories are AI-promoted from DerivedAnalysis(kind=research).
    Chat-extracted rows use source='chat' and must not enter research context.
    """

    __tablename__ = "memories"
    id = Column(Integer, primary_key=True)
    user_id = Column(Integer, ForeignKey("users.id"), nullable=False)
    project_id = Column(Integer, ForeignKey("projects.id"), nullable=True)
    fact = Column(Text, nullable=False)
    importance = Column(Integer, default=3)
    created_at = Column(DateTime, default=lambda: datetime.now(timezone.utc))
    # Sprint C research memory
    kind = Column(String(30), default="fact")  # finding|claim|contradiction|open_question|insight|fact
    source = Column(String(20), default="chat")  # research|compare|gaps|manual|chat
    source_ref = Column(String(80), default="")
    payload = Column(Text, default="{}")  # JSON
    pinned = Column(Integer, default=0)
    status = Column(String(20), default="active")  # active|archived|deleted
    claim_hash = Column(String(64), default="")


class UserFile(Base):
    __tablename__ = "files"
    id = Column(Integer, primary_key=True)
    user_id = Column(Integer, ForeignKey("users.id"), nullable=False)
    project_id = Column(Integer, ForeignKey("projects.id"), nullable=True)
    conversation_id = Column(Integer, ForeignKey("conversations.id"), nullable=True)
    name = Column(String(300), nullable=False)
    mime = Column(String(120))
    kind = Column(String(20))  # image | document
    path = Column(String(500))  # on-disk path
    size = Column(Integer, default=0)
    text_len = Column(Integer, default=0)
    created_at = Column(DateTime, default=lambda: datetime.now(timezone.utc))

    # ── Knowledge Library metadata (v1.0) ──────────────────────────────
    # `name` stays the raw filename; `title` is the extracted paper title so
    # the library can show "Attention Is All You Need" instead of "1706.pdf".
    title = Column(String(500), default="")
    authors = Column(String(1000), default="")  # "Vaswani, A.; Shazeer, N."
    year = Column(String(10), default="")
    venue = Column(String(300), default="")
    doi = Column(String(200), default="")
    abstract = Column(Text, default="")
    reading_status = Column(String(20), default="unread")  # unread|reading|read
    tags = Column(Text, default="[]")  # JSON list[str], user-set
    content_hash = Column(String(64), default="")  # sha256 of extracted text
    meta_status = Column(String(20), default="pending")  # pending|done|failed

    # sha256 of the raw uploaded bytes (not the extracted text) — storage-level
    # identity used for duplicate detection and post-upload integrity checks.
    checksum_sha256 = Column(String(64), nullable=True)

    # Scholarly provenance (migration 0018 / 0021). Discover stubs use
    # source_url for the OpenAlex OA / landing link when path is empty.
    metadata_source = Column(String(30), default="extracted")  # extracted|crossref|openalex|user
    source_url = Column(String(500), default="")
    doi_verified = Column(Boolean, default=False)
    # Phase 1b — stable identity for Connect library sync
    external_provider = Column(String(30), default="")  # zotero|mendeley|…
    external_item_id = Column(String(120), default="")
    # UFTR provenance (migration 0040) — outcome, attempts, full_text_source
    fulltext_json = Column(Text, default="{}")

    chunks = relationship("Chunk", cascade="all, delete-orphan", back_populates="file")


class UploadSession(Base):
    """Tracks a presigned/multipart upload between 'client asked for a URL'
    and 'client confirmed the bytes landed' — before a UserFile row exists.
    Expired/abandoned sessions are what garbage collection cleans up."""

    __tablename__ = "upload_sessions"
    id = Column(Integer, primary_key=True)
    user_id = Column(Integer, ForeignKey("users.id"), nullable=False)
    project_id = Column(Integer, ForeignKey("projects.id"), nullable=True)
    conversation_id = Column(Integer, ForeignKey("conversations.id"), nullable=True)
    key = Column(String(300), nullable=False)  # storage object key
    name = Column(String(300), nullable=False)  # original filename
    mime = Column(String(120))
    size_expected = Column(Integer, default=0)
    checksum_sha256 = Column(String(64), nullable=True)  # client-claimed, pre-upload
    upload_id = Column(String(300), nullable=True)  # multipart only
    status = Column(String(20), default="pending")  # pending|uploaded|confirmed|expired|aborted
    created_at = Column(DateTime, default=lambda: datetime.now(timezone.utc))


class UploadBatch(Base):
    """Groups files uploaded together in one user action (drag five PDFs
    at once). Nothing creates these yet — today's upload routes handle one
    file per request — so this stays empty until a bulk-upload entry point
    exists; the FK on UploadJob is nullable for exactly that reason."""

    __tablename__ = "upload_batches"
    id = Column(Integer, primary_key=True)
    user_id = Column(Integer, ForeignKey("users.id"), nullable=False)
    project_id = Column(Integer, ForeignKey("projects.id"), nullable=True)
    conversation_id = Column(Integer, ForeignKey("conversations.id"), nullable=True)
    source = Column(String(20), default="library")  # library|chat_composer|folder_drop
    file_count = Column(Integer, default=0)
    created_at = Column(DateTime, default=lambda: datetime.now(timezone.utc))


class UploadJob(Base):
    """One row per pipeline stage per file (import | phase1_analysis |
    paper_analysis; extract_metadata is legacy drain-only) — the actual
    queue worker.py polls with FOR UPDATE SKIP LOCKED, claims, executes,
    and marks done/failed. Written by upload_file()/confirm_upload()'s
    transactional outbox and by worker.py itself when chaining follow-on
    stages — see processing-pipeline-architecture.md."""

    __tablename__ = "upload_jobs"
    id = Column(Integer, primary_key=True)
    upload_batch_id = Column(Integer, ForeignKey("upload_batches.id"), nullable=True)
    file_id = Column(Integer, ForeignKey("files.id"), nullable=True)
    user_id = Column(Integer, ForeignKey("users.id"), nullable=False)
    job_type = Column(String(40), nullable=False)  # import|extract_metadata|paper_analysis|phase1_analysis
    status = Column(String(20), default="pending")  # pending|running|done|failed
    attempts = Column(Integer, default=0)
    run_after = Column(DateTime, default=lambda: datetime.now(timezone.utc))  # due time; backoff pushes this out
    locked_by = Column(Text, nullable=True)
    locked_at = Column(DateTime, nullable=True)
    last_error = Column(Text, nullable=True)
    # References pipeline_versions(id) (backend/ai/models.py) — plain
    # column, no ORM-level FK: pipeline_versions has no Python class
    # instantiated anywhere yet (see brain.md §7), so there's no target
    # to point a SQLAlchemy ForeignKey at. migrations/0005 adds the real
    # DB-level FK constraint once that table exists.
    pipeline_version_id = Column(Integer, nullable=True)
    started_at = Column(DateTime, nullable=True)
    finished_at = Column(DateTime, nullable=True)
    created_at = Column(DateTime, default=lambda: datetime.now(timezone.utc))
    updated_at = Column(
        DateTime,
        default=lambda: datetime.now(timezone.utc),
        onupdate=lambda: datetime.now(timezone.utc),
    )


class StorageUsage(Base):
    """Live per-user storage total — updated in the same transaction as
    every upload/delete, not a periodic rollup, because quota enforcement
    (production-hardening.md §4) needs a synchronous answer before
    accepting a new file, not yesterday's number."""

    __tablename__ = "storage_usage"
    user_id = Column(Integer, ForeignKey("users.id"), primary_key=True)
    bytes_used = Column(Integer, default=0)
    file_count = Column(Integer, default=0)
    updated_at = Column(
        DateTime,
        default=lambda: datetime.now(timezone.utc),
        onupdate=lambda: datetime.now(timezone.utc),
    )


UsageLog = create_usage_log_model(Base)
quota_service = QuotaService(SessionLocal, User, StorageUsage, UsageLog, select)
# EntitlementService is constructed after SystemSettingsService (ops) below.

FeatureFlag = create_feature_flag_model(Base)
# FeatureFlagService constructed after ops settings (same block as entitlements).


class ImportSession(Base):
    """Resumable checkpoint for a long-running import — schema only for
    now. Today's extraction is one synchronous pass per file, not the
    step-by-step resumable execution this is designed for; nothing writes
    real checkpoints here until the Step Runner (processing-pipeline-
    architecture.md §5, §10) exists. Created now so that work has a table
    to land in without a later migration."""

    __tablename__ = "import_sessions"
    id = Column(Integer, primary_key=True)
    upload_job_id = Column(Integer, ForeignKey("upload_jobs.id"), nullable=False, unique=True)
    stage = Column(String(20), default="extract")  # extract|chunk|embed
    checkpoint = Column(Text, default="{}")  # JSON
    updated_at = Column(
        DateTime,
        default=lambda: datetime.now(timezone.utc),
        onupdate=lambda: datetime.now(timezone.utc),
    )


class OutboxEvent(Base):
    """Transactional outbox: written in the SAME commit as the state
    change it announces (an UploadJob being enqueued), so a Queue Worker
    polling this table never misses an event to a process crash between
    'job row committed' and 'thread started' — the failure mode the old
    threading.Thread(daemon=True) approach had no protection against.
    aggregate_id is polymorphic (aggregate_type says which table it points
    into); no FK, enforced by the writer inserting both rows in one
    transaction, not by the schema."""

    __tablename__ = "outbox_events"
    id = Column(Integer, primary_key=True)
    aggregate_type = Column(String(30), nullable=False)  # 'upload_job' | ...
    aggregate_id = Column(Integer, nullable=False)
    event_type = Column(String(50), nullable=False)  # 'job.enqueued' | ...
    payload = Column(Text, nullable=False)  # JSON
    status = Column(String(20), default="pending")  # pending|dispatched|failed
    dispatched_at = Column(DateTime, nullable=True)
    created_at = Column(DateTime, default=lambda: datetime.now(timezone.utc))


class ModelVersion(Base):
    """Our own versioning of a model choice — so "which model produced
    this row" survives an env var change. Seeded by backfill.py's Task 3
    pass (default_model/utility_model/embed_model, version 1, active)."""

    __tablename__ = "model_versions"
    id = Column(Integer, primary_key=True)
    logical_name = Column(String(50), nullable=False)
    provider_model_id = Column(String(100), nullable=False)
    version = Column(Integer, nullable=False)
    is_active = Column(Boolean, default=True)
    created_at = Column(DateTime, default=lambda: datetime.now(timezone.utc))


class AIUsageLedger(Base):
    """Append-only cost/usage record — one row per OpenAI call, written
    from the two functions that actually make the calls (responses_text,
    embed_texts), the same choke points research-intelligence.md §7
    designed this around. Currently wired up for the import->embed,
    extract_metadata, and paper_analysis paths only — see worker.py's
    docstring notes for what isn't covered yet (chat, memory extraction,
    titles, compare, gap-finder, writing assistant)."""

    __tablename__ = "ai_usage_ledger"
    id = Column(Integer, primary_key=True)
    user_id = Column(Integer, ForeignKey("users.id"), nullable=False)
    upload_job_id = Column(Integer, ForeignKey("upload_jobs.id"), nullable=True)
    kind = Column(String(30), nullable=False)  # embedding|metadata|analysis|...
    model_version_id = Column(Integer, ForeignKey("model_versions.id"), nullable=False)
    # References prompt_versions(id) — no ORM-level FK for the same
    # reason as UploadJob.pipeline_version_id above: prompt_versions has
    # no server.py-registered class (it lives under backend/ai's own
    # private Base). migrations/0006 adds the real DB-level FK.
    prompt_version_id = Column(Integer, nullable=True)
    prompt_tokens = Column(Integer, default=0)
    completion_tokens = Column(Integer, default=0)
    cost_usd = Column(Float, default=0.0)
    created_at = Column(DateTime, default=lambda: datetime.now(timezone.utc))


class WorkerHeartbeat(Base):
    """Single row (id=1), upserted by worker.py once per poll loop
    iteration — the only signal server.py has that the separate worker.py
    process is actually alive, since it isn't a thread/child process it
    can introspect directly. GET /api/worker/health compares
    last_seen_at against now(); no row at all (id=1 missing) means the
    worker has never run since this table existed."""

    __tablename__ = "worker_heartbeats"
    id = Column(Integer, primary_key=True)
    # timezone=True (-> TIMESTAMPTZ on Postgres), not a bare DateTime:
    # verified live against a real Postgres instance whose session
    # timezone isn't UTC (Asia/Karachi, UTC+5) — writing a UTC-aware
    # datetime into a naive TIMESTAMP column gets silently shifted to the
    # session's zone on write, then misread as if it already were UTC on
    # the way back out, producing a consistent 5-hour skew (a negative
    # age_seconds in GET /api/worker/health). SQLite has no real
    # per-session timezone, so this was invisible there — only showed up
    # once this was actually run against Postgres. The migration
    # (0013_worker_heartbeat.sql) already declared timestamptz correctly;
    # this column just didn't match it. See migrations/0014.
    last_seen_at = Column(DateTime(timezone=True), nullable=False)


class Chunk(Base):
    __tablename__ = "chunks"
    id = Column(Integer, primary_key=True)
    file_id = Column(Integer, ForeignKey("files.id"), nullable=False)
    idx = Column(Integer, default=0)
    content = Column(Text, nullable=False)
    embedding = Column(Text)  # JSON list of floats (null = keyword only)
    # Locators so Paper Chat can cite "p. 4 · §Methodology" rather than just
    # the filename. Nullable: chunks from formats without page structure
    # (and every chunk written by the pre-1.0 code) simply have no locator.
    page = Column(Integer, nullable=True)
    section = Column(String(200), nullable=True)
    file = relationship("UserFile", back_populates="chunks")


class Citation(Base):
    __tablename__ = "citations"
    id = Column(Integer, primary_key=True)
    user_id = Column(Integer, ForeignKey("users.id"), nullable=False)
    project_id = Column(Integer, ForeignKey("projects.id"), nullable=True)
    authors = Column(String(500), default="")
    title = Column(String(500), default="")
    year = Column(String(10), default="")
    venue = Column(String(300), default="")  # journal / conference / publisher
    doi = Column(String(200), default="")
    url = Column(String(600), default="")
    notes = Column(Text, default="")
    created_at = Column(DateTime, default=lambda: datetime.now(timezone.utc))


class SupportRequest(Base):
    __tablename__ = "support_requests"
    id = Column(Integer, primary_key=True)
    user_id = Column(Integer, ForeignKey("users.id"), nullable=True)
    email = Column(String(320), nullable=False)
    subject = Column(String(300), default="")
    category = Column(String(50), default="general")  # bug|feature|general|account
    message = Column(Text, nullable=False)
    status = Column(String(30), default="open")  # open|in_progress|closed
    created_at = Column(DateTime, default=lambda: datetime.now(timezone.utc))


# ══════════════════════════════════════════════════════════════════════════
# RESEARCH WORKSPACE — v1.0 models
# ══════════════════════════════════════════════════════════════════════════


class PaperAnalysis(Base):
    """Cached structured analysis of one paper. One row per file.

    `content_hash` is the SHA-256 of the extracted text: if a document is
    re-uploaded unchanged we reuse the analysis instead of paying for another
    model call. Regeneration happens only when the hash changes or the user
    explicitly asks for a refresh."""

    __tablename__ = "paper_analyses"
    id = Column(Integer, primary_key=True)
    file_id = Column(Integer, ForeignKey("files.id"), nullable=False, unique=True)
    user_id = Column(Integer, ForeignKey("users.id"), nullable=False)
    status = Column(String(20), default="pending")  # pending|running|done|failed
    error = Column(Text, default="")
    content_hash = Column(String(64), default="")
    model = Column(String(100), default="")
    data = Column(Text, default="")  # JSON: ANALYSIS_FIELDS
    created_at = Column(DateTime, default=lambda: datetime.now(timezone.utc))
    updated_at = Column(
        DateTime,
        default=lambda: datetime.now(timezone.utc),
        onupdate=lambda: datetime.now(timezone.utc),
    )


class AnalysisPipelineResult(Base):
    """Persisted Phase 1.1–1.7 outputs for one file (Phase 2 integration).

    `phase_results` is a JSON object keyed by phase name
    (document_understanding, classification, …). Lazy migration: rows are
    created when a document is analyzed; older files remain without a row
    until first access."""

    __tablename__ = "analysis_pipeline_results"
    id = Column(Integer, primary_key=True)
    file_id = Column(Integer, ForeignKey("files.id"), nullable=False, unique=True)
    user_id = Column(Integer, ForeignKey("users.id"), nullable=False)
    content_hash = Column(String(64), default="")
    status = Column(String(20), default="pending")  # pending|running|done|failed|partial
    error = Column(Text, default="")
    phase_results = Column(Text, default="{}")  # JSON
    pipeline_version = Column(String(50), default="")
    total_processing_time_ms = Column(Integer, default=0)
    created_at = Column(DateTime, default=lambda: datetime.now(timezone.utc))
    updated_at = Column(
        DateTime,
        default=lambda: datetime.now(timezone.utc),
        onupdate=lambda: datetime.now(timezone.utc),
    )


class DerivedAnalysis(Base):
    """Cached multi-paper output — comparison ('compare') or gap analysis
    ('gaps'). Keyed by a hash of the sorted file-id set so the same selection
    never regenerates."""

    __tablename__ = "derived_analyses"
    id = Column(Integer, primary_key=True)
    user_id = Column(Integer, ForeignKey("users.id"), nullable=False)
    project_id = Column(Integer, ForeignKey("projects.id"), nullable=True)
    kind = Column(String(20), nullable=False)  # compare|gaps|research
    selection_hash = Column(String(64), nullable=False)
    file_ids = Column(Text, default="[]")  # JSON list[int]
    data = Column(Text, default="")  # JSON
    model = Column(String(100), default="")
    created_at = Column(DateTime, default=lambda: datetime.now(timezone.utc))


class Note(Base):
    __tablename__ = "notes"
    id = Column(Integer, primary_key=True)
    user_id = Column(Integer, ForeignKey("users.id"), nullable=False)
    project_id = Column(Integer, ForeignKey("projects.id"), nullable=True)
    file_id = Column(Integer, ForeignKey("files.id"), nullable=True)  # paper note
    title = Column(String(300), default="")
    content = Column(Text, default="")
    created_at = Column(DateTime, default=lambda: datetime.now(timezone.utc))
    updated_at = Column(
        DateTime,
        default=lambda: datetime.now(timezone.utc),
        onupdate=lambda: datetime.now(timezone.utc),
    )


class WritingDocument(Base):
    __tablename__ = "documents"
    id = Column(Integer, primary_key=True)
    user_id = Column(Integer, ForeignKey("users.id"), nullable=False)
    project_id = Column(Integer, ForeignKey("projects.id"), nullable=False)
    title = Column(String(300), default="")
    content = Column(Text, default="")
    editor_kind = Column(String(20), default="markdown")  # markdown|richtext
    status = Column(String(20), default="draft")  # draft|active|archived|deleted
    current_version = Column(Integer, default=1)  # optimistic locking
    last_saved_hash = Column(String(64), default="")
    last_autosave_key = Column(String(120), default="")
    last_opened_at = Column(DateTime, nullable=True)
    word_count = Column(Integer, default=0)
    created_at = Column(DateTime, default=lambda: datetime.now(timezone.utc))
    updated_at = Column(
        DateTime,
        default=lambda: datetime.now(timezone.utc),
        onupdate=lambda: datetime.now(timezone.utc),
    )


class WritingDocumentVersion(Base):
    __tablename__ = "document_versions"
    id = Column(Integer, primary_key=True)
    document_id = Column(Integer, ForeignKey("documents.id"), nullable=False)
    user_id = Column(Integer, ForeignKey("users.id"), nullable=False)
    version_no = Column(Integer, nullable=False)
    title = Column(String(300), default="")
    content = Column(Text, default="")
    content_hash = Column(String(64), default="")
    source = Column(String(20), default="save")  # create|save|autosave|restore
    created_at = Column(DateTime, default=lambda: datetime.now(timezone.utc))


class WritingDocumentActivity(Base):
    __tablename__ = "document_activity"
    id = Column(Integer, primary_key=True)
    document_id = Column(Integer, ForeignKey("documents.id"), nullable=False)
    user_id = Column(Integer, ForeignKey("users.id"), nullable=False)
    action = Column(String(30), nullable=False)  # create|update|autosave|restore|archive
    meta_json = Column(Text, default="{}")
    created_at = Column(DateTime, default=lambda: datetime.now(timezone.utc))


class EvidenceObject(Base):
    """Canonical Evidence Layer row (Week 2 / Phase 2.2). Soft FKs to files/projects."""

    __tablename__ = "evidence_objects"
    id = Column(Integer, primary_key=True)
    user_id = Column(Integer, nullable=False)
    project_id = Column(Integer, nullable=False)
    file_id = Column(Integer, nullable=False)
    page = Column(Integer, nullable=True)
    char_start = Column(Integer, nullable=True)
    char_end = Column(Integer, nullable=True)
    section = Column(String(200), default="")
    quote = Column(Text, nullable=False, default="")
    claim = Column(Text, nullable=False, default="")
    study_type = Column(String(80), default="")
    study_quality = Column(String(40), default="")
    supports_json = Column(Text, default="[]")
    contradicts_json = Column(Text, default="[]")
    limitations_json = Column(Text, default="[]")
    confidence_band = Column(String(20), default="low")
    status = Column(String(20), default="candidate")
    pipeline_version = Column(String(40), nullable=False, default="2.2.0")
    created_by = Column(String(80), default="analysis-pipeline")
    content_hash = Column(String(64), nullable=False, default="")
    supersedes_id = Column(Integer, nullable=True)
    provenance_json = Column(Text, default="{}")
    source_kg_node_id = Column(String(120), default="")
    created_at = Column(DateTime, default=lambda: datetime.now(timezone.utc))
    updated_at = Column(
        DateTime,
        default=lambda: datetime.now(timezone.utc),
        onupdate=lambda: datetime.now(timezone.utc),
    )


class ClaimReview(Base):
    __tablename__ = "claim_reviews"
    id = Column(Integer, primary_key=True)
    evidence_object_id = Column(Integer, nullable=False)
    user_id = Column(Integer, nullable=False)
    project_id = Column(Integer, nullable=False)
    status = Column(String(20), nullable=False)
    reason = Column(Text, default="")
    edited_claim = Column(Text, nullable=True)
    edited_quote = Column(Text, nullable=True)
    reviewed_at = Column(DateTime, default=lambda: datetime.now(timezone.utc))


class ResearchDecision(Base):
    """Phase A.2 append-only ledger of researcher judgments (project memory)."""

    __tablename__ = "research_decisions"
    id = Column(Integer, primary_key=True)
    user_id = Column(Integer, nullable=False)
    project_id = Column(Integer, nullable=False)
    evidence_object_id = Column(Integer, nullable=False)
    decision_type = Column(String(40), nullable=False)
    reason = Column(Text, default="")
    reason_code = Column(String(120), default="")
    created_at = Column(DateTime, default=lambda: datetime.now(timezone.utc))


class WorkflowEvent(Base):
    """Phase A.6 append-only workflow instrumentation breadcrumbs."""

    __tablename__ = "workflow_events"
    id = Column(Integer, primary_key=True)
    user_id = Column(Integer, nullable=False)
    project_id = Column(Integer, nullable=True)
    event_name = Column(String(80), nullable=False)
    meta_json = Column(Text, default="{}")
    created_at = Column(DateTime, default=lambda: datetime.now(timezone.utc))


class WritingSentenceBinding(Base):
    __tablename__ = "writing_sentence_bindings"
    id = Column(Integer, primary_key=True)
    user_id = Column(Integer, nullable=False)
    project_id = Column(Integer, nullable=False)
    document_id = Column(Integer, nullable=False)
    evidence_object_id = Column(Integer, nullable=False)
    block_id = Column(String(120), default="")
    range_start = Column(Integer, nullable=True)
    range_end = Column(Integer, nullable=True)
    selected_text = Column(Text, default="")
    relation = Column(String(20), default="supports")
    created_by = Column(String(40), default="user")
    created_at = Column(DateTime, default=lambda: datetime.now(timezone.utc))


class EvidenceExtractionRun(Base):
    __tablename__ = "evidence_extraction_runs"
    id = Column(Integer, primary_key=True)
    user_id = Column(Integer, nullable=False)
    project_id = Column(Integer, nullable=False)
    file_id = Column(Integer, nullable=False)
    pipeline_version = Column(String(40), nullable=False)
    input_content_hash = Column(String(64), nullable=False)
    status = Column(String(20), default="queued")
    objects_created = Column(Integer, default=0)
    error_json = Column(Text, default="{}")
    job_id = Column(Integer, nullable=True)
    created_at = Column(DateTime, default=lambda: datetime.now(timezone.utc))
    finished_at = Column(DateTime, nullable=True)


class ReviewerRun(Base):
    """Durable Research Reviewer run (A-401 / EPIC-0005 A-503).

    Distinct from ClaimReview (human EvidenceObject accept/reject).
    Soft Integer FKs — ownership checked in the evidence blueprint.
    """

    __tablename__ = "reviewer_runs"
    id = Column(Integer, primary_key=True)
    user_id = Column(Integer, nullable=False)
    project_id = Column(Integer, nullable=False)
    document_id = Column(Integer, nullable=False)
    document_version_no = Column(Integer, default=1)
    writing_version = Column(String(40), default="")
    reviewer_version = Column(String(40), nullable=False)
    binder_version = Column(String(40), default="")
    status = Column(String(20), nullable=False)
    pass_rate = Column(Float, default=0.0)
    sections_checked = Column(Integer, default=0)
    sections_passed = Column(Integer, default=0)
    issue_count = Column(Integer, default=0)
    metrics_json = Column(Text, default="{}")
    input_snapshot_json = Column(Text, default="{}")
    model_version_id = Column(Integer, nullable=True)
    prompt_version_id = Column(Integer, nullable=True)
    prompt_meta_json = Column(Text, default="{}")
    created_at = Column(DateTime, default=lambda: datetime.now(timezone.utc))
    finished_at = Column(DateTime, nullable=True)


class ReviewerFinding(Base):
    """One finding row belonging to a ReviewerRun (reconstructable history)."""

    __tablename__ = "reviewer_findings"
    id = Column(Integer, primary_key=True)
    run_id = Column(Integer, nullable=False)
    code = Column(String(80), nullable=False)
    severity = Column(String(20), nullable=False)
    message = Column(Text, default="")
    section_id = Column(String(120), nullable=True)
    block_id = Column(String(120), default="")
    range_start = Column(Integer, nullable=True)
    range_end = Column(Integer, nullable=True)
    selected_text = Column(Text, default="")
    evidence_ids_json = Column(Text, default="[]")
    confidence_band = Column(String(20), default="")
    recommendation = Column(Text, default="")
    status = Column(String(20), default="open")  # open|accepted|dismissed|fixed
    resolution_rationale = Column(Text, default="")
    resolved_at = Column(DateTime, nullable=True)
    resolved_by = Column(Integer, nullable=True)
    created_at = Column(DateTime, default=lambda: datetime.now(timezone.utc))


class ProjectQuestion(Base):
    """Research question scoped to a project (Sprint A).

    User-authored tracking — not notes (freeform prose) and not memories
    (AI-curated findings). Status: open | answered | parked.
    """

    __tablename__ = "project_questions"
    id = Column(Integer, primary_key=True)
    user_id = Column(Integer, ForeignKey("users.id"), nullable=False)
    project_id = Column(Integer, ForeignKey("projects.id"), nullable=False)
    text = Column(Text, nullable=False)
    status = Column(String(20), default="open")  # open|answered|parked
    source = Column(String(20), default="manual")  # manual|ai
    linked_insight_id = Column(Integer, nullable=True)  # derived_analyses.id (no cross-Base FK)
    created_at = Column(DateTime, default=lambda: datetime.now(timezone.utc))
    updated_at = Column(
        DateTime,
        default=lambda: datetime.now(timezone.utc),
        onupdate=lambda: datetime.now(timezone.utc),
    )


class SearchIndex(Base):
    """Unified semantic index over notes, citations and chat messages.

    Papers are NOT indexed here — their `Chunk` rows already carry embeddings
    and are searched directly, which keeps a single source of truth per
    document and avoids duplicating a paper's vectors."""

    __tablename__ = "search_index"
    id = Column(Integer, primary_key=True)
    user_id = Column(Integer, ForeignKey("users.id"), nullable=False)
    project_id = Column(Integer, ForeignKey("projects.id"), nullable=True)
    kind = Column(String(20), nullable=False)  # note|citation|chat
    ref_id = Column(Integer, nullable=False)  # id in the source table
    title = Column(String(400), default="")
    snippet = Column(Text, default="")
    embedding = Column(Text)  # JSON list[float] | null
    updated_at = Column(
        DateTime,
        default=lambda: datetime.now(timezone.utc),
        onupdate=lambda: datetime.now(timezone.utc),
    )


# Closed-beta ops tables (migrations/0025) — registered on this Base so
# SQLite create_all creates them; Postgres gets them via run_migrations.
from auth.magic_link import create_magic_link_token_model
from security.ops import (
    create_email_token_models,
    create_invite_token_model,
    create_security_event_model,
    create_system_settings_model,
)
from backend.library.models import (
    create_library_collection_models,
    create_library_connection_model,
    create_library_sync_run_model,
)

SystemSetting = create_system_settings_model(Base)
SecurityEvent = create_security_event_model(Base)
InviteToken = create_invite_token_model(Base)
EmailVerificationToken, PasswordResetToken, EmailChangeToken = create_email_token_models(Base)
MagicLinkToken = create_magic_link_token_model(Base)
LibraryConnection = create_library_connection_model(Base)
LibraryCollection, LibraryCollectionPaper = create_library_collection_models(Base)
LibrarySyncRun = create_library_sync_run_model(Base)


# checkfirst=True is SQLAlchemy's own default (verified: MetaData.create_all's
# signature already defaults to it) — spelled out explicitly so it's not
# a fact a reader has to already know. It's what makes this call safe to
# run after migrations/*.sql already created these same tables: it only
# creates what's missing, never re-creates or errors on what exists.
# Skipped in production Postgres — entrypoint already ran run_migrations.py.
if _ORM_BOOTSTRAP_SCHEMA:
    Base.metadata.create_all(engine, checkfirst=True)


def ensure_columns():
    """Lightweight migrations for DBs created by a previous version.

    Each statement is attempted independently and failures are swallowed —
    "duplicate column" is the expected outcome on an already-migrated DB.
    Every column added here is nullable or has a default, so old rows stay
    valid and the app keeps working without a backfill step."""
    for stmt in (
        # ── pre-1.0 ────────────────────────────────────────────────────
        "ALTER TABLE messages ADD COLUMN attachments TEXT",
        "ALTER TABLE users ADD COLUMN custom_instructions TEXT",
        "ALTER TABLE memories ADD COLUMN importance INTEGER DEFAULT 3",
        "ALTER TABLE memories ADD COLUMN kind VARCHAR(30) DEFAULT 'fact'",
        "ALTER TABLE memories ADD COLUMN source VARCHAR(20) DEFAULT 'chat'",
        "ALTER TABLE memories ADD COLUMN source_ref VARCHAR(80) DEFAULT ''",
        "ALTER TABLE memories ADD COLUMN payload TEXT DEFAULT '{}'",
        "ALTER TABLE memories ADD COLUMN pinned INTEGER DEFAULT 0",
        "ALTER TABLE memories ADD COLUMN status VARCHAR(20) DEFAULT 'active'",
        "ALTER TABLE memories ADD COLUMN claim_hash VARCHAR(64) DEFAULT ''",
        "ALTER TABLE conversations ADD COLUMN temperature FLOAT",
        "ALTER TABLE conversations ADD COLUMN reasoning_effort VARCHAR(20)",
        "ALTER TABLE conversations ADD COLUMN memory_enabled INTEGER DEFAULT 1",
        # ── Research Workspace v1.0 ────────────────────────────────────
        "ALTER TABLE projects ADD COLUMN description TEXT",
        "ALTER TABLE conversations ADD COLUMN file_id INTEGER",
        "ALTER TABLE files ADD COLUMN title VARCHAR(500)",
        "ALTER TABLE files ADD COLUMN authors VARCHAR(1000)",
        "ALTER TABLE files ADD COLUMN year VARCHAR(10)",
        "ALTER TABLE files ADD COLUMN venue VARCHAR(300)",
        "ALTER TABLE files ADD COLUMN doi VARCHAR(200)",
        "ALTER TABLE files ADD COLUMN abstract TEXT",
        "ALTER TABLE files ADD COLUMN reading_status VARCHAR(20) DEFAULT 'unread'",
        "ALTER TABLE files ADD COLUMN tags TEXT DEFAULT '[]'",
        "ALTER TABLE files ADD COLUMN content_hash VARCHAR(64)",
        "ALTER TABLE files ADD COLUMN meta_status VARCHAR(20) DEFAULT 'pending'",
        "ALTER TABLE chunks ADD COLUMN page INTEGER",
        "ALTER TABLE chunks ADD COLUMN section VARCHAR(200)",
        # ── Storage architecture ───────────────────────────────────────
        "ALTER TABLE files ADD COLUMN checksum_sha256 VARCHAR(64)",
        "ALTER TABLE upload_jobs ADD COLUMN run_after TIMESTAMP",
        "ALTER TABLE users ADD COLUMN auth_provider VARCHAR(20) DEFAULT 'google'",
        "ALTER TABLE users ADD COLUMN storage_limit_bytes BIGINT DEFAULT 1000000000",
        "ALTER TABLE users ADD COLUMN monthly_token_used INTEGER DEFAULT 0",
        "ALTER TABLE users ADD COLUMN monthly_token_limit INTEGER DEFAULT 100000",
        "ALTER TABLE users ADD COLUMN quota_reset_at TIMESTAMP",
        # ── ORM/migration drift found while fixing run_migrations.py:
        # these columns were in migrations/0002 and 0006 but never made
        # it into the UploadJob/AIUsageLedger classes above, so a
        # SQLite dev DB (which only ever runs the ORM's create_all, never
        # migrations/*.sql) was permanently missing them.
        "ALTER TABLE upload_jobs ADD COLUMN locked_by TEXT",
        "ALTER TABLE upload_jobs ADD COLUMN locked_at TIMESTAMP",
        "ALTER TABLE upload_jobs ADD COLUMN pipeline_version_id INTEGER",
        "ALTER TABLE ai_usage_ledger ADD COLUMN prompt_version_id INTEGER",
        # ── Prompt Engine (migrations/0015, docs/prompt-engine-architecture.md) ──
        # prompt_versions lives under backend/ai's own private Base, not this
        # one — same reasoning as the two rows above: create_all() only
        # creates a table it doesn't already find, never backfills a column
        # onto one that already exists, so an existing chat_dev.db needs
        # these applied explicitly too.
        "ALTER TABLE prompt_versions ADD COLUMN description TEXT DEFAULT ''",
        "ALTER TABLE prompt_versions ADD COLUMN status TEXT DEFAULT 'draft'",
        "ALTER TABLE prompt_versions ADD COLUMN category TEXT DEFAULT ''",
        "ALTER TABLE prompt_versions ADD COLUMN examples TEXT DEFAULT '[]'",
        "ALTER TABLE prompt_versions ADD COLUMN expected_output_type TEXT DEFAULT 'text'",
        "ALTER TABLE prompt_versions ADD COLUMN author_user_id INTEGER",
        # model_registry_cost_ledger: same private-Base situation, third
        # registry (model_registry.py's own) distinct from both of the above.
        "ALTER TABLE model_registry_cost_ledger ADD COLUMN prompt_version_id INTEGER",
        # projects.instructions: not new to this task's schema — found
        # missing entirely (no migration, no ensure_columns entry) while
        # checking this table per the Prompt Engine work; the ORM column
        # (Project.instructions) has existed and been read/written for a
        # while, so this was a real, previously-undocumented gap.
        "ALTER TABLE projects ADD COLUMN instructions TEXT",
        # users.is_admin — migrations/0016, backend/prompts/routes.py's
        # admin-gated create/update routes.
        "ALTER TABLE users ADD COLUMN is_admin BOOLEAN DEFAULT 0",
        # Closed-beta ops (migrations/0025)
        "ALTER TABLE users ADD COLUMN status VARCHAR(30) DEFAULT 'active'",
        "ALTER TABLE users ADD COLUMN email_verified BOOLEAN DEFAULT 0",
        "ALTER TABLE users ADD COLUMN email_verified_at TIMESTAMP",
        "ALTER TABLE users ADD COLUMN password_hash VARCHAR(255)",
        "ALTER TABLE users ADD COLUMN plan VARCHAR(30) DEFAULT 'beta'",
        "ALTER TABLE users ADD COLUMN session_version INTEGER DEFAULT 0",
        "ALTER TABLE users ADD COLUMN monthly_cost_used FLOAT DEFAULT 0",
        "ALTER TABLE users ADD COLUMN monthly_cost_limit FLOAT DEFAULT 20",
        "ALTER TABLE users ADD COLUMN last_login_at TIMESTAMP",
        "ALTER TABLE users ADD COLUMN onboarding_completed_at TIMESTAMP",
        "ALTER TABLE users ADD COLUMN research_role VARCHAR(40)",
        "ALTER TABLE users ADD COLUMN research_fields TEXT",
        "ALTER TABLE users ADD COLUMN institution VARCHAR(200)",
        "ALTER TABLE users ADD COLUMN research_goal VARCHAR(40)",
        "ALTER TABLE users ADD COLUMN experience_level VARCHAR(20)",
        "ALTER TABLE users ADD COLUMN onboarding_json TEXT",
        "ALTER TABLE model_registry_cost_ledger ADD COLUMN estimated_cost FLOAT",
        "ALTER TABLE model_registry_cost_ledger ADD COLUMN currency VARCHAR(8) DEFAULT 'USD'",
        # ── Scholarly provider integrations (migration 0018) ─────────────
        "ALTER TABLE files ADD COLUMN doi_verified BOOLEAN DEFAULT 0",
        "ALTER TABLE files ADD COLUMN crossref_last_synced TIMESTAMP",
        "ALTER TABLE files ADD COLUMN crossref_version VARCHAR(20) DEFAULT ''",
        "ALTER TABLE files ADD COLUMN title_source VARCHAR(30) DEFAULT 'extracted'",
        "ALTER TABLE files ADD COLUMN authors_source VARCHAR(30) DEFAULT 'extracted'",
        "ALTER TABLE files ADD COLUMN year_source VARCHAR(30) DEFAULT 'extracted'",
        "ALTER TABLE files ADD COLUMN venue_source VARCHAR(30) DEFAULT 'extracted'",
        "ALTER TABLE files ADD COLUMN abstract_source VARCHAR(30) DEFAULT 'extracted'",
        "ALTER TABLE files ADD COLUMN metadata_source VARCHAR(30) DEFAULT 'extracted'",
        # ── Discover import stubs (migration 0021) ───────────────────────
        "ALTER TABLE files ADD COLUMN source_url VARCHAR(500) DEFAULT ''",
        # ── Library sync Phase 1b (migration 0030) ───────────────────────
        "ALTER TABLE files ADD COLUMN external_provider VARCHAR(30) DEFAULT ''",
        "ALTER TABLE files ADD COLUMN external_item_id VARCHAR(120) DEFAULT ''",
        "ALTER TABLE library_connections ADD COLUMN last_synced_at TIMESTAMP",
        "ALTER TABLE library_connections ADD COLUMN sync_cursor TEXT DEFAULT ''",
        # ── UFTR full-text resolution provenance (migration 0040) ───────
        "ALTER TABLE files ADD COLUMN fulltext_json TEXT DEFAULT '{}'",
    ):
        try:
            with engine.begin() as conn:
                conn.execute(sqltext(stmt))
        except Exception:
            pass

    # Indexes for the new access patterns (library listing, cache lookups,
    # semantic search). CREATE INDEX IF NOT EXISTS works on both SQLite and
    # Postgres, so this is safe to run on every boot.
    for stmt in (
        "CREATE INDEX IF NOT EXISTS ix_files_user ON files (user_id)",
        "CREATE INDEX IF NOT EXISTS ix_chunks_file ON chunks (file_id)",
        "CREATE INDEX IF NOT EXISTS ix_notes_user ON notes (user_id)",
        "CREATE INDEX IF NOT EXISTS ix_paper_analyses_file ON paper_analyses (file_id)",
        "CREATE INDEX IF NOT EXISTS ix_derived_sel ON derived_analyses (user_id, kind, selection_hash)",
        "CREATE INDEX IF NOT EXISTS ix_search_index_user ON search_index (user_id, kind)",
        "CREATE INDEX IF NOT EXISTS ix_files_user_checksum ON files (user_id, checksum_sha256)",
        "CREATE INDEX IF NOT EXISTS ix_upload_sessions_user ON upload_sessions (user_id)",
        # ── Storage foundation (database-design.md) ─────────────────────
        # ix_upload_batches_user and ix_outbox_events_pending must match
        # migrations 0001/0007's definitions exactly (column list / WHERE
        # clause), not just the name — Postgres's CREATE INDEX IF NOT
        # EXISTS only checks the name, so if this ran first with a lesser
        # definition, the migration's more complete one would silently
        # never get created (same name = skipped), not just redundantly
        # re-created. Found by actually comparing both files side by
        # side, not assumed.
        "CREATE INDEX IF NOT EXISTS ix_upload_batches_user ON upload_batches (user_id, created_at DESC)",
        "CREATE INDEX IF NOT EXISTS ix_upload_jobs_file ON upload_jobs (file_id)",
        "CREATE INDEX IF NOT EXISTS ix_upload_jobs_batch ON upload_jobs (upload_batch_id)",
        "CREATE INDEX IF NOT EXISTS ix_upload_jobs_user_status ON upload_jobs (user_id, status)",
        "CREATE INDEX IF NOT EXISTS ix_outbox_events_pending ON outbox_events (status, created_at) WHERE status = 'pending'",
        "CREATE INDEX IF NOT EXISTS ix_usage_logs_user ON usage_logs (user_id, created_at)",
        # ── Hot-path chat / library / queue indexes (migrations/0022) ──
        # Keep definitions identical to 0022 so CREATE INDEX IF NOT EXISTS
        # does not leave a weaker SQLite-only definition that blocks Postgres.
        "CREATE INDEX IF NOT EXISTS ix_messages_conversation_created ON messages (conversation_id, created_at)",
        "CREATE INDEX IF NOT EXISTS ix_conversations_user_updated ON conversations (user_id, updated_at DESC)",
        "CREATE INDEX IF NOT EXISTS ix_conversations_user_project ON conversations (user_id, project_id)",
        "CREATE INDEX IF NOT EXISTS ix_conversations_user_file ON conversations (user_id, file_id)",
        "CREATE INDEX IF NOT EXISTS ix_conversations_project ON conversations (project_id)",
        "CREATE INDEX IF NOT EXISTS ix_conversations_file ON conversations (file_id)",
        "CREATE INDEX IF NOT EXISTS ix_files_user_project ON files (user_id, project_id)",
        "CREATE INDEX IF NOT EXISTS ix_files_conversation ON files (conversation_id)",
        "CREATE INDEX IF NOT EXISTS ix_upload_jobs_type_status ON upload_jobs (job_type, status)",
        "CREATE INDEX IF NOT EXISTS ix_upload_jobs_status_created ON upload_jobs (status, created_at)",
        "CREATE INDEX IF NOT EXISTS ix_upload_jobs_file_type ON upload_jobs (file_id, job_type)",
        "CREATE INDEX IF NOT EXISTS ix_outbox_events_status_created ON outbox_events (status, created_at)",
        "CREATE INDEX IF NOT EXISTS ix_projects_user ON projects (user_id)",
        "CREATE INDEX IF NOT EXISTS ix_memories_user ON memories (user_id)",
        "CREATE INDEX IF NOT EXISTS ix_memories_user_project ON memories (user_id, project_id)",
        "CREATE INDEX IF NOT EXISTS ix_memories_project_status ON memories (project_id, status)",
        "CREATE INDEX IF NOT EXISTS ix_memories_claim_hash ON memories (user_id, project_id, kind, claim_hash)",
        "CREATE INDEX IF NOT EXISTS ix_citations_user ON citations (user_id)",
        # ── Library search (Phase 1.5) ───────────────────────────────────
        "CREATE INDEX IF NOT EXISTS ix_files_user_created ON files (user_id, created_at)",
        "CREATE INDEX IF NOT EXISTS ix_files_user_year ON files (user_id, year)",
        "CREATE INDEX IF NOT EXISTS ix_files_user_reading ON files (user_id, reading_status)",
        "CREATE INDEX IF NOT EXISTS ix_files_user_meta ON files (user_id, meta_status)",
        "CREATE INDEX IF NOT EXISTS ix_files_user_doi ON files (user_id, doi)",
    ):
        try:
            with engine.begin() as conn:
                conn.execute(sqltext(stmt))
        except Exception:
            pass


if _ORM_BOOTSTRAP_SCHEMA:
    ensure_columns()
else:
    logging.getLogger(__name__).info(
        "Skipping ORM create_all/ensure_columns in production "
        "(schema owned by migrations/*.sql)."
    )

# ------------------------------------------------------------------ auth
oauth = OAuth(app)
google = oauth.register(
    name="google",
    client_id=GOOGLE_CLIENT_ID,
    client_secret=GOOGLE_CLIENT_SECRET,
    server_metadata_url="https://accounts.google.com/.well-known/openid-configuration",
    client_kwargs={"scope": "openid email profile"},
)


def login_required(f):
    @wraps(f)
    def wrapper(*args, **kwargs):
        if "user_id" not in session:
            if request.path.startswith("/api/"):
                return jsonify({"error": "not_authenticated"}), 401
            return redirect(url_for("login_page"))
        return f(*args, **kwargs)

    return wrapper


_LOGIN_ERRORS = {
    "not_invited": "Access denied — this account is not allowed to sign in.",
    "oauth_email": "Could not read your Google account email.",
    "oauth_failed": "Google sign-in failed or expired. Please try again.",
    "oauth_busy": "Sign-in is temporarily busy. Wait a moment and try again.",
    "login_failed": "Sign-in failed. Please try again.",
}


def _ecosystem_catalog_for_landing():
    """Same SoT as Settings Integrations / public catalog API."""
    from backend.ecosystem.catalog import public_catalog

    return public_catalog()


def _vite_ecosystem_island_tags():
    """Script/link tags for the Research Ecosystem Icon Cloud React island.

    Prod: hashed assets from frontend/dist/.vite/manifest.json (ecosystem.html entry).
    Dev: VITE_DEV_SERVER / FLASK_DEBUG / DHUND_VITE_ECOSYSTEM → Vite HMR on :5173.
    """
    force_vite = (os.environ.get("DHUND_VITE_ECOSYSTEM") or "").strip() == "1"
    vite_dev = (os.environ.get("VITE_DEV_SERVER") or "").strip().rstrip("/")
    flask_debug = (os.environ.get("FLASK_DEBUG") or "").strip() in {"1", "true", "True"}

    if force_vite or vite_dev:
        base = vite_dev or "http://localhost:5173"
        return {
            "mode": "dev",
            "scripts": [
                f"{base}/@vite/client",
                f"{base}/src/marketing/ecosystem-mount.tsx",
            ],
            "styles": [],
        }

    dist = os.path.join(os.path.dirname(os.path.abspath(__file__)), "frontend", "dist")
    for rel in (".vite/manifest.json", "manifest.json"):
        manifest_path = os.path.join(dist, rel)
        if not os.path.isfile(manifest_path):
            continue
        try:
            with open(manifest_path, encoding="utf-8") as f:
                manifest = json.load(f)
        except Exception:
            continue
        entry = manifest.get("ecosystem.html") or manifest.get("src/marketing/ecosystem-mount.tsx")
        if not isinstance(entry, dict):
            continue
        file_path = (entry.get("file") or "").lstrip("/")
        if not file_path:
            continue

        # Collect CSS from the entry and imported chunks (Vite often puts CSS on chunks).
        styles: list[str] = []
        seen_css: set[str] = set()
        queue = [entry]
        visited: set[int] = set()
        while queue:
            node = queue.pop()
            if id(node) in visited:
                continue
            visited.add(id(node))
            for c in node.get("css") or []:
                href = f"/{str(c).lstrip('/')}"
                if href not in seen_css:
                    seen_css.add(href)
                    styles.append(href)
            for imp in node.get("imports") or []:
                child = manifest.get(imp)
                if isinstance(child, dict):
                    queue.append(child)

        return {"mode": "prod", "scripts": [f"/{file_path}"], "styles": styles}

    if flask_debug:
        base = "http://localhost:5173"
        return {
            "mode": "dev",
            "scripts": [
                f"{base}/@vite/client",
                f"{base}/src/marketing/ecosystem-mount.tsx",
            ],
            "styles": [],
        }

    return {"mode": "missing", "scripts": [], "styles": []}


def _render_login_landing():
    """Public Research OS landing + sign-in (login.html). Shared by / and /login."""
    raw_error = request.args.get("error")
    if raw_error and raw_error in _LOGIN_ERRORS:
        error_msg = _LOGIN_ERRORS[raw_error]
    elif raw_error and len(raw_error) > 24 and " " in raw_error:
        error_msg = raw_error  # legacy full-sentence query params
    elif raw_error:
        error_msg = _LOGIN_ERRORS.get("login_failed")
    else:
        error_msg = None

    return render_template(
        "login.html",
        active="signin",
        oauth_ready=bool(GOOGLE_CLIENT_ID and GOOGLE_CLIENT_SECRET),
        closed_beta=CLOSED_BETA,
        error=error_msg,
        verified=request.args.get("verified") == "1",
        app_base_url=APP_BASE_URL,
        ecosystem_catalog=_ecosystem_catalog_for_landing(),
        ecosystem_island=_vite_ecosystem_island_tags(),
    )


@app.route("/login")
def login_page():
    if "user_id" in session:
        return redirect("/")

    # ── DEV_AUTO_LOGIN bypass ─────────────────────────────────────────────────
    # Development only. If this is set on Railway without a healthy schema,
    # the old path 500'd the whole /login page — catch and fall through.
    if DEV_AUTO_LOGIN and not IS_PRODUCTION:
        db = SessionLocal()
        try:
            dev_email = "dev@localhost"
            user = db.execute(select(User).where(User.email == dev_email)).scalar_one_or_none()
            if not user:
                user = User(email=dev_email, name="Dev User", picture="", auth_provider="dev")
                db.add(user)
                db.commit()
            session["user_id"] = user.id
            session["user_email"] = user.email
            session["session_version"] = int(getattr(user, "session_version", 0) or 0)
            access, refresh = create_jwt(
                user.id, session_version=int(getattr(user, "session_version", 0) or 0)
            )
            session["jwt"] = {"access": access, "refresh": refresh}
            mark_session_login(session)
            _record_user_login(user.id)
            return redirect("/")
        except Exception:
            logging.getLogger(__name__).exception(
                "DEV_AUTO_LOGIN failed — showing login page. "
                "On Railway: delete DEV_AUTO_LOGIN and set FLASK_ENV=production."
            )
            try:
                db.rollback()
            except Exception:
                pass
        finally:
            db.close()
    # ─────────────────────────────────────────────────────────────────────────

    # Preserve legacy query params on the new auth UI.
    qs = request.query_string.decode("utf-8") if request.query_string else ""
    target = "/auth/sign-in"
    if qs:
        target = f"{target}?{qs}"
    return redirect(target)


def _marketing_ctx(active: str, **extra):
    return {"active": active, "closed_beta": CLOSED_BETA, **extra}


def _serve_spa_index():
    dist = os.path.join(os.path.dirname(os.path.abspath(__file__)), "frontend", "dist")
    index_path = os.path.join(dist, "index.html")
    if not os.path.exists(index_path):
        return ("Frontend build not found — run `npm run build` in frontend/.", 501)
    return send_from_directory(dist, "index.html")


@app.route("/")
def root():
    """Signed-in users get the SPA; logged-out visitors get the login landing (retired marketing home)."""
    if "user_id" in session:
        return _serve_spa_index()
    return _render_login_landing()


@app.route("/product")
def marketing_product():
    return render_template("marketing/product.html", **_marketing_ctx("product"))


@app.route("/how-it-works")
def marketing_how_it_works():
    return render_template("marketing/how_it_works.html", **_marketing_ctx("how"))


@app.route("/research")
def marketing_research():
    return render_template("marketing/research.html", **_marketing_ctx("research"))


@app.route("/research/literature-review")
def marketing_guide_literature_review():
    return render_template(
        "marketing/guides/literature_review.html", **_marketing_ctx("research")
    )


@app.route("/research/evidence-backed-writing")
def marketing_guide_evidence_backed_writing():
    return render_template(
        "marketing/guides/evidence_backed_writing.html", **_marketing_ctx("research")
    )


@app.route("/research/ai-hallucinations")
def marketing_guide_ai_hallucinations():
    return render_template(
        "marketing/guides/ai_hallucinations.html", **_marketing_ctx("research")
    )


@app.route("/research/systematic-reviews")
def marketing_guide_systematic_reviews():
    return render_template(
        "marketing/guides/systematic_reviews.html", **_marketing_ctx("research")
    )


@app.route("/early-access")
def marketing_early_access():
    return render_template("marketing/early_access.html", **_marketing_ctx("access"))


@app.route("/pricing")
def marketing_pricing():
    return render_template("marketing/pricing.html", **_marketing_ctx("pricing"))


def _oauth_redirect_uri() -> str:
    """Absolute /auth/callback URL for Google OAuth.

    Railway terminates TLS at the edge; without an explicit https base,
    Authlib/url_for can emit http://… which Google rejects (redirect_uri_mismatch).
    """
    base = (APP_BASE_URL or "").strip().rstrip("/")
    railway = (os.environ.get("RAILWAY_PUBLIC_DOMAIN") or "").strip()
    if railway and "://" not in railway:
        railway = f"https://{railway}"

    # Upgrade accidental http:// production bases to https://
    if base.startswith("http://") and not any(
        h in base for h in ("localhost", "127.0.0.1")
    ):
        base = "https://" + base[len("http://") :]

    if base.startswith("https://") or (
        base.startswith("http://") and any(h in base for h in ("localhost", "127.0.0.1"))
    ):
        return f"{base}/auth/callback"

    if railway.startswith("https://"):
        return f"{railway.rstrip('/')}/auth/callback"

    host = (request.host or "").split(":")[0]
    if host in {"localhost", "127.0.0.1"} or host.startswith("localhost"):
        return url_for("auth_callback", _external=True)
    return f"https://{request.host}/auth/callback"


@app.route("/auth/google")
@limiter.limit("30 per minute")
def auth_google():
    # Mobile browsers (esp. Chrome/Android) drop non-permanent session cookies
    # across the Google redirect more often than desktop. Make the cookie
    # persistent before Authlib stores OAuth state in the session.
    session.permanent = True
    session["_oauth_flow"] = "google"
    return google.authorize_redirect(_oauth_redirect_uri())


@app.route("/auth/")
@app.route("/auth")
def auth_root_redirect():
    """Avoid bare /auth/ dead-ends (mobile truncates / bookmarks)."""
    return redirect(url_for("login_page"))


@app.route("/auth/callback")
@limiter.limit("60 per minute")
def auth_callback():
    """Google OAuth callback.

    Always redirect away from /auth/callback after handling. Leaving the
    user on this URL with a one-time ?code=… makes mobile refresh / retry
    hit authorize_access_token again → uncaught Internal Server Error.
    """
    try:
        token = google.authorize_access_token()
    except Exception:
        logging.getLogger(__name__).exception("OAuth token exchange failed")
        return redirect(url_for("login_page", error="oauth_failed"))

    info = token.get("userinfo") or {}
    email = (info.get("email") or "").lower()
    if not email:
        return redirect(url_for("login_page", error="oauth_email"))

    try:
        ok, reason = _signup_allowed(email)
    except Exception:
        logging.getLogger(__name__).exception("signup gate failed for oauth")
        return redirect(url_for("login_page", error="login_failed"))

    if not ok:
        try:
            log_security_event("oauth_denied", email=email, reason=reason)
            _ops_events.record(
                "oauth_denied", email=email, reason=reason, ip=request.remote_addr or ""
            )
        except Exception:
            logging.getLogger(__name__).exception("oauth_denied logging failed")
        return redirect(url_for("login_page", error="not_invited"))

    db = SessionLocal()
    try:
        user = db.execute(select(User).where(User.email == email)).scalar_one_or_none()
        if not user:
            user = User(email=email, auth_provider="google")
            db.add(user)
            db.flush()  # need user.id before invite_accepted / JWT
        user.name = info.get("name") or email
        user.picture = info.get("picture") or ""
        # Google email is verified by the provider
        user.email_verified = True
        user.email_verified_at = user.email_verified_at or datetime.now(timezone.utc)
        if not user.status or user.status == "pending_verification":
            user.status = "active"
        if reason == "invite":
            _ops_invites.consume_invite_for_email(email)
            _ops_events.record("invite_accepted", user_id=user.id, email=email)
        db.commit()
        session["user_id"] = user.id
        session["user_email"] = user.email
        session["session_version"] = int(getattr(user, "session_version", 0) or 0)
        access, refresh = create_jwt(
            user.id, session_version=int(getattr(user, "session_version", 0) or 0)
        )
        session["jwt"] = {"access": access, "refresh": refresh}
        mark_session_login(session)
        _record_user_login(user.id)
        try:
            _ops_events.record("google_login", user_id=user.id, email=email)
        except Exception:
            pass
    except Exception:
        logging.getLogger(__name__).exception("OAuth user upsert failed")
        try:
            db.rollback()
        except Exception:
            pass
        return redirect(url_for("login_page", error="login_failed"))
    finally:
        db.close()
    return redirect("/")


@app.route("/logout")
def logout():
    uid = session.get("user_id")
    session.clear()
    if uid:
        log_security_event("logout", user_id=uid)
    return redirect(url_for("login_page"))


@app.route("/api/dev-login", methods=["POST"])
@limiter.limit("10 per hour")
def dev_login():
    """Sign in as the local dev user without Google OAuth.

    Requires DEV_AUTO_LOGIN and a non-production environment.
    Returns 403 when either gate fails (Phase 2 / F2.2).
    """
    if not (DEV_AUTO_LOGIN and not IS_PRODUCTION):
        return (
            jsonify(
                {
                    "error": "dev_login_disabled",
                    "detail": "Dev login requires DEV_AUTO_LOGIN=1 and a non-production environment.",
                }
            ),
            403,
        )
    db = SessionLocal()
    try:
        dev_email = "dev@localhost"
        user = db.execute(select(User).where(User.email == dev_email)).scalar_one_or_none()
        if not user:
            user = User(email=dev_email, name="Dev User", picture="")
            db.add(user)
            db.commit()
        session["user_id"] = user.id
        session["user_email"] = user.email
        session["session_version"] = int(getattr(user, "session_version", 0) or 0)
        access, refresh = create_jwt(
            user.id, session_version=int(getattr(user, "session_version", 0) or 0)
        )
        session["jwt"] = {"access": access, "refresh": refresh}
        mark_session_login(session)
        _record_user_login(user.id)
        log_security_event("dev_login", user_id=user.id)
        return jsonify({"ok": True, "user_id": user.id})
    finally:
        db.close()


# Magic-link blueprint is registered after closed-beta ops services are
# wired (needs signup_allowed_fn) — see below near create_ops_blueprint.


from backend.storage import get_storage_backend
from backend.upload.routes import create_documents_blueprint

# POST /api/documents/upload — Bearer-JWT-authenticated upload entry point
# alongside the existing session-based POST /api/files. Reuses UserFile/
# UploadJob/OutboxEvent (no new Document model) and quota_service — see
# backend/upload/routes.py's module docstring for the full reasoning.
# Registration itself is deferred to just after the Prompt Engine wiring
# further down (model_router/PromptExecution don't exist yet at this
# point in the file) — see there.
from backend.upload.validation import MAX_DOCUMENT_UPLOAD_MB

# Flask rejects an over-limit request body before any route code runs —
# this route's own limit must not be shadowed by the app-wide one.
app.config["MAX_CONTENT_LENGTH"] = max(MAX_FILE_MB, MAX_DOCUMENT_UPLOAD_MB) * 1024 * 1024

# POST /api/uploads/bulk, GET /api/uploads/batch/<id>/status — bulk upload,
# writing to UploadBatch (schema-only until now, see its own docstring)
# via the same UserFile/UploadJob/OutboxEvent primitives as the single-file
# route above. A batch can total far more than one file's MAX_CONTENT_LENGTH,
# so the app-wide cap has to grow to fit MAX_BATCH_SIZE files at once.
from backend.upload.bulk import MAX_BATCH_SIZE, create_bulk_upload_blueprint

app.config["MAX_CONTENT_LENGTH"] = max(
    app.config["MAX_CONTENT_LENGTH"], MAX_BATCH_SIZE * MAX_DOCUMENT_UPLOAD_MB * 1024 * 1024
)
app.register_blueprint(
    create_bulk_upload_blueprint(
        SessionLocal=SessionLocal,
        UserFile=UserFile,
        UploadBatch=UploadBatch,
        UploadJob=UploadJob,
        OutboxEvent=OutboxEvent,
        quota_service=quota_service,
        storage_backend=get_storage_backend(),
        limiter=limiter,
    )
)

# GET /api/documents/search, POST /api/rag — Bearer-JWT counterparts to
# the existing session-based POST /api/search (below), same relationship
# as /api/documents/upload has to /api/files. Search the exact same
# Chunk.embedding data /api/search already uses for papers — see
# backend/search/routes.py's module docstring for why this isn't a
# second search engine. Registration itself is deferred to just after the
# Prompt Engine wiring further down (get_prompt_builder/model_router/
# PromptExecution don't exist yet at this point in the file) — see there.
from backend.search.routes import create_search_blueprint

# Unified user lookup — session first, Bearer JWT second, None if
# neither. A helper for routes that should accept either auth method;
# @login_required (session-only) and @jwt_required() (JWT-only, from
# auth/decorators.py) are both untouched and still the right choice for
# routes that should accept exactly one.
get_current_user = create_get_current_user(SessionLocal, User)


def _jwt_session_version_check(claims, identity):
    """Reject JWTs after logout-all / password reset (Phase 2 / F2.1)."""
    try:
        user_id = int(identity)
    except (TypeError, ValueError):
        return jsonify({"error": "token_revoked", "detail": "reauthenticate"}), 401
    db = SessionLocal()
    try:
        user = db.get(User, user_id)
        if user is None:
            return jsonify({"error": "token_revoked", "detail": "reauthenticate"}), 401
        current_sv = int(getattr(user, "session_version", 0) or 0)
        if not session_version_matches(claims or {}, current_sv):
            return jsonify({"error": "token_revoked", "detail": "reauthenticate"}), 401
    finally:
        db.close()
    return None


set_jwt_session_version_checker(_jwt_session_version_check)


# ══════════════════════════════════════════════════════════════════════════
# JWT — extra capability alongside session/OAuth login, for API/programmatic
# clients that can't hold a browser cookie. Neither route below touches the
# session-based login flow above.
#
# Not built: a headless OAuth-code-exchange endpoint (a client does its own
# Google OAuth, hands us the raw code, we exchange it for a JWT without a
# browser session ever existing). Nothing in this project is a headless
# client today — no CLI, no mobile app — building that integration surface
# ahead of an actual consumer is exactly the kind of speculative work this
# project's own docs (upload-architecture.md §11 on Feature Flags, among
# others) have consistently deferred until there's a real one.
# ══════════════════════════════════════════════════════════════════════════


@app.route("/api/auth/jwt")
@login_required
def get_session_jwt():
    """For a client that already has the browser session (just logged in
    via Google OAuth) and wants a portable Bearer token instead. Returns
    the token minted at login if it's still valid; mints a fresh one
    (access tokens are short-lived by design — 15 min default) if
    the session doesn't have one yet or it's expired — a client calling
    this should never get back something already broken."""
    stored = session.get("jwt")
    if stored:
        try:
            claims = decode_jwt(stored["access"])
            # Reject cached JWT if logout-all bumped session_version.
            db = SessionLocal()
            try:
                user = db.get(User, session["user_id"])
                cur = int(getattr(user, "session_version", 0) or 0) if user else -1
            finally:
                db.close()
            if session_version_matches(claims, cur):
                return jsonify({"access_token": stored["access"], "refresh_token": stored["refresh"]})
        except JWTError:
            pass  # expired/invalid — fall through and mint a fresh pair
    db = SessionLocal()
    try:
        user = db.get(User, session["user_id"])
        sv = int(getattr(user, "session_version", 0) or 0) if user else 0
    finally:
        db.close()
    access, refresh = create_jwt(session["user_id"], session_version=sv)
    session["jwt"] = {"access": access, "refresh": refresh}
    return jsonify({"access_token": access, "refresh_token": refresh})


@app.route("/api/auth/token", methods=["POST"])
@limiter.limit("30 per minute")
def refresh_jwt():
    """Exchange a refresh token for a new access token. The email+password
    variant of this endpoint doesn't apply — this app has no password-based
    accounts, only Google OAuth — so this covers just the refresh_token
    grant, which every client holding a JWT eventually needs regardless of
    how the original token was issued.

    Phase 2 / F2.1: refresh tokens must carry a matching ``sv`` claim or
    they are treated as revoked (logout-all / password reset).
    """
    data = request.get_json(silent=True) or {}
    refresh_token = data.get("refresh_token")
    if not refresh_token:
        return jsonify({"error": "refresh_token_required"}), 400
    try:
        claims = decode_jwt(refresh_token)
    except JWTError:
        return jsonify({"error": "invalid_refresh_token"}), 401
    if claims.get("type") != "refresh":
        return (
            jsonify({"error": "invalid_refresh_token", "detail": "not a refresh token"}),
            401,
        )
    try:
        user_id = int(claims.get("sub"))
    except (TypeError, ValueError):
        return jsonify({"error": "invalid_refresh_token"}), 401

    db = SessionLocal()
    try:
        user = db.get(User, user_id)
        if not user:
            return jsonify({"error": "invalid_refresh_token"}), 401
        current_sv = int(getattr(user, "session_version", 0) or 0)
        if not session_version_matches(claims, current_sv):
            log_security_event(
                "jwt_refresh_revoked",
                user_id=user_id,
                reason="session_version_mismatch",
            )
            return jsonify({"error": "token_revoked", "detail": "reauthenticate"}), 401
        access, new_refresh = create_jwt(user_id, session_version=current_sv)
    finally:
        db.close()

    return jsonify({"access_token": access, "refresh_token": new_refresh})


@app.route("/robots.txt")
def robots():
    return send_from_directory("static", "robots.txt")


# ══════════════════════════════════════════════════════════════════════════
# AI layer (backend/ai/) — prompt registry, multi-provider model registry,
# cost ledger. get_prompt_registry()/get_model_registry() are request-scoped
# factories, NOT instances built once at startup: PromptRegistry/ModelRegistry
# each hold one open SQLAlchemy Session for their lifetime (their own
# constructors take a Session instance, not a factory) — sharing one across
# concurrent requests would be a real bug, not a style issue. Every other
# route in this file already opens/closes its own `db = SessionLocal()` per
# request; these two follow that same convention rather than introducing
# Flask's g/app.extensions for just this one feature. CostLedger is the one
# genuine startup-time singleton here — it holds no session at all
# (estimate_cost() is pure; log() takes a session as a call argument).
#
# prompt_versions/pipeline_versions/model_registry_cost_ledger only exist
# under backend/ai's own private declarative Bases, never server.py's — see
# prompt_registry.py's and model_registry.py's own docstrings for why. This
# project has no Alembic (migrations/*.sql are hand-written, run via
# run_migrations.py — 00-constitution.md), so there's no autogenerate step
# to feed; what actually matters is these tables existing before anything
# queries them. checkfirst=True is a no-op everywhere the real Postgres
# migration already ran — it only creates anything on a fresh SQLite dev DB
# (verified: prompt_versions/pipeline_versions didn't exist there before this).
# ══════════════════════════════════════════════════════════════════════════
from backend.ai import CostLedger, ModelError, ModelRegistry, PromptRegistry, PromptVersion, TemplateError
from backend.ai.model_registry import CostLedgerEntry as _CostLedgerEntry
from backend.ai.model_registry import _Base as _ai_model_base
from backend.ai.prompt_registry import _Base as _ai_prompt_base

if _ORM_BOOTSTRAP_SCHEMA:
    _ai_prompt_base.metadata.create_all(engine, checkfirst=True)
    _ai_model_base.metadata.create_all(engine, checkfirst=True)

_cost_ledger = CostLedger(_CostLedgerEntry)


def get_prompt_registry(db_session):
    return PromptRegistry(db_session)


def get_model_registry(db_session):
    return ModelRegistry(db_session)


def get_cost_ledger():
    return _cost_ledger


# ── Closed-beta ops services (AI gate, kill switch, invites, password auth) ─
from security.ops import (
    AiAccessGate,
    BetaMetricsService,
    InviteService,
    PasswordAuthService,
    SecurityEventStore,
    SystemSettingsService,
    record_last_login,
)
from security.ops.estimates import estimate_chat_tokens
from security.ops.invites import signup_allowed
from security.ops.routes import create_ops_blueprint

_ops_events = SecurityEventStore(SessionLocal, SecurityEvent, log_fn=log_security_event)
_ops_settings = SystemSettingsService(SessionLocal, SystemSetting)
_ops_invites = InviteService(SessionLocal, InviteToken)
_ops_password = PasswordAuthService(
    SessionLocal,
    User,
    EmailVerificationToken,
    PasswordResetToken,
    EmailChangeToken=EmailChangeToken,
    email_service=email_service,
    app_base_url=APP_BASE_URL,
    events=_ops_events,
    auth_from=AUTH_EMAIL_FROM,
    noreply_from=NOREPLY_EMAIL_FROM,
    invite_service=_ops_invites,
)
entitlement_service = EntitlementService(
    SessionLocal=SessionLocal,
    User=User,
    StorageUsage=StorageUsage,
    UsageLog=UsageLog,
    select=select,
    quota_service=quota_service,
    settings=_ops_settings,
    events=_ops_events,
)
feature_flag_service = FeatureFlagService(SessionLocal, FeatureFlag, select)
ai_gate = AiAccessGate(
    SessionLocal=SessionLocal,
    User=User,
    settings=_ops_settings,
    quota_service=quota_service,
    events=_ops_events,
    select=select,
    entitlements=entitlement_service,
)
_beta_metrics = BetaMetricsService(
    SessionLocal,
    User,
    Project,
    UserFile,
    DerivedAnalysis,
    Memory,
    select,
)


def _record_user_login(user_id: int) -> None:
    record_last_login(SessionLocal, User, user_id)


def _signup_allowed(email: str) -> tuple[bool, str]:
    return signup_allowed(
        email,
        allowed_emails=ALLOWED_EMAILS,
        invite_service=_ops_invites,
        require_invite=BETA_INVITE_ONLY,
    )


_ops_password.signup_allowed_fn = _signup_allowed

from auth.magic_link import create_magic_link_blueprint

app.register_blueprint(
    create_magic_link_blueprint(
        secret_key=app.secret_key,
        limiter=limiter,
        email_service=email_service,
        SessionLocal=SessionLocal,
        User=User,
        select=select,
        ALLOWED_EMAILS=ALLOWED_EMAILS,
        APP_BASE_URL=APP_BASE_URL,
        create_jwt=create_jwt,
        log_security_event=log_security_event,
        MagicLinkToken=MagicLinkToken,
        signup_allowed_fn=_signup_allowed,
        on_user_created=lambda user, email: (
            _ops_invites.consume_invite_for_email(email),
            _ops_events.record("invite_accepted", user_id=user.id, email=email),
            _record_user_login(user.id),
        ),
        record_last_login_fn=_record_user_login,
    )
)


from auth.decorators import create_admin_required
from backend.ai.analytics import PromptAnalytics
from backend.ai.domain_registry import DomainRegistry
from backend.ai.gateway import AIGateway, validate_registry
from backend.ai.memory_engine import MemoryEngine
from backend.ai.model_router import ModelRouter
from backend.ai.persona_engine import PersonaEngine
from backend.ai.prompt_builder import CHAT_SYSTEM_FALLBACK, PromptBuilder

# ══════════════════════════════════════════════════════════════════════════
# Prompt Engine (docs/prompt-engine-architecture.md) — SystemPromptManager,
# PersonaEngine, MemoryEngine, ModelRouter, PromptBuilder, built on top of
# the AI layer just above. Persona/PromptExecution live under
# prompt_registry.py's own private Base too (see that module) — already
# covered by _ai_prompt_base.metadata.create_all() above, no separate
# create_all() call needed here.
# ══════════════════════════════════════════════════════════════════════════
from backend.ai.prompt_registry import Persona, PromptExecution
from backend.ai.system_prompt import SystemPromptManager
from backend.prompts.routes import create_prompts_blueprint

# ModelRouter is a genuine startup-time singleton — no DB session at all
# (see model_router.py's own docstring), same category as CostLedger
# above. task_name -> model string falls back through the SAME
# env-derived constants (UTILITY_MODEL/EMBED_MODEL/DEFAULT_MODEL) this
# app already reads, rather than inventing a fourth config surface.
# "rag" isn't in the task brief's own example dict (that dict was labeled
# "e.g.", not exhaustive) but backend/search/routes.py's rag_answer() is
# required to call get_model_for_task("rag") specifically — mapped to
# UTILITY_MODEL to preserve exactly what RAG already used before this
# task, unless a RAG_MODEL env var is set to override it.
model_router = ModelRouter(
    defaults={
        "chat": UTILITY_MODEL,
        "paper_analysis": UTILITY_MODEL,
        "rag": UTILITY_MODEL,
        "embedding": EMBED_MODEL,
        "_default": DEFAULT_MODEL,
    }
)
ai_model_gateway = AIGateway(model_router)
validate_registry()  # logs tier/model summary; pass openai_client to verify API access

# Also a genuine startup-time singleton — DomainRegistry is pure Python,
# no DB session, same category as ModelRouter (see that class's own
# comment just above, and domain_registry.py's own module docstring).
domain_registry = DomainRegistry()


# get_persona_engine/get_memory_engine/get_system_prompt_manager/
# get_prompt_builder are request-scoped factories, NOT startup
# singletons — same reasoning as get_prompt_registry/get_model_registry
# above: PersonaEngine/MemoryEngine/SystemPromptManager hold (and
# PromptBuilder transitively holds, via all three) one open Session for
# their lifetime; one shared instance across concurrent requests would
# mean concurrent requests sharing one session.
def get_persona_engine(db_session):
    return PersonaEngine(db_session, Persona)


def get_memory_engine(db_session):
    return MemoryEngine(db_session, Memory)


def get_system_prompt_manager(db_session):
    return SystemPromptManager(get_prompt_registry(db_session))


def get_prompt_builder(db_session):
    return PromptBuilder(
        system_prompt_manager=get_system_prompt_manager(db_session),
        persona_engine=get_persona_engine(db_session),
        memory_engine=get_memory_engine(db_session),
        prompt_registry=get_prompt_registry(db_session),
        SessionLocal=SessionLocal,
        Project=Project,
        domain_registry=domain_registry,
    )


def get_prompt_analytics(db_session):
    return PromptAnalytics(db_session, AIUsageLedger, ModelVersion)


admin_required = create_admin_required(SessionLocal, User)

app.register_blueprint(
    create_ops_blueprint(
        settings_service=_ops_settings,
        event_store=_ops_events,
        invite_service=_ops_invites,
        password_auth=_ops_password,
        ai_gate=ai_gate,
        quota_service=quota_service,
        entitlement_service=entitlement_service,
        feature_flag_service=feature_flag_service,
        beta_metrics=_beta_metrics,
        email_service=email_service,
        app_base_url=APP_BASE_URL,
        login_required=login_required,
        admin_required=admin_required,
        mark_session_login=mark_session_login,
        create_jwt=create_jwt,
        record_last_login_fn=_record_user_login,
        limiter=limiter,
        oauth_ready=bool(GOOGLE_CLIENT_ID and GOOGLE_CLIENT_SECRET),
        closed_beta=CLOSED_BETA,
    )
)

# Library Bridge — BibTeX/RIS + Connect Library + Collections + Sync (Phase 1b)
from backend.library.collections import CollectionService
from backend.library.routes import create_library_bridge_blueprint
from backend.library.service import LibraryImportService
from backend.library.sync import LibrarySyncService
from backend.writing.api.errors import WritingDomainError
from backend.writing.events import make_writing_event, publish_writing_event
from backend.writing.services import (
    build_version_conflict_payload,
    is_idempotent_replay,
    normalize_editor_kind,
    normalize_idempotency_key,
    normalize_status_filter,
    next_version_number,
    require_owned_document,
    require_owned_project,
)
from backend.writing.validation import ensure_transition_allowed
from backend.writing.validation.schemas import normalize_document_mutation
from backend.writing.services.logging import log_writing_metric
from backend.scholarly.crossref import enrich_file_from_doi as _enrich_file_from_doi

_collection_service = CollectionService(
    SessionLocal,
    LibraryCollection,
    LibraryCollectionPaper,
    UserFile,
    select,
)
_library_import = LibraryImportService(
    SessionLocal,
    UserFile,
    Project,
    select,
    enrich_file_from_doi=_enrich_file_from_doi,
    collection_service=_collection_service,
)
_library_sync = LibrarySyncService(
    SessionLocal,
    UserFile,
    LibraryConnection,
    LibrarySyncRun,
    select,
    _library_import,
    enrich_file_from_doi=_enrich_file_from_doi,
)
app.register_blueprint(
    create_library_bridge_blueprint(
        import_service=_library_import,
        sync_service=_library_sync,
        SessionLocal=SessionLocal,
        UserFile=UserFile,
        LibraryConnection=LibraryConnection,
        Project=Project,
        select_fn=select,
        login_required=login_required,
        app_base_url=APP_BASE_URL,
        enrich_file_from_doi=_enrich_file_from_doi,
        limiter=limiter,
        collection_service=_collection_service,
        storage=storage,
        enqueue_import=lambda db, uid, fid: _enqueue_job(db, uid, fid, "import"),
        enqueue_phase1=lambda db, uid, fid: _enqueue_job(db, uid, fid, "phase1_analysis"),
        # Deferred: _file_to_dict is defined later in this module.
        file_to_dict=lambda x: _file_to_dict(x),
        upload_dir=UPLOAD_DIR,
        max_file_mb=MAX_FILE_MB,
        allowed_extensions=None,  # attach route defaults to PDF
        LibrarySyncRun=LibrarySyncRun,
        token_secret_key=app.secret_key,
        UploadJob=UploadJob,
        OutboxEvent=OutboxEvent,
    )
)

from backend.ecosystem import create_integrations_catalog_blueprint

app.register_blueprint(
    create_integrations_catalog_blueprint(
        SessionLocal=SessionLocal,
        UserFile=UserFile,
        LibraryConnection=LibraryConnection,
        LibrarySyncRun=LibrarySyncRun,
        select_fn=select,
        login_required=login_required,
    )
)

app.register_blueprint(
    create_prompts_blueprint(
        SessionLocal=SessionLocal,
        PromptVersion=PromptVersion,
        Persona=Persona,
        PromptRegistry=PromptRegistry,
        PersonaEngine=PersonaEngine,
        get_prompt_builder=get_prompt_builder,
        login_required=login_required,
        admin_required=admin_required,
    )
)

# backend/upload's and backend/search's blueprint registrations,
# deferred to here (see the comments at their imports above) since both
# need model_router/PromptExecution (backend/upload also get_prompt_builder
# via backend/search), all defined just above.
app.register_blueprint(
    create_documents_blueprint(
        SessionLocal=SessionLocal,
        UserFile=UserFile,
        UploadBatch=UploadBatch,
        UploadJob=UploadJob,
        OutboxEvent=OutboxEvent,
        PaperAnalysis=PaperAnalysis,
        PromptExecution=PromptExecution,
        quota_service=quota_service,
        storage_backend=get_storage_backend(),
        model_router=model_router,
        ai_gateway=ai_model_gateway,
        get_prompt_builder=get_prompt_builder,
        domain_registry=domain_registry,
        AnalysisPipelineResult=AnalysisPipelineResult,
        limiter=limiter,
    )
)

app.register_blueprint(
    create_search_blueprint(
        SessionLocal=SessionLocal,
        UserFile=UserFile,
        Chunk=Chunk,
        get_prompt_builder=get_prompt_builder,
        model_router=model_router,
        ai_gateway=ai_model_gateway,
        PromptExecution=PromptExecution,
        limiter=limiter,
        ai_gate=ai_gate,
    )
)

# Sprint A — Project workspace hub (single read model). CRUD stays below
# until a later slice migrates those routes onto ProjectService.
from backend.projects import create_project_service
from backend.projects.research import create_project_research_service
from backend.projects.memory import create_memory_promotion_service
from backend.projects.routes import create_projects_blueprint
from backend.analysis_pipeline.summary import build_phase1_prompt_context

project_service = create_project_service(
    SessionLocal=SessionLocal,
    select=select,
    Project=Project,
    UserFile=UserFile,
    Note=Note,
    Memory=Memory,
    Conversation=Conversation,
    DerivedAnalysis=DerivedAnalysis,
    ProjectQuestion=ProjectQuestion,
    AnalysisPipelineResult=AnalysisPipelineResult,
)
# project_research_service + projects blueprint registered after responses_text().


def _parse_usage_date_range():
    """start_date/end_date query params, YYYY-MM-DD — naive datetimes on
    purpose: SQLite doesn't preserve tzinfo on created_at (the same
    limitation quotas/service.py's _ensure_reset already documents for
    this exact database), so comparing a caller-supplied aware datetime
    against a naive-in-SQLite column would be inconsistent; day-level
    granularity doesn't need finer precision than this anyway. Returns
    (None, None) on an unparseable date, for the route to turn into 400.
    Defaults to the trailing 30 days when neither param is given."""
    end_str = request.args.get("end_date")
    start_str = request.args.get("start_date")
    now = datetime.now(timezone.utc).replace(tzinfo=None)
    try:
        end = datetime.fromisoformat(end_str) if end_str else now
        start = datetime.fromisoformat(start_str) if start_str else end - timedelta(days=30)
    except ValueError:
        return None, None
    return start, end


# GET /api/prompt-usage[/by-prompt|/by-model|/by-user] — admin-only cost/
# usage analytics (docs/prompt-engine-architecture.md, backend/ai/analytics.py).
# No /by-project route: PromptAnalytics.get_usage_by_project() exists (the
# class-level requirement asked for it) but wasn't in this task's own
# route list, so it isn't surfaced as one yet — not an oversight.
@app.route("/api/prompt-usage")
@login_required
@admin_required
def prompt_usage_summary():
    start, end = _parse_usage_date_range()
    if start is None:
        return jsonify({"error": "invalid_date", "message": "start_date/end_date must be YYYY-MM-DD"}), 400
    db = SessionLocal()
    try:
        by_model = get_prompt_analytics(db).get_usage_by_model(start, end)
    finally:
        db.close()
    return jsonify(
        {
            "start_date": start.isoformat(),
            "end_date": end.isoformat(),
            "calls": sum(m["calls"] for m in by_model),
            "total_tokens": sum(m["total_tokens"] for m in by_model),
            "cost_usd": round(sum(m["cost_usd"] for m in by_model), 6),
            "models_used": len(by_model),
        }
    )


@app.route("/api/prompt-usage/by-prompt")
@login_required
@admin_required
def prompt_usage_by_prompt():
    start, end = _parse_usage_date_range()
    if start is None:
        return jsonify({"error": "invalid_date", "message": "start_date/end_date must be YYYY-MM-DD"}), 400
    db = SessionLocal()
    try:
        rows = get_prompt_analytics(db).get_usage_by_prompt(start, end)
    finally:
        db.close()
    return jsonify({"start_date": start.isoformat(), "end_date": end.isoformat(), "prompts": rows})


@app.route("/api/prompt-usage/by-model")
@login_required
@admin_required
def prompt_usage_by_model():
    start, end = _parse_usage_date_range()
    if start is None:
        return jsonify({"error": "invalid_date", "message": "start_date/end_date must be YYYY-MM-DD"}), 400
    db = SessionLocal()
    try:
        rows = get_prompt_analytics(db).get_usage_by_model(start, end)
    finally:
        db.close()
    return jsonify({"start_date": start.isoformat(), "end_date": end.isoformat(), "models": rows})


@app.route("/api/prompt-usage/by-user")
@login_required
@admin_required
def prompt_usage_by_user():
    start, end = _parse_usage_date_range()
    if start is None:
        return jsonify({"error": "invalid_date", "message": "start_date/end_date must be YYYY-MM-DD"}), 400
    db = SessionLocal()
    try:
        rows = get_prompt_analytics(db).get_usage_by_user(start, end)
    finally:
        db.close()
    return jsonify({"start_date": start.isoformat(), "end_date": end.isoformat(), "users": rows})


@app.route("/metrics")
def metrics():
    """Prometheus scrape target. Gated by METRICS_TOKEN (Bearer) or
    loopback-only when the token is unset — see security/metrics_access.py.
    Set METRICS_ALLOW_UNAUTHENTICATED=1 only for local/dev open scrapes."""
    allowed, reason = check_metrics_access(
        authorization=request.headers.get("Authorization"),
        remote_addr=request.remote_addr,
        environ=os.environ,
    )
    if not allowed:
        log_security_event(
            "metrics_access_denied",
            reason=reason,
            remote=request.remote_addr or "",
            path=request.path,
        )
        return jsonify({"error": "unauthorized", "message": "Metrics endpoint requires authentication"}), 401

    db = SessionLocal()
    try:
        for status in ("pending", "running", "failed", "done"):
            count = db.execute(
                select(func.count()).select_from(UploadJob).where(UploadJob.status == status)
            ).scalar_one()
            UPLOAD_QUEUE_LENGTH.labels(status=status).set(count)
    finally:
        db.close()
    return Response(render_metrics(), mimetype="text/plain; version=0.0.4")


@app.route("/api/worker/health")
def worker_health():
    """Unauthenticated on purpose — an ops liveness check (uptime monitor,
    orchestrator probe), not a user-facing route, same class as a plain
    /healthz. Reports on worker.py specifically, not this Flask process:
    server.py being up says nothing about whether the separate worker.py
    process is still polling upload_jobs."""
    db = SessionLocal()
    try:
        hb = db.get(WorkerHeartbeat, 1)
    finally:
        db.close()

    if hb is None:
        return (
            jsonify(
                {
                    "status": "unknown",
                    "message": "worker has not reported in since this deploy — "
                    "either it has never run, or it started before this "
                    "endpoint existed",
                }
            ),
            503,
        )

    last_seen = hb.last_seen_at
    if last_seen.tzinfo is None:
        last_seen = last_seen.replace(tzinfo=timezone.utc)
    age_seconds = (datetime.now(timezone.utc) - last_seen).total_seconds()
    healthy = age_seconds <= WORKER_HEALTH_THRESHOLD_SECONDS

    return jsonify(
        {
            "status": "ok" if healthy else "down",
            "last_seen_at": hb.last_seen_at.isoformat(),
            "age_seconds": round(age_seconds, 1),
            "threshold_seconds": WORKER_HEALTH_THRESHOLD_SECONDS,
        }
    ), (200 if healthy else 503)


@app.route("/api/health/providers")
def providers_health_endpoint():
    """Ops check for scholarly providers — unauthenticated like worker health.

    Returns per-provider status (healthy | circuit_open | disabled),
    24h cache_hit_rate, and in-process bulkhead stats.
    """
    from backend.scholarly import providers_health
    db = SessionLocal()
    try:
        payload = providers_health(db)
        return jsonify(payload), 200
    except Exception as exc:
        app.logger.warning("providers_health failed: %s", exc)
        return jsonify({
            "crossref": "unknown",
            "openalex": "unknown",
            "semantic_scholar": "unknown",
            "cache_hit_rate": 0.0,
            "error": "health_unavailable",
        }), 503
    finally:
        db.close()


@app.route("/api/ai/prompts")
@login_required
def list_ai_prompts():
    db = SessionLocal()
    try:
        registry = get_prompt_registry(db)
        prompts = registry.list_prompts()
        return jsonify(
            {
                "prompts": [
                    {
                        "name": p.name,
                        "version": p.version,
                        "template": p.template,
                        "is_active": p.is_active,
                        "created_at": p.created_at.isoformat() if p.created_at else None,
                    }
                    for p in prompts
                ]
            }
        )
    finally:
        db.close()


@app.route("/api/ai/test", methods=["POST"])
@login_required
def test_ai_call():
    """Dev-only: exercises ModelRegistry.call() directly against a real
    provider. Gated on IS_PRODUCTION — "optional, for dev" shouldn't mean
    an unrestricted call-any-model-with-any-prompt endpoint shipped without
    a guard; that's a real cost/abuse surface, not just a nicety to skip."""
    if IS_PRODUCTION:
        return jsonify({"error": "disabled_in_production"}), 403

    data = request.get_json(silent=True) or {}
    model = data.get("model") or DEFAULT_MODEL
    message = data.get("message") or "Say hello in one short sentence."

    db = SessionLocal()
    try:
        registry = get_model_registry(db)
        result = registry.call(
            model,
            [{"role": "user", "content": message}],
            user_id=session["user_id"],
            max_tokens=data.get("max_tokens", 100),
        )
        return jsonify(result)
    except ModelError as exc:
        return jsonify({"error": "model_call_failed", "message": str(exc)}), 502
    finally:
        db.close()


# ------------------------------------------------------------------ text extraction / chunking / embeddings

# extract_text(path, mime, name) -> str: '' = no readable text, '[...]' =
# a bracketed note (unsupported/unparseable format), or the extracted
# text. Implemented by the Import Engine (imports/) — one Importer class
# per format behind a registry, replacing what used to be an if/elif
# chain of _extract_pdf/_extract_docx/etc. functions in this file.
from imports import extract_text


def pdf_page_images(path, max_pages=6, zoom=2.0):
    """Rasterise the first pages of a PDF to PNG data-URLs (for scanned PDFs
    with no text layer, so the vision model can still read them)."""
    import fitz  # PyMuPDF

    urls = []
    doc = fitz.open(path)
    try:
        for i, page in enumerate(doc):
            if i >= max_pages:
                break
            pix = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom))
            b64 = base64.b64encode(pix.tobytes("png")).decode()
            urls.append(f"data:image/png;base64,{b64}")
    finally:
        doc.close()
    return urls


def chunk_text(text, size=1500, overlap=200):
    """Legacy plain-text chunker kept for non-document paths (pptx, xlsx,
    zip members, raw text).  Returns list[str] with no locators."""
    text = text.strip()
    if not text:
        return []
    chunks, i = [], 0
    while i < len(text):
        chunks.append(text[i : i + size])
        i += size - overlap
    return chunks[:400]


import re as _re

_PAGE_RE = _re.compile(r"\x00PAGE(\d+)\x00")
_SEC_RE = _re.compile(r"\x00SEC\d+:([^\x00]*)\x00")


def chunk_document(text, size=1500, overlap=200):
    """Sentinel-aware chunker for PDFs and DOCX files.

    Returns list[dict] with keys:
      content  – clean text, no sentinels
      page     – int | None  (1-based page number for PDFs)
      section  – str | None  (heading text for DOCX)

    Algorithm
    ---------
    1. Walk the text character-by-character tracking the *current* page and
       section from any sentinel we pass.
    2. Strip all sentinels from the content so the model never sees them.
    3. Slide a window of `size` characters (with `overlap`) over the clean
       text, attaching the page/section that was active when that slice
       started.
    """
    if not text:
        return []

    # ── pass 1: build a sentinel-stripped clean text plus a position map ──
    # pos_map[clean_pos] = (page, section) at the time we wrote that char
    clean_chars = []
    pos_meta = []  # parallel list: meta at each clean char position
    cur_page = None
    cur_section = None
    idx = 0
    n = len(text)

    while idx < n:
        # Check for sentinel starting here
        page_m = _PAGE_RE.match(text, idx)
        sec_m = _SEC_RE.match(text, idx)
        if page_m:
            cur_page = int(page_m.group(1))
            idx = page_m.end()
            continue
        if sec_m:
            cur_section = sec_m.group(1).strip()
            idx = sec_m.end()
            continue
        ch = text[idx]
        clean_chars.append(ch)
        pos_meta.append((cur_page, cur_section))
        idx += 1

    clean = "".join(clean_chars).strip()
    if not clean:
        return []

    # ── pass 2: sliding window over clean text ──
    chunks = []
    i = 0
    total = len(clean)
    while i < total:
        end = min(i + size, total)
        page, section = pos_meta[i] if i < len(pos_meta) else (None, None)
        chunks.append(
            {
                "content": clean[i:end],
                "page": page,
                "section": section,
            }
        )
        i += size - overlap

    return chunks[:400]  # safety cap


def embed_texts(texts, user_id=None):
    """Returns list of embeddings or None per text (None = embedding failed).

    `user_id`: when given, logs token usage to ai_usage_ledger (kind=
    "embedding") — optional so existing call sites that don't have a
    user_id handy (or don't need cost tracking) are unaffected."""
    try:
        out = []
        total_tokens = 0
        for i in range(0, len(texts), 64):
            resp = client.embeddings.create(model=EMBED_MODEL, input=texts[i : i + 64])
            out.extend([d.embedding for d in resp.data])
            total_tokens += getattr(resp.usage, "prompt_tokens", 0) or 0
        record_ai_call(EMBED_MODEL, prompt_tokens=total_tokens)
        if user_id:
            _log_ai_usage(user_id, "embedding", "embed_model", total_tokens, 0)
        return out
    except Exception:
        return [None] * len(texts)


def cosine(a, b):
    dot = sum(x * y for x, y in zip(a, b))
    na = math.sqrt(sum(x * x for x in a)) or 1e-9
    nb = math.sqrt(sum(x * x for x in b)) or 1e-9
    return dot / (na * nb)


def keyword_score(query, text):
    q = set(w for w in query.lower().split() if len(w) > 3)
    if not q:
        return 0.0
    t = text.lower()
    return sum(1 for w in q if w in t) / len(q)


def rag_retrieve(user_id, conversation_id, project_id, query, top_k=6, file_id=None):
    """Top chunks from files scoped to this conversation or its project.

    When file_id is given the retrieval is hard-scoped to that single file —
    used by Paper Chat (M7) so the AI draws only from the paper being discussed.

    Delegates to ``backend.research.research_retrieve`` (W2 unified spine).
    Returns prompt-oriented dicts (legacy shape + file_id/chunk_id).
    """
    from backend.research import ResearchScope, research_retrieve

    db = SessionLocal()
    try:
        scope = ResearchScope.for_chat(
            user_id=user_id,
            conversation_id=conversation_id,
            project_id=project_id,
            file_id=file_id,
        )
        hits = research_retrieve(
            db,
            UserFile=UserFile,
            Chunk=Chunk,
            select=select,
            scope=scope,
            query=query,
            embed_texts=embed_texts,
            top_k=top_k,
            LibraryCollectionPaper=LibraryCollectionPaper,
        )
        return [h.to_prompt_dict() for h in hits]
    finally:
        db.close()


def _research_passages_for_chat(
    user_id,
    conversation_id,
    project_id,
    query,
    file_id=None,
    collection_id=None,
    search_mode="off",
    top_k=6,
):
    """W1/W2: PassageHit list + scope for Trust Chat citations."""
    from backend.research import ResearchScope, research_retrieve

    db = SessionLocal()
    try:
        scope = ResearchScope.for_chat(
            user_id=user_id,
            conversation_id=conversation_id,
            project_id=project_id,
            file_id=file_id,
            search_mode=search_mode if not file_id else "off",
            collection_id=collection_id,
        )
        hits = research_retrieve(
            db,
            UserFile=UserFile,
            Chunk=Chunk,
            select=select,
            scope=scope,
            query=query,
            embed_texts=embed_texts,
            top_k=top_k,
            LibraryCollectionPaper=LibraryCollectionPaper,
        )
        return hits, scope
    finally:
        db.close()
# ------------------------------------------------------------------ web search
def web_search(query, max_results=5):
    try:
        from ddgs import DDGS
    except ImportError:
        from duckduckgo_search import DDGS
    out = []
    try:
        with DDGS() as ddgs:
            for r in ddgs.text(query, max_results=max_results):
                out.append(
                    {
                        "title": r.get("title", ""),
                        "url": r.get("href") or r.get("url", ""),
                        "snippet": r.get("body", ""),
                    }
                )
    except Exception as e:
        out.append({"title": "search error", "url": "", "snippet": str(e)})
    return out


TOOL_WEB_SEARCH = {
    "type": "function",
    "name": "web_search",
    "description": (
        "Search the web for current or factual information "
        "(news, papers, prices, dates, anything after your "
        "knowledge cutoff or that you are unsure about)."
    ),
    "parameters": {
        "type": "object",
        "properties": {"query": {"type": "string"}},
        "required": ["query"],
    },
}
TOOL_SAVE_CITATION = {
    "type": "function",
    "name": "save_citation",
    "description": (
        "Save an academic reference to the user's citation "
        "manager. Use when the user asks to save/cite a paper, "
        "or when they clearly want to keep a reference found "
        "via search."
    ),
    "parameters": {
        "type": "object",
        "properties": {
            "authors": {"type": "string", "description": "e.g. Smith, J.; Doe, A."},
            "title": {"type": "string"},
            "year": {"type": "string"},
            "venue": {"type": "string", "description": "journal/conference"},
            "doi": {"type": "string"},
            "url": {"type": "string"},
        },
        "required": ["title"],
    },
}


# ------------------------------------------------------------------ memory (selective)
MEMORY_PROMPT = """You maintain long-term memory about a user of a research \
assistant. From the conversation below, extract NEW durable facts worth \
remembering long-term.

WORTH remembering: thesis/research topic and field, methodology, preferred \
citation style, programming languages/tools, preferred tone or writing \
style, name/role/institution, long-term goals.
NOT worth remembering: one-off requests, temporary questions, trivia they \
asked about, anything about the assistant's own answers.

Do NOT repeat facts already known.
Already known: {known}

Conversation:
{transcript}

Reply ONLY with JSON: {{"facts": ["fact 1", ...]}} — empty list if nothing \
is worth remembering (this is common and fine)."""


def responses_text(prompt, json_mode=False, kind=None, user_id=None):
    """`kind`/`user_id`: when both are given, logs token usage to
    ai_usage_ledger. Optional — most of this function's call sites don't
    pass them yet (chat, memory extraction, titles, compare, gap-finder,
    writing assistant); only extract_metadata and paper_analysis do
    today. Not logging is the safe default, not an error."""
    kwargs = dict(model=UTILITY_MODEL, input=prompt, store=False)
    if json_mode:
        kwargs["text"] = {"format": {"type": "json_object"}}
    resp = client.responses.create(**kwargs)
    usage = getattr(resp, "usage", None)
    record_ai_call(
        UTILITY_MODEL,
        prompt_tokens=getattr(usage, "input_tokens", 0) or 0,
        completion_tokens=getattr(usage, "output_tokens", 0) or 0,
    )
    if kind and user_id:
        _log_ai_usage(
            user_id,
            kind,
            "utility_model",
            getattr(usage, "input_tokens", 0) or 0,
            getattr(usage, "output_tokens", 0) or 0,
        )
    return resp.output_text


# Sprint B/C — project research + memory promotion (needs responses_text above).
memory_promotion_service = create_memory_promotion_service(
    SessionLocal=SessionLocal,
    select=select,
    Project=Project,
    UserFile=UserFile,
    Memory=Memory,
    DerivedAnalysis=DerivedAnalysis,
)
project_research_service = create_project_research_service(
    SessionLocal=SessionLocal,
    select=select,
    Project=Project,
    UserFile=UserFile,
    PaperAnalysis=PaperAnalysis,
    DerivedAnalysis=DerivedAnalysis,
    AnalysisPipelineResult=AnalysisPipelineResult,
    get_prompt_builder=get_prompt_builder,
    responses_text=responses_text,
    utility_model=UTILITY_MODEL,
    build_phase1_prompt_context=build_phase1_prompt_context,
    memory_promotion_service=memory_promotion_service,
    ai_gate=ai_gate,
    cost_ledger=_cost_ledger,
    events=_ops_events,
    max_active_research=int(os.environ.get("MAX_ACTIVE_RESEARCH", "5")),
)
app.register_blueprint(
    create_projects_blueprint(
        project_service=project_service,
        project_research_service=project_research_service,
        memory_promotion_service=memory_promotion_service,
        login_required=login_required,
        limiter=limiter,
    )
)


# ══════════════════════════════════════════════════════════════════════════
# RESEARCH WORKSPACE — Milestone 3: automatic paper metadata extraction
# ══════════════════════════════════════════════════════════════════════════

import hashlib as _hashlib
import re as _re_meta

# The prompt is deliberately strict: return ONLY the JSON object, with null
# for any field that cannot be found in the text.  We do NOT ask it to
# invent, guess, or web-search — only extract from what is present.
_META_PROMPT = """You are a metadata extractor for academic papers.

Given the first portion of a research document, extract the following fields
exactly as they appear.  Return ONLY a JSON object — no markdown, no prose.
Use null for any field you cannot find with high confidence.

Fields:
  title       – full paper title (string | null)
  authors     – semicolon-separated author names, "Last, F." style (string | null)
  year        – 4-digit publication year (string | null)
  venue       – journal, conference, or publisher name (string | null)
  doi         – DOI string without "https://doi.org/" prefix (string | null)
  abstract    – full abstract text verbatim (string | null)
  keywords    – comma-separated keywords if listed (string | null)

Document excerpt (first 3 000 chars):
{excerpt}
"""


def _sha256(text: str) -> str:
    return _hashlib.sha256(text.encode("utf-8", errors="replace")).hexdigest()


def _extract_meta_from_text(text: str, user_id=None) -> dict:
    """Call the utility model to extract bibliographic metadata from `text`.

    Returns a dict with keys: title, authors, year, venue, doi, abstract,
    keywords — all strings or None.  Never raises: on any failure returns an
    empty dict so the caller can degrade gracefully."""
    excerpt = text[:3000]
    try:
        raw = responses_text(
            _META_PROMPT.format(excerpt=excerpt),
            json_mode=True,
            kind="metadata",
            user_id=user_id,
        )
        data = json.loads(raw)
    except Exception:
        return {}

    clean = {}
    for key in ("title", "authors", "year", "venue", "doi", "abstract", "keywords"):
        val = data.get(key)
        clean[key] = str(val).strip() if val and str(val).strip() not in ("null", "None", "") else None
    return clean


def _apply_metadata(file_id: int, text: str, content_hash: str, job_id=None) -> None:
    """DEPRECATED: prefer AnalysisPipelineService / phase1_analysis job.

    Background task: extract metadata and write it to the UserFile row.

    Runs in a daemon thread so the upload HTTP response is already sent by
    the time this does its model call.  It is idempotent: if the content_hash
    already matches the stored one we skip the model call entirely.

    `job_id`: pass an already-claimed UploadJob id (the queue worker does)
    to skip creating/finishing a second, duplicate tracking row — the
    worker owns that row's lifecycle instead. Left None (default) for the
    legacy thread-spawned callers, which still manage their own row here."""
    from backend.analysis_pipeline.deprecation import warn_legacy

    warn_legacy("server._apply_metadata")
    db = SessionLocal()
    owns_job = job_id is None
    try:
        uf = db.get(UserFile, file_id)
        if not uf:
            return

        # Idempotency: if a previous run already processed this exact content,
        # do nothing.  Covers the case where the same paper is re-uploaded.
        if uf.content_hash == content_hash and uf.meta_status == "done":
            return

        if owns_job:
            job_id = _start_upload_job(db, uf.user_id, file_id, "extract_metadata")
        uf.meta_status = "running"
        db.commit()

        meta = _extract_meta_from_text(text, user_id=uf.user_id)

        # Re-fetch in case another thread touched the row while we were waiting
        uf = db.get(UserFile, file_id)
        if not uf:
            return

        uf.content_hash = content_hash
        uf.meta_status = "done"
        if meta.get("title"):
            uf.title = meta["title"][:500]
        if meta.get("authors"):
            uf.authors = meta["authors"][:1000]
        if meta.get("year"):
            # Validate: keep only if it looks like a 4-digit year
            y = _re_meta.search(r"(19|20)\d{2}", meta["year"] or "")
            if y:
                uf.year = y.group(0)
        if meta.get("venue"):
            uf.venue = meta["venue"][:300]
        if meta.get("doi"):
            uf.doi = meta["doi"][:200]
        if meta.get("abstract"):
            uf.abstract = meta["abstract"][:8000]
        db.commit()
        if owns_job:
            _finish_upload_job(db, job_id, ok=True)

    except Exception as exc:
        try:
            uf2 = db.get(UserFile, file_id)
            if uf2:
                uf2.meta_status = "failed"
                db.commit()
            if owns_job:
                _finish_upload_job(db, job_id, ok=False, error=exc)
        except Exception:
            pass
        if not owns_job:
            raise  # let the queue worker's own try/except apply retry/backoff
        logging.getLogger(__name__).warning("metadata extraction failed for file %s: %s", file_id, exc)
    finally:
        db.close()


def extract_metadata(file_id: int, text: str, content_hash: str) -> None:
    """DEPRECATED: enqueue phase1_analysis (no daemon threads).

    Bibliographic fields come from Phase 1.1 via the worker chain
    ``import → phase1_analysis → paper_analysis``. Kept as a named
    entry point so any leftover callers still hit the queue.
    """
    from backend.analysis_pipeline.deprecation import warn_legacy

    warn_legacy("server.extract_metadata")
    db = SessionLocal()
    try:
        uf = db.get(UserFile, file_id)
        if not uf:
            return
        if content_hash and not uf.content_hash:
            uf.content_hash = content_hash
        _enqueue_job(db, uf.user_id, file_id, "phase1_analysis")
        db.commit()
    finally:
        db.close()


def extract_metadata_sync(file_id: int, text: str, content_hash: str) -> None:
    """DEPRECATED: same as extract_metadata — enqueue phase1_analysis only."""
    extract_metadata(file_id, text, content_hash)


# ══════════════════════════════════════════════════════════════════════════
# RESEARCH WORKSPACE — Milestone 4: Automatic Paper Analysis (14 fields)
# ══════════════════════════════════════════════════════════════════════════

# Every section is a concise, specific label — the model fills in the value.
# "null" is the correct answer when a section genuinely doesn't exist in the
# paper (e.g. a theory paper has no Dataset section).
ANALYSIS_FIELDS = [
    "executive_summary",  # 3–5 sentence plain-English overview
    "abstract_explained",  # abstract rewritten for a smart non-specialist
    "research_objective",  # what the paper is trying to achieve
    "problem_statement",  # the gap or problem it addresses
    "methodology",  # how the study was conducted
    "dataset",  # data used; null if not applicable
    "experiments",  # key experiments or evaluations run
    "results",  # main quantitative/qualitative findings
    "key_contributions",  # list of specific, novel contributions
    "strengths",  # what the paper does well
    "limitations",  # weaknesses, threats to validity, open questions
    "future_work",  # directions the authors or community should pursue
    "keywords",  # 5–10 technical keywords
    "important_terms",  # glossary: {term: one-line definition}
]

_ANALYSIS_PROMPT = """You are an expert research analyst. Analyse the paper below and return ONLY a JSON object — no markdown fences, no prose outside the object.

Each key maps to the content described. Use null when a section genuinely does not apply (e.g. no dataset for a pure theory paper). Never fabricate details not present in the text.

Keys and what to put in them:
  executive_summary   – 3–5 sentences: what is this paper, why does it matter
  abstract_explained  – rewrite the abstract for a smart non-specialist
  research_objective  – one sentence: what the paper is trying to achieve
  problem_statement   – the specific gap or problem being addressed
  methodology         – how the study was conducted (approach, framework, steps)
  dataset             – datasets used, sizes, sources (null if not applicable)
  experiments         – key experiments or evaluations described
  results             – main findings; include numbers if stated
  key_contributions   – JSON array of strings, each a distinct novel contribution
  strengths           – what the paper does particularly well (array of strings)
  limitations         – weaknesses, assumptions, threats to validity (array)
  future_work         – next steps suggested by authors or implied by gaps (array)
  keywords            – 5–10 technical keywords as a JSON array
  important_terms     – JSON object {{term: one-line definition}} for key jargon

Paper text (first {max_chars} characters):
{text}
"""

_ANALYSIS_MAX_CHARS = 12_000  # covers most papers; keeps prompt cost bounded


def _run_paper_analysis(file_id: int, text: str, content_hash: str, job_id=None) -> None:
    """DEPRECATED: prefer worker phase1_analysis → paper_analysis chain.

    Background worker: generate and persist the 14-field paper analysis.

    Idempotent on content_hash: if the stored hash matches and status=='done'
    we skip the model call. 'force' refreshes bypass this check (see route).

    `job_id`: pass an already-claimed UploadJob id (the queue worker does)
    to skip creating/finishing a second, duplicate tracking row. Left
    None (default) for the legacy thread-spawned callers.
    """
    from backend.analysis_pipeline.deprecation import warn_legacy

    warn_legacy("server._run_paper_analysis")
    db = SessionLocal()
    owns_job = job_id is None
    try:
        pa = db.execute(select(PaperAnalysis).where(PaperAnalysis.file_id == file_id)).scalar_one_or_none()

        if pa is None:
            pa = PaperAnalysis(file_id=file_id, user_id=db.get(UserFile, file_id).user_id)
            db.add(pa)
            db.commit()

        # Idempotency check
        if pa.content_hash == content_hash and pa.status == "done":
            return

        if owns_job:
            job_id = _start_upload_job(db, pa.user_id, file_id, "paper_analysis")
        pa.status = "running"
        pa.error = ""
        db.commit()

        prompt = _ANALYSIS_PROMPT.format(
            max_chars=_ANALYSIS_MAX_CHARS,
            text=text[:_ANALYSIS_MAX_CHARS],
        )
        raw = responses_text(prompt, json_mode=True, kind="analysis", user_id=pa.user_id)
        data = json.loads(raw)

        # Normalise: ensure array fields are lists, terms dict is a dict
        for arr_field in (
            "key_contributions",
            "strengths",
            "limitations",
            "future_work",
            "keywords",
        ):
            v = data.get(arr_field)
            if isinstance(v, str):
                data[arr_field] = [v] if v else []
            elif not isinstance(v, list):
                data[arr_field] = []

        if not isinstance(data.get("important_terms"), dict):
            data["important_terms"] = {}

        # Re-fetch in case a concurrent request modified the row
        pa = db.execute(select(PaperAnalysis).where(PaperAnalysis.file_id == file_id)).scalar_one_or_none()
        if pa is None:
            return

        pa.status = "done"
        pa.content_hash = content_hash
        pa.model = UTILITY_MODEL
        pa.data = json.dumps(data, ensure_ascii=False)
        pa.error = ""
        db.commit()
        if owns_job:
            _finish_upload_job(db, job_id, ok=True)

    except Exception as exc:
        try:
            pa2 = db.execute(select(PaperAnalysis).where(PaperAnalysis.file_id == file_id)).scalar_one_or_none()
            if pa2:
                pa2.status = "failed"
                pa2.error = str(exc)[:500]
                db.commit()
            if owns_job:
                _finish_upload_job(db, job_id, ok=False, error=exc)
        except Exception:
            pass
        if not owns_job:
            raise  # let the queue worker's own try/except apply retry/backoff
        logging.getLogger(__name__).warning("paper analysis failed for file %s: %s", file_id, exc)
    finally:
        db.close()


def trigger_paper_analysis(file_id: int, text: str, content_hash: str, sync: bool = False) -> None:
    """DEPRECATED: enqueue paper_analysis (or run legacy inline when sync=True).

    Prefer worker ``phase1_analysis → paper_analysis``. The async path no
    longer spawns daemon threads — it writes UploadJob + OutboxEvent.
    """
    from backend.analysis_pipeline.deprecation import warn_legacy

    warn_legacy("server.trigger_paper_analysis")
    if sync:
        _run_paper_analysis(file_id, text, content_hash)
        return
    db = SessionLocal()
    try:
        uf = db.get(UserFile, file_id)
        if not uf:
            return
        if content_hash and not uf.content_hash:
            uf.content_hash = content_hash
        _enqueue_job(db, uf.user_id, file_id, "paper_analysis")
        db.commit()
    finally:
        db.close()


def _analysis_to_dict(pa: PaperAnalysis) -> dict:
    """Serialise a PaperAnalysis row to the public API shape."""
    data = {}
    if pa.data:
        try:
            data = json.loads(pa.data)
        except Exception:
            pass
    return {
        "file_id": pa.file_id,
        "status": pa.status,
        "error": pa.error or "",
        "model": pa.model or "",
        "updated_at": pa.updated_at.isoformat() if pa.updated_at else None,
        "data": data,
    }


def _file_to_dict(x: UserFile) -> dict:
    """Serialise a UserFile to the JSON shape the frontend expects.

    Centralising this means every route (upload, list, patch, …) returns
    exactly the same shape and there is one place to add fields."""
    from backend.library.readiness import readiness_payload

    n_chunks = len(x.chunks)
    payload = {
        "id": x.id,
        "name": x.name,
        "kind": x.kind,
        "size": x.size,
        "project_id": x.project_id,
        "conversation_id": x.conversation_id,
        "chunks": n_chunks,
        # ── research metadata ──
        "title": x.title or "",
        "authors": x.authors or "",
        "year": x.year or "",
        "venue": x.venue or "",
        "doi": x.doi or "",
        "abstract": x.abstract or "",
        "reading_status": x.reading_status or "unread",
        "tags": json.loads(x.tags) if x.tags else [],
        "meta_status": x.meta_status or "pending",
        "created_at": x.created_at.isoformat() if x.created_at else None,
        # ── scholarly enrichment provenance ──
        "doi_verified": bool(getattr(x, "doi_verified", False)),
        "metadata_source": getattr(x, "metadata_source", "extracted") or "extracted",
        "source_url": getattr(x, "source_url", "") or "",
        "external_provider": getattr(x, "external_provider", "") or "",
        "external_item_id": getattr(x, "external_item_id", "") or "",
        "crossref_last_synced": (
            getattr(x, "crossref_last_synced", None).isoformat()
            if getattr(x, "crossref_last_synced", None) else None
        ),
    }
    payload.update(readiness_payload(x, chunk_count=n_chunks))
    try:
        from backend.scholarly.uftr.state import fulltext_payload, lifecycle_label

        ft = fulltext_payload(x)
        if ft is not None:
            payload["fulltext"] = ft
        payload["lifecycle_label"] = lifecycle_label(x)
    except Exception:
        pass
    return payload


def extract_memories(user_id, project_id, convo_messages):
    db = SessionLocal()
    try:
        existing = [m.fact for m in db.execute(select(Memory).where(Memory.user_id == user_id)).scalars()]
        transcript = "\n".join(f"{m['role']}: {str(m['content'])[:500]}" for m in convo_messages[-10:])
        text = responses_text(
            MEMORY_PROMPT.format(known=json.dumps(existing), transcript=transcript),
            json_mode=True,
        )
        facts = json.loads(text).get("facts", [])
        for f in facts[:5]:
            if f and f not in existing:
                db.add(
                    Memory(
                        user_id=user_id,
                        project_id=project_id,
                        fact=f[:1000],
                        kind="fact",
                        source="chat",
                        status="active",
                        payload="{}",
                    )
                )
        db.commit()
    except Exception:
        db.rollback()
    finally:
        db.close()


def generate_title(first_user_msg, first_reply):
    try:
        t = responses_text(
            "Write a short title (2-5 words, no quotes, no punctuation at the "
            "end) for a chat that starts like this:\n"
            f"user: {str(first_user_msg)[:400]}\nassistant: {first_reply[:400]}"
        )
        return (t or "").strip().strip('"')[:60] or None
    except Exception:
        return None


def build_paper_chat_prompt(user, paper, now=None):
    """Focused system prompt for Paper Chat (M7).

    Canonical text lives in ``backend.ai_core.prompts.legacy_paper_chat``
    (``LEGACY_PAPER_CHAT_PROMPT_VERSION``). Stage 1 pipeline path resolves the
    same string via ``PromptRouter.route_legacy_paper_chat``. Prefer
    ``build_paper_chat_system_prompt`` from /api/chat (PromptBuilder path).
    """
    from backend.ai_core.prompts.legacy_paper_chat import render_legacy_paper_chat_prompt

    return render_legacy_paper_chat_prompt(
        user_name=user.name,
        paper_title=paper.title or paper.name,
        authors=paper.authors,
        year=paper.year,
        venue=paper.venue,
        now=now,
    )


def _paper_chat_pipeline_mode():
    from backend.ai_core.paper_chat import paper_chat_pipeline_mode

    return paper_chat_pipeline_mode()


def _paper_chat_use_prompt_builder() -> bool:
    return os.environ.get("PAPER_CHAT_USE_PROMPT_BUILDER", "true").strip().lower() in (
        "1",
        "true",
        "yes",
        "on",
    )


def _paper_chat_phase1_context_enabled() -> bool:
    """Inject persisted Phase 1 JSON as a developer message (does not replace RAG)."""
    return os.environ.get("PAPER_CHAT_PHASE1_CONTEXT", "true").strip().lower() in (
        "1",
        "true",
        "yes",
        "on",
    )


def _load_paper_phase1_context(db, file_id: int) -> str:
    """Compact Phase 1 block for paper chat, or empty if missing/failed."""
    if not file_id or not _paper_chat_phase1_context_enabled():
        return ""
    try:
        from backend.analysis_pipeline.persistence import load_analysis_result
        from backend.analysis_pipeline.summary import build_phase1_prompt_context

        result = load_analysis_result(db, AnalysisPipelineResult, file_id)
        if not result or not result.phase_results:
            return ""
        return build_phase1_prompt_context(result.phase_results) or ""
    except Exception:
        logging.getLogger(__name__).warning(
            "paper_chat phase1 context load failed file_id=%s", file_id, exc_info=True
        )
        return ""


def build_paper_chat_system_prompt(user, paper, now=None, phase1_context: str = ""):
    """Paper chat system instructions via PromptBuilder (legacy text parity).

    Returns ``(system_prompt, phase1_for_developer_message)``. Phase 1 is never
    folded into the system string so Stage 1 shadow hashes stay comparable.
    """
    if not _paper_chat_use_prompt_builder():
        return build_paper_chat_prompt(user, paper, now=now), (phase1_context or "").strip()

    db = SessionLocal()
    try:
        builder = get_prompt_builder(db)
        assembled = builder.build_paper_chat_instructions(
            user_name=user.name or "",
            paper_title=paper.title or paper.name or "",
            authors=paper.authors,
            year=paper.year,
            venue=paper.venue,
            phase1_context=phase1_context or "",
            now=now,
        )
        return assembled.final, (assembled.rag or "").strip()
    finally:
        db.close()


# Static opening sentence only — everything else build_system_prompt()
# assembles below (user name, date, custom instructions, project,
# memories) is computed per-request. Opening text lives in PromptRegistry
# ("chat_system") with CHAT_SYSTEM_FALLBACK when unseeded.
# Phase A: normal chat assembly is PromptBuilder.build_chat_instructions().
# CHAT_USE_PROMPT_BUILDER (default true) can force the legacy assembler
# for emergency rollback during soak; remove after stability is confirmed.
_CHAT_SYSTEM_FALLBACK = CHAT_SYSTEM_FALLBACK


def _chat_use_prompt_builder() -> bool:
    return os.environ.get("CHAT_USE_PROMPT_BUILDER", "true").strip().lower() in (
        "1",
        "true",
        "yes",
        "on",
    )


def _get_chat_system_opening(db):
    try:
        text, _prompt_version = get_prompt_registry(db).get_prompt("chat_system")
        return text
    except (ValueError, TemplateError):
        return _CHAT_SYSTEM_FALLBACK
    except Exception:
        logging.getLogger(__name__).warning("chat_system prompt fetch failed, using fallback", exc_info=True)
        return _CHAT_SYSTEM_FALLBACK


def _build_system_prompt_legacy(user, project, memory_enabled=True, now=None):
    """Legacy flat chat assembler (rollback when CHAT_USE_PROMPT_BUILDER=false).

    PR2: never inject another user's project instructions even if the caller
    passed a cross-owned Project row.
    """
    if project is not None and not project_owned_by_user(project, user.id):
        project = None
    global_mems, proj_mems = [], []
    db = SessionLocal()
    try:
        if memory_enabled:
            global_mems = [
                m.fact
                for m in db.execute(
                    select(Memory).where(Memory.user_id == user.id, Memory.project_id.is_(None))
                ).scalars()
            ]
            if project:
                proj_mems = [
                    m.fact
                    for m in db.execute(
                        select(Memory).where(Memory.user_id == user.id, Memory.project_id == project.id)
                    ).scalars()
                ]
        opening = _get_chat_system_opening(db)
    finally:
        db.close()

    when = now or datetime.now()
    parts = [
        opening,
        f"The user's name is {user.name}.",
        f"Current date/time: {when.strftime('%Y-%m-%d %H:%M')}.",
    ]
    if user.custom_instructions:
        parts.append("The user's custom instructions (always follow):\n" + user.custom_instructions)
    if project:
        parts.append(f'Current project: "{project.name}".')
        if project.instructions:
            parts.append("Project instructions from the user:\n" + project.instructions)
    if global_mems:
        parts.append("Things you remember about the user:\n" + "\n".join(f"- {m}" for m in global_mems))
    if proj_mems:
        parts.append("Things you remember in this project:\n" + "\n".join(f"- {m}" for m in proj_mems))
    return "\n\n".join(parts)


def build_system_prompt(user, project, memory_enabled=True):
    """Normal-chat system instructions. Default path: PromptBuilder chat parity."""
    if not _chat_use_prompt_builder():
        return _build_system_prompt_legacy(user, project, memory_enabled=memory_enabled)

    db = SessionLocal()
    try:
        builder = get_prompt_builder(db)
        assembled = builder.build_chat_instructions(
            user_id=user.id,
            user_name=user.name or "",
            custom_instructions=user.custom_instructions or "",
            project_id=project.id if project else None,
            memory_enabled=memory_enabled,
        )
        return assembled.final
    finally:
        db.close()


def preview_chat_prompt_builder_migration(user, project, memory_enabled=True):
    """Compare legacy flat assembler vs PromptBuilder.build_chat_instructions()."""
    now = datetime.now()
    legacy = _build_system_prompt_legacy(user, project, memory_enabled=memory_enabled, now=now)

    db = SessionLocal()
    try:
        builder = get_prompt_builder(db)
        assembled = builder.build_chat_instructions(
            user_id=user.id,
            user_name=user.name or "",
            custom_instructions=user.custom_instructions or "",
            project_id=project.id if project else None,
            memory_enabled=memory_enabled,
            now=now,
        )
    finally:
        db.close()

    return {
        "legacy_system_prompt": legacy,
        "prompt_builder_final": assembled.final,
        "prompt_builder_system": assembled.system,
        "prompt_builder_task": assembled.task,
        "prompt_builder_project_context": assembled.project_context,
        "prompt_builder_memory": assembled.memory,
        "parity_match": legacy == assembled.final,
    }


def _log_chat_cost(user_id, model, usage):
    """Best-effort — a logging failure must never break an otherwise-
    successful chat response, same reasoning as every other best-effort
    cost-logging call site in this app. /api/chat calls
    client.responses.create() directly (never went through
    responses_text()), so unlike extract_metadata/paper_analysis this
    route had NO cost logging at all until now — a real gap being
    closed, not a duplicate of anything."""
    if not usage:
        return
    prompt_tokens = getattr(usage, "input_tokens", 0) or 0
    completion_tokens = getattr(usage, "output_tokens", 0) or 0
    record_ai_call(model, prompt_tokens=prompt_tokens, completion_tokens=completion_tokens)
    cost = 0.0
    try:
        ledger = get_cost_ledger()
        cost = ledger.estimate_cost(model, prompt_tokens, completion_tokens)
        db = SessionLocal()
        try:
            ledger.log(
                db,
                user_id=user_id,
                model=model,
                prompt_tokens=prompt_tokens,
                completion_tokens=completion_tokens,
                total_tokens=prompt_tokens + completion_tokens,
                cost=cost,
                action="chat",
            )
        finally:
            db.close()
    except Exception:
        logging.getLogger(__name__).warning("chat cost logging failed", exc_info=True)
    try:
        ai_gate.record_usage(
            user_id,
            tokens=prompt_tokens + completion_tokens,
            cost_usd=cost,
            operation="chat",
        )
    except Exception:
        logging.getLogger(__name__).warning("chat quota increment failed", exc_info=True)


# ------------------------------------------------------------------ API: profile / models
def _onboarding_payload(user) -> dict:
    fields_raw = getattr(user, "research_fields", None) or ""
    fields = [f for f in str(fields_raw).split(",") if f]
    return {
        "research_role": getattr(user, "research_role", None) or "",
        "research_fields": fields,
        "institution": getattr(user, "institution", None) or "",
        "research_goal": getattr(user, "research_goal", None) or "",
        "experience_level": getattr(user, "experience_level", None) or "",
        # Convenience aliases for launchpad copy
        "research_focus": getattr(user, "institution", None) or "",
        "goal": getattr(user, "research_goal", None) or "",
    }


@app.route("/api/me")
@login_required
def api_me():
    db = SessionLocal()
    try:
        user = db.get(User, session["user_id"])
        return jsonify(
            {
                "id": user.id,
                "name": user.name,
                "email": user.email,
                "picture": user.picture or "",
                "custom_instructions": user.custom_instructions or "",
                "default_model": DEFAULT_MODEL,
                "beta_mode": CLOSED_BETA,
                "auth_provider": getattr(user, "auth_provider", None) or "password",
                "has_password": bool(getattr(user, "password_hash", None)),
                "onboarding_completed": bool(getattr(user, "onboarding_completed_at", None)),
                "onboarding": _onboarding_payload(user),
                "is_admin": bool(getattr(user, "is_admin", False)),
            }
        )
    finally:
        db.close()


@app.route("/api/profile", methods=["PATCH"])
@login_required
def update_profile():
    data = request.get_json(silent=True) or {}
    db = SessionLocal()
    try:
        u = db.get(User, session["user_id"])
        if "custom_instructions" in data:
            u.custom_instructions = str(data["custom_instructions"])[:4000]
        if "name" in data:
            name = str(data["name"] or "").strip()[:200]
            if name:
                u.name = name
        db.commit()
        return jsonify({"ok": True, "name": u.name, "email": u.email})
    finally:
        db.close()


@app.route("/api/models")
@login_required
def api_models():
    force = request.args.get("refresh") == "1"
    return jsonify({"models": get_models(force=force), "default": DEFAULT_MODEL})


# ------------------------------------------------------------------ API: files
IMAGE_EXT = (".png", ".jpg", ".jpeg", ".gif", ".webp")  # vision-API formats
MAX_UPLOAD_BYTES = MAX_FILE_MB * 1024 * 1024

from backend.upload.validation import (  # noqa: E402
    ALLOWED_EXTENSIONS,
    ValidationError as UploadValidationError,
    kind_for_extension,
    validate_extension,
    validate_upload_path,
)

# ---- storage architecture: presigned/multipart upload config -------------
MULTIPART_THRESHOLD_BYTES = int(os.environ.get("MULTIPART_THRESHOLD_MB", "25")) * 1024 * 1024
UPLOAD_PART_BYTES = int(os.environ.get("UPLOAD_PART_SIZE_MB", "8")) * 1024 * 1024
UPLOAD_SESSION_TTL_SECONDS = int(os.environ.get("UPLOAD_SESSION_TTL_MINUTES", "60")) * 60


def _find_duplicate_file(db, user_id, checksum):
    """Storage-level dedup, scoped per-user (not global) — a global
    content-addressed store would need reference counting before a delete
    could ever remove the underlying object, which isn't worth it for a
    personal-scale library and would blur the per-user isolation this app
    otherwise guarantees."""
    if not checksum:
        return None
    existing = (
        db.execute(select(UserFile).where(UserFile.user_id == user_id, UserFile.checksum_sha256 == checksum))
        .scalars()
        .first()
    )
    if not existing:
        return None
    # A DB row whose object has since gone missing from storage (deleted
    # out-of-band, reconciliation hasn't caught it yet) must not be handed
    # back as if the bytes still exist.
    if storage.storage_manager.provider.head(existing.path) is None:
        return None
    return existing


def _start_upload_job(db, user_id, file_id, job_type):
    """Create an UploadJob row and return its id. This tracks today's
    still-synchronous/threading-based processing — it is the foundation a
    future queue (processing-pipeline-architecture.md) reads from, not a
    queue itself; execution here is unchanged, only observed."""
    job = UploadJob(
        user_id=user_id,
        file_id=file_id,
        job_type=job_type,
        status="running",
        started_at=datetime.now(timezone.utc),
    )
    db.add(job)
    db.commit()
    return job.id


def _finish_upload_job(db, job_id, ok, error=None):
    job = db.get(UploadJob, job_id)
    if not job:
        return
    job.status = "done" if ok else "failed"
    job.finished_at = datetime.now(timezone.utc)
    if error is not None:
        job.attempts = (job.attempts or 0) + 1
        job.last_error = str(error)[:2000]
    db.commit()


def _enqueue_job(db, user_id, file_id, job_type, upload_batch_id=None):
    """Create an UploadJob + its paired OutboxEvent in one transaction —
    the same transactional-outbox pattern upload_file() / confirm_upload()
    use, factored out so the queue worker can chain follow-on stages
    (import → phase1_analysis → paper_analysis) the same way instead of
    spawning threads. Does not commit: caller folds this into its own
    transaction."""
    from backend.jobs.outbox import enqueue_upload_job_with_outbox

    job = enqueue_upload_job_with_outbox(
        db,
        UploadJob=UploadJob,
        OutboxEvent=OutboxEvent,
        user_id=user_id,
        file_id=file_id,
        job_type=job_type,
        upload_batch_id=upload_batch_id,
    )
    return job.id


from backend.analysis_pipeline.routes import create_analysis_pipeline_blueprint

app.register_blueprint(
    create_analysis_pipeline_blueprint(
        SessionLocal=SessionLocal,
        UserFile=UserFile,
        UploadJob=UploadJob,
        OutboxEvent=OutboxEvent,
        AnalysisPipelineResult=AnalysisPipelineResult,
        enqueue_job=_enqueue_job,
        storage_backend=get_storage_backend(),
    )
)


_redis_client = None


def _get_redis():
    """Lazy singleton; returns None if Redis isn't configured or isn't
    reachable. Every caller must treat None as "cache unavailable, fall
    back to Postgres" — deliberately not memoized as permanently
    unavailable, so a Redis instance that comes up later is picked up on
    the very next call instead of staying disabled for the process
    lifetime."""
    global _redis_client
    if not REDIS_URL:
        return None
    if _redis_client is not None:
        return _redis_client
    try:
        import redis as redis_lib

        client = redis_lib.from_url(REDIS_URL, socket_connect_timeout=2, socket_timeout=2, decode_responses=True)
        client.ping()
        _redis_client = client
        return _redis_client
    except Exception:
        logging.getLogger(__name__).warning("Redis unavailable — job-status cache disabled")
        return None


def _set_job_status_cache(job_id, status=None, progress=None, updated_at=None, user_id=None, *, payload=None):
    """Cache job status hash for GET /api/jobs/<id>/status.

    Preferred path: `_cache_upload_job_status(job)` → `payload=` + `user_id`.
    Legacy positional `(job_id, status, progress, updated_at, user_id)` still works.
    """
    r = _get_redis()
    if not r:
        return
    key = f"job:{job_id}:status"
    try:
        from backend.jobs.observability import job_status_cache_mapping

        if payload is not None:
            if user_id is None:
                raise ValueError("user_id required when caching job status payload")
            mapping = job_status_cache_mapping(payload, user_id=int(user_id))
        else:
            mapping = {
                "status": str(status or ""),
                "progress": str(int(progress or 0)),
                "updated_at": updated_at.isoformat() if updated_at else "",
                "user_id": str(int(user_id)),
                "job_type": "",
                "attempts": "0",
                "last_error": "",
                "payload_json": "",
            }
        r.hset(key, mapping=mapping)
        r.expire(key, JOB_STATUS_CACHE_TTL_SECONDS)
    except Exception:
        logging.getLogger(__name__).warning("job-status cache write failed", exc_info=True)


def _get_job_status_cache(job_id):
    r = _get_redis()
    if not r:
        return None
    try:
        data = r.hgetall(f"job:{job_id}:status")
        return data or None
    except Exception:
        logging.getLogger(__name__).warning("job-status cache read failed", exc_info=True)
        return None


def _cache_upload_job_status(job):
    """Write full A-404 status payload into Redis from an UploadJob row."""
    from backend.jobs.observability import DEFAULT_MAX_ATTEMPTS, serialize_job_status

    payload = serialize_job_status(job, max_attempts=DEFAULT_MAX_ATTEMPTS, cached=False)
    _set_job_status_cache(job.id, payload=payload, user_id=job.user_id)


# Small static per-model price table (USD per 1K tokens) — update by hand
# when OpenAI changes pricing, same approach devops-observability.md §3
# recommends over integrating a billing API for this scale of app.
#
# gpt-5-mini was previously listed here with a specific rate while
# backend/ai/cost_ledger.py's CostLedger.PRICING deliberately excluded the
# whole gpt-5 family as unverified (its own docstring: "a wrong-but-
# confident-looking dollar figure is worse than an honest unknown") — two
# tables disagreeing about whether the same number was trustworthy.
# Reconciled toward the more conservative policy: removed here too, so an
# unpriced model returns cost=0.0 everywhere in this app, consistently,
# rather than a real number in one place and a deliberate "we don't know"
# in the other. Re-add with a real rate once OpenAI's published gpt-5-mini
# pricing is actually confirmed against their pricing page.
_PRICE_PER_1K_TOKENS = {
    "gpt-4o-mini": {"prompt": 0.00015, "completion": 0.0006},
    "text-embedding-3-small": {"prompt": 0.00002, "completion": 0.0},
}


def _get_active_model_version_id(logical_name):
    """Look up the seeded/active model_versions row for a logical name
    (backfill.py's Task 3 pass seeds default_model/utility_model/
    embed_model, version 1, active). Returns None if none is seeded —
    callers must treat that as "don't log this usage", not fail the
    actual AI call over missing cost-tracking metadata."""
    db = SessionLocal()
    try:
        row = (
            db.execute(
                select(ModelVersion).where(
                    ModelVersion.logical_name == logical_name,
                    ModelVersion.is_active == True,  # noqa: E712 (SQLAlchemy Column needs == True)
                )
            )
            .scalars()
            .first()
        )
        return row.id if row else None
    finally:
        db.close()


def _log_ai_usage(user_id, kind, logical_model_name, prompt_tokens, completion_tokens):
    """Best-effort: never let a cost-tracking failure break the AI call it's
    tracking. Silently does nothing if no active model_versions row is
    seeded (e.g. backfill.py hasn't been run) or the write itself fails."""
    if not user_id:
        return
    try:
        model_version_id = _get_active_model_version_id(logical_model_name)
        if not model_version_id:
            return
        cost = 0.0
        db = SessionLocal()
        try:
            mv = db.get(ModelVersion, model_version_id)
            prices = _PRICE_PER_1K_TOKENS.get(mv.provider_model_id) if mv else None
            if prices:
                cost = (prompt_tokens / 1000) * prices["prompt"] + (completion_tokens / 1000) * prices["completion"]
            db.add(
                AIUsageLedger(
                    user_id=user_id,
                    kind=kind,
                    model_version_id=model_version_id,
                    prompt_tokens=prompt_tokens,
                    completion_tokens=completion_tokens,
                    cost_usd=cost,
                )
            )
            db.commit()
        finally:
            db.close()
    except Exception:
        logging.getLogger(__name__).warning("ai usage logging failed", exc_info=True)


def _adjust_storage_usage(db, user_id, delta_bytes, delta_files):
    """Live per-user storage total, updated in the same transaction as the
    upload/delete that changes it — get-or-create, then increment. Does
    NOT commit: callers fold this into their own transaction boundary
    (some need it atomic with several other writes), and must commit
    themselves."""
    usage = db.get(StorageUsage, user_id)
    if not usage:
        usage = StorageUsage(user_id=user_id, bytes_used=0, file_count=0)
        db.add(usage)
    usage.bytes_used = max(0, (usage.bytes_used or 0) + delta_bytes)
    usage.file_count = max(0, (usage.file_count or 0) + delta_files)


def _process_document(db, uf, path, name, mime, job_id=None, on_processed=None):
    """Extract, chunk, embed, and persist Chunk rows for a document that's
    already stored. Shared by the queue worker (and any caller that already
    holds a local copy). HTTP upload routes enqueue an ``import`` job and
    never call this in-request.

    Returns the user-facing `note` (None on a normal successful index).

    `job_id`: pass an already-claimed UploadJob id (the queue worker does)
    to skip creating/finishing a second, duplicate tracking row — the
    worker owns that row's lifecycle instead.

    `on_processed(file_id, text, content_hash)`: called once real text has
    been extracted and chunked. The queue worker passes one that enqueues
    ``phase1_analysis``. When omitted, this function enqueues
    ``phase1_analysis`` itself (queue-only — no daemon threads)."""
    owns_job = job_id is None
    if owns_job:
        job_id = _start_upload_job(db, uf.user_id, uf.id, "import")
    try:
        lower = name.lower()
        text = extract_text(path, mime, name)
        is_note = bool(text) and text.startswith("[") and text.endswith("]") and len(text) < 400
        note = None
        n_chunks = 0

        if not text:
            # No readable text — e.g. a scanned/image PDF or a binary blob.
            uf.text_len = 0
            note = "scanned_pdf" if (lower.endswith(".pdf") or "pdf" in (mime or "")) else "no_text"
        elif is_note:
            uf.text_len = 0
            note = text.strip("[]")
        else:
            uf.text_len = len(text)
            # Use the locator-aware chunker for PDFs and DOCX so every chunk
            # knows its page / section; fall back to plain chunking for
            # everything else (pptx, xlsx, txt …).
            is_locatable = lower.endswith(".pdf") or "pdf" in (mime or "") or lower.endswith(".docx")
            if is_locatable:
                doc_chunks = chunk_document(text)
                pieces = [c["content"] for c in doc_chunks]
                embs = embed_texts(pieces, user_id=uf.user_id) if pieces else []
                for i, (ch_dict, e) in enumerate(zip(doc_chunks, embs)):
                    db.add(
                        Chunk(
                            file_id=uf.id,
                            idx=i,
                            content=ch_dict["content"],
                            embedding=json.dumps(e) if e else None,
                            page=ch_dict.get("page"),
                            section=ch_dict.get("section"),
                        )
                    )
            else:
                pieces = chunk_text(text)
                embs = embed_texts(pieces, user_id=uf.user_id) if pieces else []
                for i, (p, e) in enumerate(zip(pieces, embs)):
                    db.add(
                        Chunk(
                            file_id=uf.id,
                            idx=i,
                            content=p,
                            embedding=json.dumps(e) if e else None,
                        )
                    )
            n_chunks = len(pieces)
        db.commit()

        # Queue follow-up analysis — never spawn daemon threads here.
        if text and not is_note and n_chunks > 0:
            h = _sha256(text)
            # Persist hash immediately so follow-up jobs can use it for
            # their idempotency checks.
            db2 = SessionLocal()
            try:
                uf2 = db2.get(UserFile, uf.id)
                if uf2:
                    uf2.content_hash = h
                    db2.commit()
            finally:
                db2.close()
            if on_processed:
                on_processed(uf.id, text, h)
            else:
                # Fallback for any non-worker caller: same chain as worker.
                _enqueue_job(db, uf.user_id, uf.id, "phase1_analysis")
                db.commit()

        if owns_job:
            _finish_upload_job(db, job_id, ok=True)
        return note
    except Exception as exc:
        if owns_job:
            _finish_upload_job(db, job_id, ok=False, error=exc)
        else:
            raise  # let the queue worker's own try/except apply retry/backoff


# ── A-302 File Processing: direct upload + job status ─────────────────────
from backend.upload.processing_routes import create_processing_upload_blueprint

app.register_blueprint(
    create_processing_upload_blueprint(
        SessionLocal=SessionLocal,
        Project=Project,
        User=User,
        UserFile=UserFile,
        UploadBatch=UploadBatch,
        UploadJob=UploadJob,
        OutboxEvent=OutboxEvent,
        StorageUsage=StorageUsage,
        ALLOWED_EXTENSIONS=ALLOWED_EXTENSIONS,
        MAX_FILE_MB=MAX_FILE_MB,
        upload_dir=UPLOAD_DIR,
        storage=storage,
        login_required=login_required,
        limiter=limiter,
        resolve_owned_project_id=resolve_owned_project_id,
        log_security_event=log_security_event,
        find_duplicate_file=lambda *a, **k: _find_duplicate_file(*a, **k),
        file_to_dict=lambda *a, **k: _file_to_dict(*a, **k),
        validate_extension=validate_extension,
        UploadValidationError=UploadValidationError,
        validate_upload_path=lambda *a, **k: validate_upload_path(*a, **k),
        kind_for_extension=kind_for_extension,
        adjust_storage_usage=lambda *a, **k: _adjust_storage_usage(*a, **k),
        get_job_status_cache=lambda *a, **k: _get_job_status_cache(*a, **k),
        set_job_status_cache=lambda job: _cache_upload_job_status(job),
        default_storage_limit_bytes=quota_service.DEFAULT_STORAGE_LIMIT_BYTES,
    )
)


# ══════════════════════════════════════════════════════════════════════════
# Storage architecture — presigned / multipart uploads
#
# Bytes go straight from the browser to storage instead of proxying through
# this server twice (browser→Flask, then Flask→R2). Three calls:
#   1. POST /api/uploads/presign            → get a URL (or multipart part
#                                              URLs) to PUT bytes to
#   2. POST /api/uploads/multipart/complete → multipart only
#   3. POST /api/uploads/confirm            → server verifies the object
#                                              landed, then creates the
#                                              UserFile row and processes it
#
# Not yet wired into the frontend (Composer.tsx still uses the direct
# /api/files POST above) — this is the backend half of that migration.
# ══════════════════════════════════════════════════════════════════════════


from backend.upload.presign_routes import create_presign_upload_blueprint

app.register_blueprint(
    create_presign_upload_blueprint(
        SessionLocal=SessionLocal,
        Project=Project,
        UploadSession=UploadSession,
        UploadBatch=UploadBatch,
        UserFile=UserFile,
        ALLOWED_EXTENSIONS=ALLOWED_EXTENSIONS,
        MAX_UPLOAD_BYTES=MAX_UPLOAD_BYTES,
        MAX_FILE_MB=MAX_FILE_MB,
        MULTIPART_THRESHOLD_BYTES=MULTIPART_THRESHOLD_BYTES,
        UPLOAD_PART_BYTES=UPLOAD_PART_BYTES,
        UPLOAD_SESSION_TTL_SECONDS=UPLOAD_SESSION_TTL_SECONDS,
        storage=storage,
        login_required=login_required,
        limiter=limiter,
        resolve_owned_project_id=resolve_owned_project_id,
        log_security_event=log_security_event,
        find_duplicate_file=_find_duplicate_file,
        file_to_dict=_file_to_dict,
        validate_extension=validate_extension,
        UploadValidationError=UploadValidationError,
        validate_upload_path=lambda *args, **kwargs: validate_upload_path(*args, **kwargs),
        kind_for_extension=kind_for_extension,
        adjust_storage_usage=_adjust_storage_usage,
        enqueue_job=_enqueue_job,
        upload_dir=UPLOAD_DIR,
    )
)


# ------------------------------------------------------------------ storage maintenance CLI
# Run with `flask --app server sweep-temp` / `gc-storage` / `reconcile-storage`.


@app.cli.command("sweep-temp")
def sweep_temp_cmd():
    """Delete stray files left in UPLOAD_DIR by a request that crashed
    before its own cleanup ran."""
    removed = storage.sweep_temp_dir(UPLOAD_DIR, max_age_seconds=UPLOAD_SESSION_TTL_SECONDS)
    click.echo(f"sweep-temp: removed {len(removed)} stale temp file(s)")


@app.cli.command("gc-storage")
def gc_storage_cmd():
    """Delete storage objects for upload sessions that never got confirmed
    (abandoned presigned/multipart uploads)."""
    cutoff = datetime.now(timezone.utc) - timedelta(seconds=UPLOAD_SESSION_TTL_SECONDS)
    db = SessionLocal()
    try:
        stale = (
            db.execute(
                select(UploadSession).where(
                    UploadSession.status.in_(["pending", "uploaded"]),
                    UploadSession.created_at < cutoff,
                )
            )
            .scalars()
            .all()
        )
        report = storage.garbage_collect(storage.storage_manager.provider, [s.key for s in stale])
        for s in stale:
            s.status = "expired"
        db.commit()
        click.echo(
            f"gc-storage: {len(report.deleted)} deleted, "
            f"{len(report.failed)} failed, {len(stale)} session(s) expired"
        )
    finally:
        db.close()


@app.cli.command("reconcile-storage")
@click.option(
    "--apply",
    is_flag=True,
    help="Actually delete orphaned objects (default: dry-run report only).",
)
def reconcile_storage_cmd(apply):
    """Compare what's actually in storage against what the DB references."""
    db = SessionLocal()
    try:
        known_keys = {row[0] for row in db.execute(select(UserFile.path)).all() if row[0]}
        report = storage.reconcile(storage.storage_manager.provider, known_keys, dry_run=not apply)
        click.echo(f"orphaned: {len(report.orphaned_keys)}  missing: {len(report.missing_keys)}")
        for k in report.orphaned_keys:
            click.echo(f"  orphan   {k}")
        for k in report.missing_keys:
            click.echo(f"  missing  {k}")
        if apply:
            click.echo(f"deleted: {len(report.deleted)}")
    finally:
        db.close()


# ══════════════════════════════════════════════════════════════════════════
# RESEARCH WORKSPACE — Milestone 5: Knowledge Library API
# ══════════════════════════════════════════════════════════════════════════


from backend.library.library_overview_routes import create_library_overview_blueprint

app.register_blueprint(
    create_library_overview_blueprint(
        SessionLocal=SessionLocal,
        UserFile=UserFile,
        PaperAnalysis=PaperAnalysis,
        Conversation=Conversation,
        Citation=Citation,
        Project=Project,
        select_fn=select,
        login_required=login_required,
        file_to_dict=_file_to_dict,
        collection_service=_collection_service,
    )
)


from backend.library.file_detail_routes import create_file_detail_blueprint

app.register_blueprint(
    create_file_detail_blueprint(
        SessionLocal=SessionLocal,
        UserFile=UserFile,
        Project=Project,
        PaperAnalysis=PaperAnalysis,
        select_fn=select,
        login_required=login_required,
        limiter=limiter,
        storage=storage,
        extract_text=extract_text,
        sha256_fn=_sha256,
        enqueue_job=_enqueue_job,
        adjust_storage_usage=_adjust_storage_usage,
        file_to_dict=_file_to_dict,
        analysis_to_dict=_analysis_to_dict,
        app_logger=app.logger,
    )
)


# ── Scholarly provider endpoints ─────────────────────────────────────────────
from backend.library.discover_routes import create_discover_blueprint

app.register_blueprint(
    create_discover_blueprint(
        SessionLocal=SessionLocal,
        UserFile=UserFile,
        Project=Project,
        select_fn=select,
        login_required=login_required,
        file_to_dict=_file_to_dict,
        app_logger=app.logger,
        feature_flag_service=feature_flag_service,
        discover_flag=FLAG_DISCOVER_SEARCH,
        storage=storage,
        upload_dir=UPLOAD_DIR,
        enqueue_import=lambda db, uid, fid: _enqueue_job(db, uid, fid, "import"),
        max_file_mb=MAX_FILE_MB,
    )
)

from backend.projects.notes_routes import create_notes_blueprint

app.register_blueprint(
    create_notes_blueprint(
        SessionLocal=SessionLocal,
        Note=Note,
        Project=Project,
        UserFile=UserFile,
        select_fn=select,
        login_required=login_required,
        resolve_owned_project_id=resolve_owned_project_id,
        log_security_event=log_security_event,
    )
)


def _doc_hash(content: str) -> str:
    return _hashlib.sha256((content or "").encode("utf-8")).hexdigest()


def _word_count(content: str) -> int:
    return len([w for w in (content or "").split() if w.strip()])


def _writing_doc_to_dict(d: WritingDocument) -> dict:
    return {
        "id": d.id,
        "title": d.title or "",
        "content": d.content or "",
        "project_id": d.project_id,
        "editor_kind": d.editor_kind or "markdown",
        "status": d.status or "draft",
        "current_version": int(d.current_version or 1),
        "last_opened_at": d.last_opened_at.isoformat() if d.last_opened_at else None,
        "word_count": int(d.word_count or 0),
        "created_at": d.created_at.isoformat() if d.created_at else None,
        "updated_at": d.updated_at.isoformat() if d.updated_at else None,
    }


def _writing_doc_version_to_dict(v: WritingDocumentVersion) -> dict:
    return {
        "id": v.id,
        "document_id": v.document_id,
        "version_no": int(v.version_no or 0),
        "title": v.title or "",
        "source": v.source or "save",
        "created_at": v.created_at.isoformat() if v.created_at else None,
    }


def _log_document_activity(
    db,
    *,
    uid: int,
    document_id: int,
    action: str,
    meta: dict | None = None,
) -> None:
    db.add(
        WritingDocumentActivity(
            user_id=uid,
            document_id=document_id,
            action=action,
            meta_json=json.dumps(meta or {}),
        )
    )


def _append_document_version(
    db,
    *,
    uid: int,
    doc: WritingDocument,
    source: str,
) -> WritingDocumentVersion:
    content = doc.content or ""
    version = WritingDocumentVersion(
        document_id=doc.id,
        user_id=uid,
        version_no=int(doc.current_version or 1),
        title=(doc.title or "")[:300],
        content=content,
        content_hash=_doc_hash(content),
        source=source,
    )
    db.add(version)
    return version


def _apply_writing_status_transition(doc: WritingDocument, target_status: str) -> None:
    current = (doc.status or "draft").strip().lower()
    target = (target_status or "").strip().lower()
    if not target or target == current:
        return
    ensure_transition_allowed(current, target)
    doc.status = target


def _emit_writing_observability(event_name: str, *, uid: int, doc: WritingDocument, metadata: dict | None = None) -> None:
    payload = metadata or {}
    publish_writing_event(
        make_writing_event(
            event_name,
            user_id=uid,
            document_id=doc.id,
            metadata=payload,
        )
    )
    log_writing_metric(
        event_name,
        user_id=uid,
        document_id=doc.id,
        status=doc.status,
        current_version=int(doc.current_version or 1),
    )


from backend.writing.api.document_routes import create_writing_document_blueprint

app.register_blueprint(
    create_writing_document_blueprint(
        SessionLocal=SessionLocal,
        WritingDocument=WritingDocument,
        WritingDocumentVersion=WritingDocumentVersion,
        Project=Project,
        select_fn=select,
        login_required=login_required,
        limiter=limiter,
        WritingDomainError=WritingDomainError,
        normalize_status_filter=normalize_status_filter,
        require_owned_project=require_owned_project,
        writing_doc_to_dict=_writing_doc_to_dict,
        normalize_document_mutation=normalize_document_mutation,
        normalize_editor_kind=normalize_editor_kind,
        resolve_owned_project_id=resolve_owned_project_id,
        log_security_event=log_security_event,
        doc_hash_fn=_doc_hash,
        word_count_fn=_word_count,
        append_document_version=_append_document_version,
        log_document_activity=_log_document_activity,
        emit_writing_observability=_emit_writing_observability,
        require_owned_document=require_owned_document,
        build_version_conflict_payload=build_version_conflict_payload,
        apply_writing_status_transition=_apply_writing_status_transition,
        next_version_number=next_version_number,
        normalize_idempotency_key=normalize_idempotency_key,
        is_idempotent_replay=is_idempotent_replay,
        writing_doc_version_to_dict=_writing_doc_version_to_dict,
    )
)


from backend.analysis_pipeline.persistence import load_analysis_result as _load_analysis_result
from backend.evidence.api.routes import create_evidence_blueprint

app.register_blueprint(
    create_evidence_blueprint(
        SessionLocal=SessionLocal,
        Project=Project,
        UserFile=UserFile,
        WritingDocument=WritingDocument,
        EvidenceObject=EvidenceObject,
        ClaimReview=ClaimReview,
        ResearchDecision=ResearchDecision,
        WritingSentenceBinding=WritingSentenceBinding,
        EvidenceExtractionRun=EvidenceExtractionRun,
        ReviewerRun=ReviewerRun,
        ReviewerFinding=ReviewerFinding,
        AnalysisPipelineResult=AnalysisPipelineResult,
        UploadJob=UploadJob,
        OutboxEvent=OutboxEvent,
        select=select,
        login_required=login_required,
        limiter=limiter,
        load_analysis_result=_load_analysis_result,
        enqueue_job=_enqueue_job,
        ai_gateway=ai_model_gateway,
        get_model_registry=get_model_registry,
        PaperAnalysis=PaperAnalysis,
        WorkflowEvent=WorkflowEvent,
        ai_gate=ai_gate,
        feature_flag_service=feature_flag_service,
        writing_intelligence_flag=FLAG_WRITING_INTELLIGENCE,
    )
)

from backend.research.routes import create_research_blueprint

app.register_blueprint(
    create_research_blueprint(
        SessionLocal=SessionLocal,
        Project=Project,
        UserFile=UserFile,
        WritingDocument=WritingDocument,
        EvidenceObject=EvidenceObject,
        WritingSentenceBinding=WritingSentenceBinding,
        ReviewerRun=ReviewerRun,
        ReviewerFinding=ReviewerFinding,
        AnalysisPipelineResult=AnalysisPipelineResult,
        UploadJob=UploadJob,
        OutboxEvent=OutboxEvent,
        PaperAnalysis=PaperAnalysis,
        select=select,
        login_required=login_required,
        limiter=limiter,
        load_analysis_result=_load_analysis_result,
        enqueue_job=_enqueue_job,
        ai_gateway=ai_model_gateway,
        get_model_registry=get_model_registry,
    )
)

from backend.workflow.routes import create_workflow_blueprint

app.register_blueprint(
    create_workflow_blueprint(
        SessionLocal=SessionLocal,
        Project=Project,
        WorkflowEvent=WorkflowEvent,
        select=select,
        login_required=login_required,
        limiter=limiter,
    )
)


# ── A-304 Metadata: Citation Manager ──────────────────────────────────────
from backend.library.citation_routes import create_citation_blueprint

app.register_blueprint(
    create_citation_blueprint(
        SessionLocal=SessionLocal,
        Citation=Citation,
        Project=Project,
        UserFile=UserFile,
        select_fn=select,
        login_required=login_required,
        resolve_owned_project_id=resolve_owned_project_id,
        log_security_event=log_security_event,
        app_logger=app.logger,
        EvidenceObject=EvidenceObject,
    )
)


# ------------------------------------------------------------------ API: projects (CRUD)
from backend.projects.crud_routes import create_project_crud_blueprint

app.register_blueprint(
    create_project_crud_blueprint(
        SessionLocal=SessionLocal,
        Project=Project,
        Conversation=Conversation,
        Memory=Memory,
        select_fn=select,
        login_required=login_required,
        project_service=project_service,
    )
)


# ------------------------------------------------------------------ API: conversations (Phase 4)
from backend.chat.conversation_routes import create_conversation_blueprint

app.register_blueprint(
    create_conversation_blueprint(
        SessionLocal=SessionLocal,
        Conversation=Conversation,
        Project=Project,
        UserFile=UserFile,
        select_fn=select,
        login_required=login_required,
        resolve_owned_project_id=resolve_owned_project_id,
        log_security_event=log_security_event,
        get_models=lambda force=False: get_models(force=force),
        default_model=DEFAULT_MODEL,
        remove_file_row=lambda db, uf: _remove_file_row(db, uf),
    )
)


def _remove_file_row(db, uf):
    """Delete a UserFile completely: its chunks, its R2 object, and the row."""
    storage.delete(uf.path)
    db.execute(delete(Chunk).where(Chunk.file_id == uf.id))
    db.delete(uf)


# ------------------------------------------------------------------ export
def _role_label(role):
    return {
        "user": "You",
        "assistant": "Assistant",
        "developer": "System",
        "system": "System",
    }.get(role, role.title())


def _collect_export(db, uid, conversation_id=None):
    from backend.research import normalize_sources_for_api

    q = select(Conversation).where(Conversation.user_id == uid)
    if conversation_id:
        q = q.where(Conversation.id == conversation_id)
    convos = db.execute(q.order_by(Conversation.created_at)).scalars().all()
    data = []
    for c in convos:
        msgs = [
            {
                "role": m.role,
                "content": m.content,
                "created_at": m.created_at.isoformat() if m.created_at else None,
                "attachments": json.loads(m.attachments) if m.attachments else [],
                **normalize_sources_for_api(m.sources),
            }
            for m in c.messages
        ]
        data.append(
            {
                "id": c.id,
                "title": c.title or "Untitled",
                "model": c.model,
                "created_at": c.created_at.isoformat() if c.created_at else None,
                "messages": msgs,
            }
        )
    cites = db.execute(select(Citation).where(Citation.user_id == uid).order_by(Citation.created_at)).scalars().all()
    cite_list = [
        {
            "authors": ct.authors,
            "title": ct.title,
            "year": ct.year,
            "venue": ct.venue,
            "doi": ct.doi,
            "url": ct.url,
        }
        for ct in cites
    ]
    return data, cite_list


def _fmt_cite(ct):
    bits = [
        b
        for b in [
            ct.get("authors"),
            f"({ct['year']})" if ct.get("year") else "",
            ct.get("title"),
            ct.get("venue"),
            ct.get("doi"),
            ct.get("url"),
        ]
        if b
    ]
    return ". ".join(bits)


def _export_markdown(data, cites, user, plain=False):
    b = "" if plain else "**"
    stamp = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    lines = [
        f"{'' if plain else '# '}Chat export — {user['name']}",
        f"Exported {stamp}",
        "",
    ]
    for c in data:
        lines += [
            "",
            f"{'' if plain else '## '}{c['title']}",
            f"{c['created_at']} · {c['model']}",
            "",
        ]
        for m in c["messages"]:
            if m["role"] == "developer":
                continue
            who = _role_label(m["role"])
            att = ""
            if m["attachments"]:
                att = " (attached: " + ", ".join(a.get("name", "") for a in m["attachments"]) + ")"
            lines += [f"{b}{who}{att}:{b}", "", m["content"], ""]
    if cites:
        lines += ["", f"{'' if plain else '## '}Citations", ""]
        lines += [("- " if not plain else "• ") + _fmt_cite(ct) for ct in cites]
    return "\n".join(lines)


def _export_docx(data, cites, user):
    import docx

    d = docx.Document()
    d.add_heading(f"Chat export — {user['name']}", 0)
    d.add_paragraph(datetime.now(timezone.utc).strftime("Exported %Y-%m-%d %H:%M UTC"))
    for c in data:
        d.add_heading(c["title"], level=1)
        meta = d.add_paragraph(f"{c['created_at']} · {c['model']}")
        meta.runs[0].italic = True
        for m in c["messages"]:
            if m["role"] == "developer":
                continue
            p = d.add_paragraph()
            r = p.add_run(f"{_role_label(m['role'])}: ")
            r.bold = True
            p.add_run(m["content"])
            if m["attachments"]:
                names = ", ".join(a.get("name", "") for a in m["attachments"])
                a = d.add_paragraph(f"Attached: {names}")
                a.runs[0].italic = True
    if cites:
        d.add_heading("Citations", level=1)
        for ct in cites:
            d.add_paragraph(_fmt_cite(ct), style="List Bullet")
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def _text_to_pdf(title, body):
    """Simple word-wrapped, paginated PDF built with PyMuPDF (no extra deps)."""
    import fitz

    W, H, margin, fs, lh = 595.0, 842.0, 50.0, 10.0, 14.0  # A4
    font = fitz.Font("helv")
    max_w = W - 2 * margin
    doc = fitz.open()

    def wrap(line):
        if not line:
            return [""]
        words, out, cur = line.split(" "), [], ""
        for w in words:
            trial = (cur + " " + w).strip()
            if font.text_length(trial, fs) <= max_w or not cur:
                cur = trial
            else:
                out.append(cur)
                cur = w
        out.append(cur)
        return out

    rows = [(title, 15.0)]
    rows.append(("", fs))
    for raw in body.split("\n"):
        for wl in wrap(raw.replace("\t", "    ")):
            rows.append((wl, fs))

    page = doc.new_page(width=W, height=H)
    y = margin
    for txt, size in rows:
        step = 22.0 if size > 12 else lh
        if y + step > H - margin:
            page = doc.new_page(width=W, height=H)
            y = margin
        try:
            page.insert_text((margin, y), txt, fontsize=size, fontname="helv")
        except Exception:
            page.insert_text(
                (margin, y),
                txt.encode("latin-1", "replace").decode("latin-1"),
                fontsize=size,
                fontname="helv",
            )
        y += step
    out = doc.tobytes()
    doc.close()
    return out


@app.route("/api/export")
@login_required
@limiter.limit("60 per hour")
def export_data():
    fmt = (request.args.get("format") or "json").lower()
    conversation_id = request.args.get("conversation_id", type=int)
    uid = session["user_id"]
    db = SessionLocal()
    try:
        user = db.get(User, uid)
        if conversation_id:
            c = db.get(Conversation, conversation_id)
            if not c or c.user_id != uid:
                return jsonify({"error": "not_found"}), 404
        uinfo = {"name": user.name, "email": user.email}
        data, cites = _collect_export(db, uid, conversation_id)
    finally:
        db.close()

    stamp = datetime.now(timezone.utc).strftime("%Y%m%d-%H%M")
    base = f"dhund-export-{stamp}"
    title = f"Chat export — {uinfo['name']}"

    if fmt == "json":
        payload = json.dumps(
            {
                "user": uinfo,
                "exported_at": datetime.now(timezone.utc).isoformat(),
                "conversations": data,
                "citations": cites,
            },
            ensure_ascii=False,
            indent=2,
        ).encode("utf-8")
        return _download(payload, base + ".json", "application/json")
    if fmt in ("md", "markdown"):
        return _download(
            _export_markdown(data, cites, uinfo).encode("utf-8"),
            base + ".md",
            "text/markdown",
        )
    if fmt == "txt":
        return _download(
            _export_markdown(data, cites, uinfo, plain=True).encode("utf-8"),
            base + ".txt",
            "text/plain",
        )
    if fmt == "docx":
        return _download(
            _export_docx(data, cites, uinfo),
            base + ".docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    if fmt == "pdf":
        body = _export_markdown(data, cites, uinfo, plain=True)
        return _download(_text_to_pdf(title, body), base + ".pdf", "application/pdf")
    return (
        jsonify({"error": "bad_format", "detail": "format must be json|md|txt|docx|pdf"}),
        400,
    )


def _download(payload, filename, mimetype):
    return send_file(
        io.BytesIO(payload),
        mimetype=mimetype,
        as_attachment=True,
        download_name=filename,
    )


# ------------------------------------------------------------------ delete account
@app.route("/api/account", methods=["DELETE"])
@login_required
@limiter.limit("5 per hour")
def delete_account():
    """Permanently delete the signed-in account — requires step-up reauth (#16).

    Body (JSON):
      - confirm: must be the literal string "DELETE"
      - password: required when the account has a password
      - email: required (must match) when the account has no password (OAuth/magic)
    """
    from security.ops.step_up import STEP_UP_ERROR_DETAIL, authorize_account_delete
    from werkzeug.security import check_password_hash

    uid = session["user_id"]
    body = request.get_json(silent=True) or {}
    db = SessionLocal()
    try:
        user = db.get(User, uid)
        if not user:
            return jsonify({"error": "not_found"}), 404

        has_password = bool(getattr(user, "password_hash", None))
        password_matches = None
        if has_password:
            password_matches = check_password_hash(
                user.password_hash, str(body.get("password") or "")
            )

        ok, reason = authorize_account_delete(
            has_password=has_password,
            user_email=user.email or "",
            body=body,
            password_matches=password_matches,
        )
        if not ok:
            status = 403 if reason == "wrong_password" else 400
            log_security_event(
                "account_delete_denied",
                user=uid,
                reason=reason,
            )
            return (
                jsonify(
                    {
                        "error": reason,
                        "detail": STEP_UP_ERROR_DETAIL.get(reason, reason),
                    }
                ),
                status,
            )

        # Files (chunks + on-disk blobs)
        for uf in db.execute(select(UserFile).where(UserFile.user_id == uid)).scalars().all():
            _remove_file_row(db, uf)
        # Conversations (messages cascade), memories, citations, projects
        for conv in db.execute(select(Conversation).where(Conversation.user_id == uid)).scalars().all():
            db.delete(conv)
        db.execute(delete(Memory).where(Memory.user_id == uid))
        db.execute(delete(Citation).where(Citation.user_id == uid))
        db.execute(delete(Project).where(Project.user_id == uid))
        db.delete(user)
        db.commit()
    finally:
        db.close()
    log_security_event("account_deleted", user=uid)
    session.clear()
    return jsonify({"ok": True})


# ------------------------------------------------------------------ support
SUPPORT_CATEGORIES = {"general", "bug", "feature", "account", "beta"}


def _valid_email(e):
    import re

    return bool(re.match(r"^[^@\s]+@[^@\s]+\.[^@\s]+$", e or ""))


def _support_ack_html(ticket_id, subject, message):
    return (
        f"<p>Thanks for reaching out to Dhund. We've logged your "
        f"message and will get back to you soon.</p>"
        f"<p><b>Ticket:</b> #{ticket_id}<br><b>Subject:</b> "
        f"{subject or '(none)'}</p><hr>"
        f"<p style='white-space:pre-wrap;color:#555'>{message}</p>"
    )


def _support_notify_html(ticket_id, email, category, subject, message):
    return (
        f"<p>New support request <b>#{ticket_id}</b></p>"
        f"<p><b>From:</b> {email}<br><b>Category:</b> {category}<br>"
        f"<b>Subject:</b> {subject or '(none)'}</p><hr>"
        f"<p style='white-space:pre-wrap'>{message}</p>"
    )


@app.route("/api/support", methods=["POST"])
@limiter.limit("6 per hour;30 per day")
def submit_support():
    """Public contact/support endpoint (works logged-in or anonymous)."""
    data = request.get_json(silent=True) or {}
    uid = session.get("user_id")
    email = (data.get("email") or session.get("user_email") or "").strip()
    subject = (data.get("subject") or "").strip()[:300]
    category = (data.get("category") or "general").strip().lower()
    message = (data.get("message") or "").strip()
    if category not in SUPPORT_CATEGORIES:
        category = "general"
    if not _valid_email(email):
        return (
            jsonify({"error": "invalid_email", "detail": "A valid email is required."}),
            400,
        )
    if len(message) < 5:
        return (
            jsonify({"error": "empty_message", "detail": "Please describe your issue."}),
            400,
        )
    message = message[:5000]

    db = SessionLocal()
    try:
        sr = SupportRequest(
            user_id=uid,
            email=email,
            subject=subject,
            category=category,
            message=message,
        )
        db.add(sr)
        db.commit()
        ticket_id = sr.id
    finally:
        db.close()

    email_service.send(
        email,
        f"We received your message (#{ticket_id})",
        _support_ack_html(ticket_id, subject, message),
        reply_to=SUPPORT_EMAIL or None,
    )
    if SUPPORT_EMAIL:
        email_service.send(
            SUPPORT_EMAIL,
            f"[{category}] {subject or 'New support request'} (#{ticket_id})",
            _support_notify_html(ticket_id, email, category, subject, message),
            reply_to=email,
        )
    log_security_event("support_submitted", ticket=ticket_id, category=category)
    return jsonify({"ok": True, "ticket": ticket_id})


from backend.chat.memory_routes import create_memory_blueprint

app.register_blueprint(
    create_memory_blueprint(
        SessionLocal=SessionLocal,
        Memory=Memory,
        select_fn=select,
        login_required=login_required,
    )
)


def sse(event, data):
    return f"event: {event}\ndata: {json.dumps(data)}\n\n"


def image_data_url(path, mime):
    with open(path, "rb") as f:
        b64 = base64.b64encode(f.read()).decode()
    return f"data:{mime or 'image/png'};base64,{b64}"


def run_tool(name, args, user_id, project_id):
    if name == "web_search":
        results = web_search(args.get("query", ""))
        return json.dumps(results, ensure_ascii=False), results
    if name == "save_citation":
        db = SessionLocal()
        try:
            c = Citation(
                user_id=user_id,
                project_id=project_id,
                authors=str(args.get("authors", ""))[:500],
                title=str(args.get("title", ""))[:500],
                year=str(args.get("year", ""))[:10],
                venue=str(args.get("venue", ""))[:300],
                doi=str(args.get("doi", ""))[:200],
                url=str(args.get("url", ""))[:600],
            )
            db.add(c)
            db.commit()
            return json.dumps({"saved": True, "id": c.id}), []
        finally:
            db.close()
    return json.dumps({"error": f"unknown tool {name}"}), []


@app.route("/api/chat", methods=["POST"])
@login_required
@limiter.limit("20 per minute")
def chat():
    data = request.get_json(silent=True) or {}
    cid = data.get("conversation_id")
    user_message = (data.get("message") or "").strip()
    attachment_ids = data.get("attachments") or []
    regenerate = bool(data.get("regenerate"))
    search_mode = data.get("search", "auto")
    research_skill_raw = data.get("skill") or data.get("research_skill") or "ask"
    user_id = session["user_id"]

    # Phase 3 / F4.1 — hard cap before token estimates / model calls.
    MAX_CHAT_MESSAGE_CHARS = 32_000
    if len(user_message) > MAX_CHAT_MESSAGE_CHARS:
        return (
            jsonify(
                {
                    "error": "message_too_long",
                    "message": f"Message must be at most {MAX_CHAT_MESSAGE_CHARS} characters",
                }
            ),
            400,
        )

    # Unified AI gate: kill switch, verification, daily budget, quotas
    from security.ops.gate import AiAccessDenied

    try:
        ai_gate.preflight(
            user_id,
            token_estimate=estimate_chat_tokens(user_message),
            cost_estimate=0.01,
            operation="chat",
        )
    except AiAccessDenied as exc:
        body = {"error": exc.code, "detail": exc.message}
        if getattr(exc, "payload", None):
            body["quota"] = exc.payload
        return jsonify(body), exc.http_status

    db = SessionLocal()
    try:
        user = db.get(User, user_id)
        convo = db.get(Conversation, cid) if cid else None
        if not convo or convo.user_id != user_id:
            return jsonify({"error": "conversation_not_found"}), 404
        model = data.get("model") if data.get("model") in get_models() else convo.model
        project = db.get(Project, convo.project_id) if convo.project_id else None
        if project is not None and not project_owned_by_user(project, user_id):
            log_security_event(
                "authz_denied",
                resource="project",
                action="chat",
                user_id=user_id,
                project_id=convo.project_id,
                conversation_id=convo.id,
            )
            _ops_events.record(
                "authz_denied",
                user_id=user_id,
                resource="project",
                action="chat",
                project_id=convo.project_id,
            )
            project = None

        atts = []
        for fid in attachment_ids[:8]:
            x = db.get(UserFile, fid)
            if x and x.user_id == user_id:
                x.conversation_id = convo.id
                if convo.project_id and not x.project_id:
                    x.project_id = convo.project_id
                atts.append(
                    {
                        "id": x.id,
                        "name": x.name,
                        "mime": x.mime,
                        "kind": x.kind,
                        "path": x.path,
                        "text_len": x.text_len,
                    }
                )

        if regenerate:
            msgs = list(convo.messages)
            if msgs and msgs[-1].role == "assistant":
                db.delete(msgs[-1])
                db.commit()
        else:
            if not user_message and not atts:
                return jsonify({"error": "empty_message"}), 400
            db.add(
                Message(
                    conversation_id=convo.id,
                    role="user",
                    content=user_message or "(see attached files)",
                    attachments=(
                        json.dumps([{k: a[k] for k in ("id", "name", "mime", "kind")} for a in atts]) if atts else None
                    ),
                )
            )
            convo.model = model
            db.commit()

        raw_msgs = db.get(Conversation, convo.id).messages
        history = []
        for m in raw_msgs:
            content = m.content
            m_atts = json.loads(m.attachments) if m.attachments else []
            if m_atts:
                names = ", ".join(a["name"] for a in m_atts)
                content = f"[attached files: {names}]\n{content}"
            history.append({"role": m.role, "content": content})
        memory_enabled = bool(convo.memory_enabled) if convo.memory_enabled is not None else True
        temperature = convo.temperature
        reasoning_effort = convo.reasoning_effort
        paper_file_id = convo.file_id  # M7: paper chat scope (may be None)
        paper_plan = None
        paper_pipeline_mode = "false"
        paper_phase1_context = ""

        # M7: if this is a paper chat, use a focused system prompt and
        # hard-scope retrieval to the single paper.
        # PromptBuilder assembles the legacy M7 system text (parity).
        # Stage 1: optional ai_core pipeline (PAPER_CHAT_PIPELINE_ENABLED).
        # Soak-safe: any plan failure falls back to legacy so shadow/true
        # never take Paper Chat down. Phase 1 JSON is injected as a
        # developer message (does not alter Stage 1 system-prompt hashes).
        if paper_file_id:
            paper = db.get(UserFile, paper_file_id)
            if paper and paper.user_id == user_id:
                paper_phase1_context = _load_paper_phase1_context(db, paper_file_id)
                try:
                    from backend.ai_core.paper_chat import resolve_paper_chat_system_prompt

                    legacy_base, paper_phase1_context = build_paper_chat_system_prompt(
                        user, paper, phase1_context=paper_phase1_context
                    )
                    system_prompt, paper_plan, paper_pipeline_mode = resolve_paper_chat_system_prompt(
                        user_name=user.name,
                        paper_title=paper.title or paper.name,
                        authors=paper.authors,
                        year=paper.year,
                        venue=paper.venue,
                        file_id=paper_file_id,
                        project_id=convo.project_id,
                        question=(user_message or "")[:500] or None,
                        legacy_prompt=legacy_base,
                    )
                except Exception:
                    logging.getLogger(__name__).exception(
                        "paper_chat_stage1_plan_failed; falling back to legacy"
                    )
                    system_prompt, paper_phase1_context = build_paper_chat_system_prompt(
                        user, paper, phase1_context=paper_phase1_context
                    )
                    paper_plan = None
                    paper_pipeline_mode = "false"
            else:
                paper_file_id = None  # safety: invalid file, fall back
                system_prompt = build_system_prompt(user, project, memory_enabled)
        else:
            system_prompt = build_system_prompt(user, project, memory_enabled)

        convo_id = convo.id
        needs_title = not convo.title_generated
        project_id = convo.project_id
        project_name = (project.name if project is not None else None) or None
        research_skill_for_scope = research_skill_raw
    finally:
        db.close()

    # Stage 1 executor — only used when pipeline mode is ``true`` for paper chat.
    paper_executor = None
    if paper_plan is not None and paper_pipeline_mode == "true":
        from backend.ai_core.orchestration import AIExecutor, OpenAIResponsesStreamClient

        paper_executor = AIExecutor(
            stream_client=OpenAIResponsesStreamClient(client),
            default_model=model,
        )
    def generate():
        input_items = list(history)
        sources = []
        full_text = ""
        workspace_references = []
        research_scope = None
        passages = []
        skill = None
        stage1_started = time.perf_counter() if paper_executor is not None else None
        try:
            last_query = user_message or (history[-1]["content"] if history else "")

            # Prompt Gateway — Research Scope (ADR-0017). Soft redirect/clarify
            # without calling the LLM for asks that don't advance research.
            from backend.ai.research_scope import evaluate_research_scope

            scope_decision = evaluate_research_scope(
                last_query,
                project_name=project_name,
                paper_scoped=bool(paper_file_id),
                research_skill=research_skill_for_scope,
            )
            if scope_decision.blocks_llm and scope_decision.user_message:
                full_text = scope_decision.user_message
                yield sse("status", {"text": "Research scope check…"})
                # Stream as a normal assistant turn so the UI stays consistent.
                chunk_size = 48
                for i in range(0, len(full_text), chunk_size):
                    yield sse("delta", {"text": full_text[i : i + chunk_size]})
                gate_blob = json.dumps(scope_decision.to_gate_dict())
                dbi = SessionLocal()
                try:
                    dbi.add(
                        Message(
                            conversation_id=convo_id,
                            role="assistant",
                            content=full_text,
                            sources=gate_blob,
                        )
                    )
                    c2 = dbi.get(Conversation, convo_id)
                    if c2:
                        c2.updated_at = datetime.now(timezone.utc)
                    dbi.commit()
                finally:
                    dbi.close()
                yield sse(
                    "done",
                    {
                        "sources": [],
                        "references": [],
                        "scope": None,
                        "confidence": None,
                        "warnings": [],
                        "skill": research_skill_for_scope or "ask",
                        "scope_gate": scope_decision.to_gate_dict().get("scope_gate"),
                    },
                )
                return

            doc_atts = [a for a in atts if a["kind"] == "document"]
            img_atts = [a for a in atts if a["kind"] == "image"]

            # The user turn we're answering — multimodal images attach here.
            user_idx = max(
                (i for i, m in enumerate(input_items) if m["role"] == "user"),
                default=None,
            )

            # Optional Phase 1 structured context (before RAG) — grounds
            # answers without replacing excerpt retrieval.
            if paper_file_id and paper_phase1_context:
                input_items.append(
                    {
                        "role": "developer",
                        "content": (
                            "Structured analysis of this paper (Phase 1 pipeline). "
                            "Use it to orient your answer (domain, study design, "
                            "entities, evidence). Still ground specific claims in "
                            "the retrieved excerpts that follow — do not invent "
                            "details that are only implied by this summary.\n\n"
                            + paper_phase1_context
                        ),
                    }
                )

            # M7 / W1–W4: research spine + skills + grounding
            from backend.research import (
                dump_message_sources,
                get_skill,
                passages_to_workspace_references,
                verify_chat_grounding,
            )

            skill = get_skill(research_skill_raw)
            passages, research_scope = _research_passages_for_chat(
                user_id,
                convo_id,
                project_id,
                last_query[:500],
                file_id=paper_file_id,
                search_mode=search_mode,
                top_k=skill.top_k,
            )
            excerpts = [p.to_prompt_dict() for p in passages]
            workspace_references = passages_to_workspace_references(
                passages,
                primary_file_id=paper_file_id,
            )
            if skill.instruction:
                input_items.append(
                    {
                        "role": "developer",
                        "content": skill.instruction,
                    }
                )
            # W5 — prefer typed medical_understanding table for extract skill
            if skill.id == "extract":
                try:
                    from backend.analysis_pipeline.persistence import load_analysis_result
                    from backend.research.structured_extract import (
                        build_structured_extract_table,
                        table_prompt_block,
                    )

                    papers = []
                    file_ids = []
                    if paper_file_id:
                        file_ids = [int(paper_file_id)]
                    elif project_id:
                        with SessionLocal() as _db:
                            file_ids = [
                                int(x)
                                for x in _db.execute(
                                    select(UserFile.id).where(
                                        UserFile.user_id == user_id,
                                        UserFile.project_id == int(project_id),
                                    )
                                )
                                .scalars()
                                .all()
                            ][:20]
                    if file_ids:
                        with SessionLocal() as _db:
                            for fid in file_ids:
                                uf = _db.get(UserFile, fid)
                                if not uf or int(uf.user_id) != int(user_id):
                                    continue
                                analysis = load_analysis_result(
                                    _db, AnalysisPipelineResult, fid
                                )
                                medical = None
                                if analysis and isinstance(analysis.phase_results, dict):
                                    medical = analysis.phase_results.get(
                                        "medical_understanding"
                                    )
                                papers.append(
                                    {
                                        "file_id": fid,
                                        "paper_title": (uf.title or uf.name or f"#{fid}"),
                                        "paper_year": getattr(uf, "year", None) or "",
                                        "medical": medical
                                        if isinstance(medical, dict)
                                        else None,
                                    }
                                )
                        if papers:
                            table = build_structured_extract_table(
                                project_id=project_id,
                                papers=papers,
                            )
                            block = table_prompt_block(table)
                            if block:
                                input_items.append(
                                    {"role": "developer", "content": block}
                                )
                except Exception:
                    pass
            if excerpts:
                yield sse("status", {"text": "Reading your documents…"})
                input_items.append(
                    {
                        "role": "developer",
                        "content": (
                            "Relevant excerpts from the user's uploaded documents.\n"
                            "Each excerpt may include 'page' (1-based PDF page) and/or "
                            "'section' (heading the excerpt falls under), plus file_id. "
                            "When citing, be specific: prefer 'p. 4, §Methodology' over "
                            "just the filename. If no locator is present, cite by filename.\n"
                            "Only claim what these excerpts support — do not invent sources.\n"
                            + json.dumps(excerpts, ensure_ascii=False)
                        ),
                    }
                )

            # Deliver each attached document's content to the model. Small docs
            # go in whole; large ones get a generous head (RAG excerpts above
            # cover the rest); scanned PDFs are rasterised for the vision model;
            # unparseable files get an explicit, honest note (never silent).
            INLINE_DOC_CHARS = 30000
            vision_urls = []
            for a in doc_atts:
                with storage.local_copy(a["path"]) as local_path:
                    txt = extract_text(local_path, a["mime"], a["name"])
                    has_text = bool(txt) and not (txt.startswith("[") and txt.endswith("]") and len(txt) < 400)
                    if has_text:
                        body = (
                            txt
                            if len(txt) <= INLINE_DOC_CHARS
                            else (
                                txt[:INLINE_DOC_CHARS] + f"\n\n[…truncated {len(txt) - INLINE_DOC_CHARS} more "
                                "characters; see the excerpts above for the rest.]"
                            )
                        )
                        input_items.append(
                            {
                                "role": "developer",
                                "content": f"Full text of attached file " f"'{a['name']}':\n{body}",
                            }
                        )
                        continue
                    is_pdf = a["name"].lower().endswith(".pdf") or "pdf" in (a["mime"] or "")
                    pages = []
                    if is_pdf:
                        try:
                            pages = pdf_page_images(local_path)
                        except Exception:
                            pages = []
                if pages:
                    vision_urls.extend(pages)
                    input_items.append(
                        {
                            "role": "developer",
                            "content": f"The attached PDF '{a['name']}' has no text "
                            f"layer (scanned). Its first {len(pages)} "
                            "page image(s) are attached to the user's "
                            "message — read them to answer.",
                        }
                    )
                else:
                    reason = (
                        txt.strip("[]")
                        if txt
                        else "no readable text could be extracted (it may be "
                        "empty, encrypted, or an unsupported binary "
                        "format)"
                    )
                    input_items.append(
                        {
                            "role": "developer",
                            "content": f"The attached file '{a['name']}' could not "
                            f"be parsed: {reason}. Do not fabricate its "
                            "contents — tell the user what happened and "
                            "suggest a fix (re-export as PDF/DOCX/XLSX or "
                            "paste the text).",
                        }
                    )

            for a in img_atts:
                try:
                    with storage.local_copy(a["path"]) as local_path:
                        vision_urls.append(image_data_url(local_path, a["mime"]))
                except Exception:
                    pass

            if vision_urls and user_idx is not None and not regenerate:
                base = input_items[user_idx]
                text_part = base["content"] if isinstance(base["content"], str) else ""
                content = [{"type": "input_text", "text": text_part}]
                for url in vision_urls[:16]:
                    content.append({"type": "input_image", "image_url": url})
                input_items[user_idx] = {"role": "user", "content": content}

            if search_mode == "on":
                yield sse("status", {"text": "Searching the web…"})
                results = web_search(last_query[:300])
                sources.extend(results)
                input_items.append(
                    {
                        "role": "developer",
                        "content": "Web search results (cite these):\n" + json.dumps(results, ensure_ascii=False),
                    }
                )

            # Paper chat: disable all external search so the AI cannot pull
            # content from outside the uploaded document.
            if paper_file_id:
                tools = [TOOL_SAVE_CITATION]
            else:
                tools = [TOOL_SAVE_CITATION]
                if search_mode == "auto":
                    tools.append(TOOL_WEB_SEARCH)

            for _round in range(4):
                use_stream = supports_streaming(model)
                kwargs = dict(
                    model=model,
                    instructions=system_prompt,
                    input=input_items,
                    store=False,
                    tools=tools,
                )
                if use_stream:
                    kwargs["stream"] = True
                if temperature is not None and supports_temperature(model):
                    kwargs["temperature"] = temperature
                effort = normalize_reasoning_effort(model, reasoning_effort)
                if effort and supports_reasoning_effort(model):
                    kwargs["reasoning"] = {"effort": effort}

                final = None
                # Stage 1 (flag true): model invocation via AIExecutor — no
                # direct responses.create on this path.
                if paper_executor is not None and paper_plan is not None:
                    stream_kwargs = {}
                    if temperature is not None and supports_temperature(model):
                        stream_kwargs["temperature"] = temperature
                    reasoning = None
                    if effort and supports_reasoning_effort(model):
                        reasoning = {"effort": effort}
                    if not use_stream:
                        yield sse(
                            "status",
                            {
                                "text": "Running Pro model — this can take a few minutes…",
                            },
                        )
                    for event in paper_executor.stream_round(
                        paper_plan,
                        input_items=input_items,
                        tools=tools,
                        model=model,
                        reasoning=reasoning,
                        **stream_kwargs,
                    ):
                        et = event.type
                        if et == "response.output_text.delta":
                            full_text += event.delta
                            yield sse("delta", {"text": event.delta})
                        elif et == "response.completed":
                            final = event.response
                        elif et == "response.failed":
                            raise RuntimeError(event.error_message or "response failed")
                else:
                    if not use_stream:
                        yield sse(
                            "status",
                            {
                                "text": "Running Pro model — this can take a few minutes…",
                            },
                        )
                        final = client.responses.create(**kwargs)
                        chunk = _responses_output_text(final)
                        if chunk:
                            full_text += chunk
                            yield sse("delta", {"text": chunk})
                    else:
                        stream = client.responses.create(**kwargs)
                        for event in stream:
                            et = getattr(event, "type", "")
                            if et == "response.output_text.delta":
                                full_text += event.delta
                                yield sse("delta", {"text": event.delta})
                            elif et == "response.completed":
                                final = event.response
                            elif et == "response.failed":
                                raise RuntimeError(
                                    getattr(
                                        getattr(event.response, "error", None),
                                        "message",
                                        "response failed",
                                    )
                                )

                _log_chat_cost(user_id, model, getattr(final, "usage", None))

                calls = [it for it in (final.output if final else []) if getattr(it, "type", "") == "function_call"]
                if calls:
                    for c in calls:
                        input_items.append(
                            {
                                "type": "function_call",
                                "call_id": c.call_id,
                                "name": c.name,
                                "arguments": c.arguments,
                            }
                        )
                        try:
                            args = json.loads(c.arguments)
                        except Exception:
                            args = {}
                        label = (
                            "Searching: " + args.get("query", "") + "…"
                            if c.name == "web_search"
                            else "Saving citation…"
                        )
                        yield sse("status", {"text": label})
                        output, src = run_tool(c.name, args, user_id, project_id)
                        sources.extend(src)
                        input_items.append(
                            {
                                "type": "function_call_output",
                                "call_id": c.call_id,
                                "output": output,
                            }
                        )
                    continue
                break

            if paper_executor is not None and paper_plan is not None:
                from backend.ai_core.paper_chat import log_stage1_execution
                from backend.ai_core.schemas.execution import TokenUsage

                usage_obj = getattr(final, "usage", None) if final else None
                usage_dict = None
                if usage_obj is not None:
                    usage_dict = {
                        "input_tokens": getattr(usage_obj, "input_tokens", None),
                        "output_tokens": getattr(usage_obj, "output_tokens", None),
                        "total_tokens": getattr(usage_obj, "total_tokens", None),
                        "prompt_tokens": getattr(usage_obj, "prompt_tokens", None),
                        "completion_tokens": getattr(usage_obj, "completion_tokens", None),
                    }
                latency_ms = (
                    int((time.perf_counter() - stage1_started) * 1000) if stage1_started is not None else 0
                )
                exec_result = paper_executor.observe_answer(
                    paper_plan,
                    full_text,
                    model=model,
                    usage=TokenUsage.from_openai(usage_dict),
                    latency_ms=latency_ms,
                    rag_excerpt_count=len(excerpts) if excerpts else 0,
                )
                log_stage1_execution(exec_result)

            new_title = None
            from backend.research import (
                dump_message_sources as _dump_sources,
                get_skill as _get_skill_fallback,
                verify_chat_grounding as _verify_grounding,
            )

            _skill = skill or _get_skill_fallback("ask")
            grounding_report = _verify_grounding(
                full_text,
                passages,
                skill=_skill.id,
            )
            grounding_dict = grounding_report.to_dict()
            grounding_dict["warnings"] = list(grounding_report.warnings)

            dbi = SessionLocal()
            try:
                sources_blob = _dump_sources(
                    web=sources,
                    references=workspace_references,
                    scope=research_scope.to_dict() if research_scope else None,
                    grounding=grounding_dict,
                )
                dbi.add(
                    Message(
                        conversation_id=convo_id,
                        role="assistant",
                        content=full_text,
                        sources=sources_blob,
                    )
                )
                c2 = dbi.get(Conversation, convo_id)
                c2.updated_at = datetime.now(timezone.utc)
                if needs_title and full_text:
                    new_title = generate_title(last_query, full_text)
                    if new_title:
                        c2.title = new_title
                        c2.title_generated = 1
                dbi.commit()
            finally:
                dbi.close()

            done_payload = {
                "sources": sources,
                "references": workspace_references,
                "scope": research_scope.to_dict() if research_scope else None,
                "confidence": grounding_report.confidence,
                "warnings": list(grounding_report.warnings),
                "skill": _skill.id,
                "grounding": grounding_dict,
            }
            if new_title:
                done_payload["title"] = new_title
            yield sse("done", done_payload)

            if memory_enabled:
                snapshot = history + [{"role": "assistant", "content": full_text}]
                threading.Thread(
                    target=extract_memories,
                    args=(user_id, project_id, snapshot),
                    daemon=True,
                ).start()

        except Exception as e:
            msg = str(e)
            if "invalid_api_key" in msg or "Incorrect API key" in msg:
                msg = "Your OpenAI API key seems invalid — check OPENAI_API_KEY in .env."
            elif "insufficient_quota" in msg:
                msg = "Your OpenAI account is out of credit."
            elif "does not exist" in msg or "model_not_found" in msg:
                msg = f"Model '{model}' isn't available — pick another from the dropdown."
            elif "temperature" in msg.lower() and "not supported" in msg.lower():
                msg = (
                    f"Model '{model}' doesn't support temperature — "
                    "clear the temperature control or switch model."
                )
            elif "image" in msg.lower() and "support" in msg.lower():
                msg = f"Model '{model}' doesn't support images — " "switch to a vision model like gpt-4o or gpt-5."
            yield sse("error", {"text": msg})

    return Response(
        generate(),
        mimetype="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


# ══════════════════════════════════════════════════════════════════════════
# RESEARCH WORKSPACE — Milestone 14: Semantic Search (Phase 3)
# ══════════════════════════════════════════════════════════════════════════
from backend.search.semantic_routes import create_semantic_search_blueprint

app.register_blueprint(
    create_semantic_search_blueprint(
        SessionLocal=SessionLocal,
        UserFile=UserFile,
        Chunk=Chunk,
        Note=Note,
        Citation=Citation,
        Message=Message,
        Conversation=Conversation,
        select_fn=select,
        login_required=login_required,
        limiter=limiter,
        embed_texts=lambda *a, **k: embed_texts(*a, **k),
    )
)


# ══════════════════════════════════════════════════════════════════════════
# RESEARCH WORKSPACE — Multi-Paper Analysis (compare + gaps) — Phase 3
# ══════════════════════════════════════════════════════════════════════════
from backend.analysis_pipeline.derived_routes import create_derived_analysis_blueprint

app.register_blueprint(
    create_derived_analysis_blueprint(
        SessionLocal=SessionLocal,
        UserFile=UserFile,
        PaperAnalysis=PaperAnalysis,
        DerivedAnalysis=DerivedAnalysis,
        select_fn=select,
        login_required=login_required,
        limiter=limiter,
        responses_text=lambda *a, **k: responses_text(*a, **k),
        utility_model=UTILITY_MODEL,
    )
)


# ══════════════════════════════════════════════════════════════════════════
# RESEARCH WORKSPACE — Milestone 15: AI Writing Assistant + Export Center
# ══════════════════════════════════════════════════════════════════════════
from backend.writing.api.assistant_routes import create_writing_assistant_blueprint

app.register_blueprint(
    create_writing_assistant_blueprint(
        login_required=login_required,
        limiter=limiter,
        responses_text=responses_text,
    )
)


# ── Export Centre ─────────────────────────────────────────────────────────────
from backend.writing.api.export_routes import create_writing_export_blueprint

app.register_blueprint(
    create_writing_export_blueprint(
        SessionLocal=SessionLocal,
        Note=Note,
        UserFile=UserFile,
        PaperAnalysis=PaperAnalysis,
        Conversation=Conversation,
        select_fn=select,
        login_required=login_required,
        WritingDocument=WritingDocument,
        ReviewerRun=ReviewerRun,
        ReviewerFinding=ReviewerFinding,
    )
)


# ------------------------------------------------------------------ SPA (React build) serving
FRONTEND_DIST = os.path.join(os.path.dirname(os.path.abspath(__file__)), "frontend", "dist")


@app.route("/assets/<path:filename>")
def spa_assets(filename):
    # Vite content-hashes asset filenames, so they can be cached forever.
    return send_from_directory(os.path.join(FRONTEND_DIST, "assets"), filename, max_age=31536000)


@app.route("/<path:path>")
def spa(path):
    # Real API/auth/static routes are matched by their explicit rules first;
    # this catch-all only fires for unmatched paths (client-side routes). The
    # guard turns a stray /api/... typo into a 404 instead of the SPA shell.
    # Marketing pages and /login are registered as explicit routes above.
    if path.startswith(("api/", "auth/", "static/", "assets/")) or path in (
        "login",
        "logout",
        "robots.txt",
        "product",
        "how-it-works",
        "research",
        "early-access",
        "pricing",
    ):
        abort(404)
    if path.startswith("research/"):
        abort(404)
    # Serve real Vite public/ files (e.g. /brand/*.svg) before falling back
    # to the SPA shell — otherwise browsers get index.html and show broken icons.
    if path:
        candidate = os.path.normpath(os.path.join(FRONTEND_DIST, path))
        dist_root = os.path.normpath(FRONTEND_DIST)
        if candidate.startswith(dist_root + os.sep) and os.path.isfile(candidate):
            return send_from_directory(FRONTEND_DIST, path)
    index_path = os.path.join(FRONTEND_DIST, "index.html")
    if not os.path.exists(index_path):
        return ("Frontend build not found — run `npm run build` in frontend/.", 501)
    return send_from_directory(FRONTEND_DIST, "index.html")


if __name__ == "__main__":
    # Railway/Render/Fly set PORT; local default stays 5000.
    port = int(os.environ.get("PORT", "5000"))
    print(f"Dhund running -> http://0.0.0.0:{port}")
    app.run(host="0.0.0.0", port=port, debug=False, threaded=True)
