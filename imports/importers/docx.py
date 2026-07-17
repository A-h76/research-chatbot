class DocxImporter:
    extensions = (".docx",)
    supports_locators = True

    def matches(self, lower_name, mime):
        return lower_name.endswith(".docx")

    def extract(self, path, mime, name):
        """NUL-sentinel section markers from headings — \\x00SEC<n>:<heading>\\x00
        so chunk_document() can record which section each chunk belongs to."""
        import docx
        d = docx.Document(path)
        parts = []
        sec_idx = 0
        for p in d.paragraphs:
            if p.style and p.style.name.startswith("Heading") and p.text.strip():
                sec_idx += 1
                parts.append(f"\x00SEC{sec_idx}:{p.text.strip()}\x00")
            if p.text.strip():
                parts.append(p.text)
        for tbl in d.tables:
            for row in tbl.rows:
                row_text = " | ".join(c.text for c in row.cells)
                if row_text.strip():
                    parts.append(row_text)
        return "\n".join(parts).strip()
